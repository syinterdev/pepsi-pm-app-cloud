#!/usr/bin/env python3
"""Convert UAT markdown checklist to Word (.docx) for customer fill-in."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            p.add_run(strip_md_inline(part))
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [strip_md_inline(c.strip()) for c in line.split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells if c)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell_text = row[ci] if ci < len(row) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    i = 0
    table_buffer: list[list[str]] = []
    in_code = False

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        data_rows = [r for r in table_buffer if not is_separator_row(r)]
        add_table(doc, data_rows)
        table_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"
            i += 1
            continue

        if stripped.startswith("|"):
            table_buffer.append(parse_table_row(stripped))
            i += 1
            continue

        flush_table()

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            doc.add_heading(strip_md_inline(stripped[2:]), level=0)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(strip_md_inline(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(strip_md_inline(stripped[4:]), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            doc.add_heading(strip_md_inline(stripped[5:]), level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            add_rich_paragraph(doc, stripped[2:], style="Intense Quote")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.add_run(strip_md_inline(m.group(2)))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(strip_md_inline(stripped[2:]))
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main(argv: list[str]) -> int:
    base = Path(__file__).resolve().parent
    if len(argv) >= 2:
        md_path = Path(argv[1])
    else:
        md_path = base / "UAT-ROUND-2-TH.md"
    if len(argv) >= 3:
        docx_path = Path(argv[2])
    else:
        docx_path = md_path.with_suffix(".docx")
    if not md_path.is_file():
        print(f"Missing: {md_path}", file=sys.stderr)
        return 1
    convert_md_to_docx(md_path, docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
