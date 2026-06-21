#!/usr/bin/env python3
"""Generate PM Pepsi App — detailed Stress Test documentation (Word .docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "STRESS-TEST-DETAILED-TH.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "FCE4D6",
    font_size: int = 9,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main() -> None:
    today = date.today().isoformat()
    doc = Document()

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PM Pepsi App\nเอกสาร Stress Test ฉบับละเอียด\n(API Load · Performance · UAT)")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายการทดสอบโหลด API (Stress Test) ของ PM-Pepsi-App "
        "รวมสถาปัตยกรรมสคริปต์ endpoint ที่ทดสอบ พารามิเตอร์ วิธีรัน "
        "ผลการทดสอบล่าสุด และเกณฑ์ผ่านก่อน UAT / ส่งมอบลูกค้า"
    )
    doc.add_paragraph(
        "เอกสารคู่กัน: UNIT-TEST-DETAILED-TH.docx (Unit Test แยกฉบับ)"
    )

    doc.add_page_break()

    # --- 1 Overview ---
    doc.add_heading("1. ภาพรวมและวัตถุประสงค์", level=1)
    doc.add_paragraph(
        "Stress Test จำลองผู้ใช้หลายคนเรียก API พร้อมกันบน hot path ของแอป "
        "เพื่อวัดความทนทานของ backend (Express + PostgreSQL) ก่อน UAT และ production"
    )
    add_bullets(
        doc,
        [
            "วัด throughput (Requests Per Second — RPS)",
            "วัด latency (p50, p95, p99, max)",
            "ตรวจอัตรา HTTP error (4xx/5xx, timeout, connection fail)",
            "ระบุ endpoint ที่เป็นคอขวด (bottleneck)",
            "ยืนยันว่า server ไม่ crash / OOM ภายใต้โหลดที่กำหนด",
        ],
    )

    doc.add_heading("1.1 ขอบเขต", level=2)
    add_table(
        doc,
        ["รายการ", "ครอบคลุม", "ไม่ครอบคลุม"],
        [
            ["ชั้นทดสอบ", "API backend เท่านั้น", "Frontend UI, browser E2E"],
            ["Authentication", "Login จริง (JWT Bearer)", "—"],
            ["Database", "PostgreSQL จริงตาม env", "Mock DB"],
            ["Network", "LAN / localhost / production URL", "CDN, WAF ภายนอก (ยกเว้น prod)"],
        ],
    )

    doc.add_heading("1.2 ไฟล์สคริปต์", level=2)
    add_table(
        doc,
        ["ไฟล์", "บทบาท"],
        [
            ["PM-Pepsi-App/backend/scripts/stress-api.ts", "สคริปต์หลัก — weighted load test"],
            ["PM-Pepsi-App/backend/package.json", "npm script: stress:api"],
            ["from customer/tools/generate_stress_test_docx.py", "สร้างเอกสารฉบับนี้"],
        ],
    )

    doc.add_page_break()

    # --- 2 Architecture ---
    doc.add_heading("2. สถาปัตยกรรมการทดสอบ", level=1)
    doc.add_paragraph("สคริปต์ทำงานตามลำดับดังนี้:")

    add_bullets(
        doc,
        [
            "1. Login ผ่าน POST /api/v1/auth/login ได้ JWT token",
            "2. สร้าง worker จำนวน STRESS_CONCURRENCY ตัว ทำงานขนานกัน",
            "3. แต่ละ worker สุ่มเลือก scenario ตาม weight จนกว่าครบ STRESS_DURATION_SEC",
            "4. บันทึก latency, HTTP status, success/fail ทุกคำขอ",
            "5. สรุป RPS, percentile latency, แยกตาม endpoint",
            "6. Exit code 1 ถ้ามี HTTP error ใด ๆ — exit 0 ถ้าผ่านทั้งหมด",
        ],
    )

    doc.add_heading("2.1 การสุ่มโหลด (Weighted Scenario)", level=2)
    doc.add_paragraph(
        "แต่ละคำขอสุ่ม endpoint ตามน้ำหนัก (weight) เพื่อจำลองการใช้งานจริง "
        "โดย calendar/events มีน้ำหนักสูงสุดเพราะเป็นหน้าที่ช่างเปิดบ่อยที่สุด"
    )
    add_table(
        doc,
        ["Scenario name", "Weight", "สัดส่วนโดยประมาณ"],
        [
            ["health", "1", "~6%"],
            ["calendar/events", "4", "~27%"],
            ["work-orders", "3", "~20%"],
            ["planning/orders", "2", "~13%"],
            ["reports/kpi", "2", "~13%"],
            ["plan-calendar/events", "2", "~13%"],
        ],
    )

    doc.add_heading("2.2 Metrics ที่วัด", level=2)
    add_table(
        doc,
        ["Metric", "คำอธิบาย", "หน่วย"],
        [
            ["Requests", "จำนวนคำขอทั้งหมดในรอบทดสอบ", "ครั้ง"],
            ["RPS", "Requests / STRESS_DURATION_SEC", "req/s"],
            ["Error %", "คำขอที่ HTTP ไม่ ok / total", "%"],
            ["p50 / p95 / p99", "Percentile latency", "ms"],
            ["max", "Latency สูงสุด", "ms"],
            ["By endpoint", "แยก n, err, p95 ต่อ scenario", "—"],
        ],
    )

    doc.add_page_break()

    # --- 3 Endpoints ---
    doc.add_heading("3. Endpoint ที่ทดสอบ (รายละเอียด)", level=1)

    doc.add_heading("3.1 health", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "GET /api/v1/health"],
            ["Auth", "Bearer JWT"],
            ["Weight", "1"],
            ["วัตถุประสงค์", "Baseline — ตรวจว่า server ตอบสนอง"],
        ],
    )

    doc.add_heading("3.2 calendar/events", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "POST /api/v1/calendar/events"],
            ["Auth", "Bearer JWT"],
            ["Weight", "4 (สูงสุด)"],
            ["Body", "year, month, filters ว่าง (activity, wktype, status, …)"],
            ["วัตถุประสงค์", "ปฏิทิน Work scheduling — query หนักสุด"],
            ["หมายเหตุ", "มักเป็น bottleneck แรก — ดู p95 เป็นหลัก"],
        ],
    )

    doc.add_heading("3.3 work-orders", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "GET /api/v1/work-orders?limit=100"],
            ["Auth", "Bearer JWT"],
            ["Weight", "3"],
            ["วัตถุประสงค์", "รายการ Work Order หน้า WO / Confirmation"],
        ],
    )

    doc.add_heading("3.4 planning/orders", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "GET /api/v1/planning/orders?status=open"],
            ["Auth", "Bearer JWT"],
            ["Weight", "2"],
            ["วัตถุประสงค์", "แผน PM/CM สถานะ open"],
        ],
    )

    doc.add_heading("3.5 plan-calendar/events", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "GET /api/v1/plan-calendar/events?year=YYYY&month=M"],
            ["Auth", "Bearer JWT"],
            ["Weight", "2"],
            ["วัตถุประสงค์", "ปฏิทินจ่ายงาน (Plan calendar)"],
        ],
    )

    doc.add_heading("3.6 reports/kpi", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["URL", "GET /api/v1/reports/kpi?weeksBack=8"],
            ["Auth", "Bearer JWT"],
            ["Weight", "2"],
            ["วัตถุประสงค์", "KPI รายงาน Engineering"],
        ],
    )

    doc.add_page_break()

    # --- 4 Parameters ---
    doc.add_heading("4. พารามิเตอร์และสภาพแวดล้อม", level=1)

    doc.add_heading("4.1 Environment Variables", level=2)
    add_table(
        doc,
        ["ตัวแปร", "ค่าเริ่มต้น", "คำอธิบาย"],
        [
            ["STRESS_BASE_URL", "http://127.0.0.1:4000", "Base URL ของ API (ไม่มี trailing /)"],
            ["STRESS_USER", "ADMIN01", "Username สำหรับ login"],
            ["STRESS_PASSWORD", "admin", "รหัสผ่าน (dev seed)"],
            ["STRESS_CONCURRENCY", "15", "จำนวน worker พร้อมกัน (= virtual users)"],
            ["STRESS_DURATION_SEC", "30", "ระยะเวลารันแต่ละรอบ (วินาที)"],
        ],
    )

    doc.add_heading("4.2 สภาพแวดล้อมที่ต้องเตรียม", level=2)
    add_table(
        doc,
        ["รายการ", "Dev local", "Production-like"],
        [
            ["Backend", "npm run dev หรือ PM2", "PM2 + IIS reverse proxy"],
            ["PostgreSQL", "DATABASE_URL local", "DB บนเซิร์ฟเวอร์จริง"],
            ["Migrations", "ครบทุกไฟล์ใน database/migrations/", "เหมือนกัน"],
            ["User ทดสอบ", "ADMIN01 (dev seed)", "บัญชี UAT ที่ลูกค้ากำหนด"],
            ["Concurrent แนะนำ", "15–25", "10–30 ตามจำนวนช่างพร้อมกัน"],
        ],
    )

    doc.add_heading("4.3 ชุดทดสอบมาตรฐาน (Test Profiles)", level=2)
    add_table(
        doc,
        ["Profile", "CONCURRENCY", "DURATION_SEC", "ใช้เมื่อ"],
        [
            ["มาตรฐาน (Smoke load)", "15", "30", "ทุกครั้งก่อน UAT / หลัง deploy"],
            ["โหลดหนัก (Peak load)", "25", "45", "จำลองช่วง peak shift"],
            ["Soak (ถ้าต้องการ)", "15", "300", "ทดสอบ memory leak / connection pool"],
        ],
    )

    doc.add_page_break()

    # --- 5 How to run ---
    doc.add_heading("5. วิธีรัน Stress Test", level=1)

    doc.add_heading("5.1 รอบมาตรฐาน", level=2)
    add_code_block(
        doc,
        "cd PM-Pepsi-App/backend\n"
        "npm run stress:api",
    )

    doc.add_heading("5.2 รอบโหลดหนัก (PowerShell)", level=2)
    add_code_block(
        doc,
        '$env:STRESS_CONCURRENCY="25"\n'
        '$env:STRESS_DURATION_SEC="45"\n'
        'cd PM-Pepsi-App/backend\n'
        "npm run stress:api",
    )

    doc.add_heading("5.3 ชี้ไป staging / production", level=2)
    add_code_block(
        doc,
        '$env:STRESS_BASE_URL="https://pm-app.example.com"\n'
        '$env:STRESS_USER="UAT_USER"\n'
        '$env:STRESS_PASSWORD="***"\n'
        '$env:STRESS_CONCURRENCY="20"\n'
        '$env:STRESS_DURATION_SEC="60"\n'
        "npm run stress:api",
    )

    doc.add_heading("5.4 ตัวอย่าง Output ที่คาดหวัง", level=2)
    add_code_block(
        doc,
        "Stress API — http://127.0.0.1:4000\n"
        "Concurrency: 15 · Duration: 30s · User: ADMIN01\n"
        "\n"
        "--- Summary ---\n"
        "Requests: 3991 · RPS: 133.0 · Errors: 0 (0.0%)\n"
        "Latency ms — p50: 90 · p95: 241 · p99: 322 · max: 482\n"
        "\n"
        "--- By endpoint ---\n"
        "calendar/events        n=1078 · err=0 · p95=260ms · statuses=—\n"
        "...\n"
        "\n"
        "PASS — no HTTP errors under load",
    )

    doc.add_heading("5.5 การตีความผล", level=2)
    add_bullets(
        doc,
        [
            "PASS: exit code 0 และข้อความ 'PASS — no HTTP errors under load'",
            "FAIL: exit code 1 — มี HTTP error อย่างน้อย 1 คำขอ (ดู Sample failures)",
            "RPS ลดลงเมื่อ p95 สูงขึ้น — บ่งชี้ saturation",
            "calendar/events p95 สูงกว่า endpoint อื่นเป็นเรื่องปกติ",
        ],
    )

    doc.add_page_break()

    # --- 6 Results ---
    doc.add_heading("6. ผลการทดสอบล่าสุด", level=1)
    doc.add_paragraph(
        f"วันที่บันทึก: {today} · สภาพแวดล้อม: Windows Dev · API: http://127.0.0.1:4000 · User: ADMIN01"
    )
    doc.add_paragraph(
        "หมายเหตุ: ผลด้านล่างจากเครื่องพัฒนา — ก่อนส่งลูกค้าควรรันซ้ำบนเซิร์ฟเวอร์จริง "
        "และบันทึกผลในตาราง §6.3"
    )

    doc.add_heading("6.1 สรุปรวม", level=2)
    add_table(
        doc,
        ["รอบ", "Profile", "Concurrent", "ระยะเวลา", "คำขอรวม", "RPS", "Error %", "p50", "p95", "p99", "max"],
        [
            ["1", "มาตรฐาน", "15", "30s", "3,991", "133.0", "0%", "90 ms", "241 ms", "322 ms", "482 ms"],
            ["2", "โหลดหนัก", "25", "45s", "5,855", "130.1", "0%", "156 ms", "380 ms", "481 ms", "715 ms"],
        ],
    )

    doc.add_heading("6.2 แยกตาม Endpoint (รอบ 2 — โหลดหนัก)", level=2)
    add_table(
        doc,
        ["Endpoint", "จำนวนคำขอ", "Error", "p95 (ms)", "หมายเหตุ"],
        [
            ["calendar/events", "1,656", "0", "460", "Bottleneck หลัก"],
            ["work-orders", "1,293", "0", "194", "—"],
            ["planning/orders", "858", "0", "207", "—"],
            ["plan-calendar/events", "837", "0", "303", "—"],
            ["reports/kpi", "779", "0", "192", "—"],
            ["health", "432", "0", "187", "Baseline"],
        ],
    )

    doc.add_heading("6.3 บันทึกผล Production / UAT (กรอกหลังทดสอบ)", level=2)
    add_table(
        doc,
        ["วันที่", "เซิร์ฟเวอร์", "Concurrent", "Duration", "RPS", "Error %", "p95 calendar", "ผ่าน/ไม่ผ่าน", "ผู้ทดสอบ"],
        [
            ["", "", "", "", "", "", "", "☐", ""],
            ["", "", "", "", "", "", "", "☐", ""],
            ["", "", "", "", "", "", "", "☐", ""],
        ],
    )

    doc.add_page_break()

    # --- 7 Pass criteria ---
    doc.add_heading("7. เกณฑ์ผ่าน Stress Test", level=1)
    add_table(
        doc,
        ["เกณฑ์", "เป้าหมาย Dev local", "เป้าหมาย Production-like"],
        [
            ["HTTP error rate", "< 1%", "< 0.5%"],
            ["p95 calendar/events", "< 800 ms", "< 1500 ms"],
            ["p95 endpoint อื่น", "< 500 ms", "< 1000 ms"],
            ["5xx ต่อเนื่อง", "0", "0"],
            ["Server crash / OOM", "ไม่เกิด", "ไม่เกิด"],
            ["Login failure", "ไม่เกิด", "ไม่เกิด"],
        ],
    )

    doc.add_heading("7.1 สิ่งที่ทำเมื่อไม่ผ่าน", level=2)
    add_bullets(
        doc,
        [
            "ตรวจ PostgreSQL: slow query, missing index, connection pool เต็ม",
            "ตรวจ backend log (PM2 / console) หา 5xx stack trace",
            "ลด STRESS_CONCURRENCY แล้วรันซ้ำ — แยกว่าเป็น capacity หรือ bug",
            "เปรียบเทียบ p95 แต่ละ endpoint — optimize query ของ calendar/events ก่อน",
            "ตรวจ migration ครบ (SCHEMA_NOT_READY ทำให้ 503)",
        ],
    )

    doc.add_page_break()

    # --- 8 Checklist ---
    doc.add_heading("8. Checklist ก่อน UAT / ส่งมอบ", level=1)
    add_table(
        doc,
        ["ลำดับ", "รายการ", "คำสั่ง / หลักฐาน", "ผ่าน"],
        [
            ["1", "Backend รันและ health OK", "GET /api/v1/health", "☐"],
            ["2", "Migration DB ครบ", "admin/health หรือ migration log", "☐"],
            ["3", "Stress รอบมาตรฐาน", "npm run stress:api", "☐"],
            ["4", "Stress รอบโหลดหนัก", "CONCURRENCY=25 DURATION=45", "☐"],
            ["5", "บันทึกผล production (ถ้ามี)", "กรอกตาราง §6.3", "☐"],
            ["6", "Unit test ผ่าน", "ดู UNIT-TEST-DETAILED-TH.docx", "☐"],
        ],
    )

    doc.add_heading("9. บทบาทและความรับผิดชอบ", level=1)
    add_table(
        doc,
        ["บทบาท", "ความรับผิดชอบ"],
        [
            ["ทีมพัฒนา", "รัน stress หลังแก้ API หนัก; แก้ regression performance"],
            ["DevOps", "รันบน staging/production-like; monitor CPU/RAM/DB connections"],
            ["QA / UAT", "ยืนยัน concurrent ตามจำนวนผู้ใช้จริง"],
            ["ลูกค้า", "อนุมัติเกณฑ์ latency และ peak concurrent"],
        ],
    )

    doc.add_heading("10. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "docs/customer-requirements/UNIT-TEST-DETAILED-TH.docx — Unit Test (เอกสารคู่)",
            "PM-Pepsi-App/backend/scripts/stress-api.ts — ซอร์สสคริปต์",
            "docs/customer-requirements/PRE-UAT-UI-PHASES.md — เกณฑ์ UI U4",
            "docs/PRE-UAT-MASTER-PHASES.md — แผน UAT รวม",
            "from customer/tools/generate_stress_test_docx.py — สคริปต์สร้างเอกสารฉบับนี้",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
