#!/usr/bin/env python3
"""Generate PM Pepsi App — Database Design documentation (Word .docx)."""
from __future__ import annotations

import re
import sys
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# Reuse schema parser from ER diagram generator
TOOLS = Path(__file__).resolve().parent
sys.path.insert(0, str(TOOLS))
from generate_er_diagram_docx import (  # noqa: E402
    ER_IMAGE,
    MODULE_MAP,
    Table,
    add_bullets,
    add_table,
    collect_fk_relationships,
    module_for,
    parse_migrations,
)


def group_tables_by_module(tables: dict[str, Table]) -> dict[str, list[tuple[str, Table]]]:
    by_mod: dict[str, list[tuple[str, Table]]] = defaultdict(list)
    for name, tbl in sorted(tables.items()):
        if tbl.is_view:
            continue
        mod, _ = module_for(name, False)
        by_mod[mod].append((name, tbl))
    return dict(sorted(by_mod.items()))

ROOT = Path(__file__).resolve().parents[2]
MIGRATIONS = ROOT / "database" / "migrations"
OUT = ROOT / "docs" / "customer-requirements" / "DATABASE-DESIGN-DETAILED-TH.docx"


@dataclass
class DbIndex:
    name: str
    table: str
    columns: str
    definition: str


@dataclass
class DbComment:
    target: str  # table or table.column
    text: str


def parse_indexes_and_comments() -> tuple[list[DbIndex], dict[str, str]]:
    indexes: list[DbIndex] = []
    comments: dict[str, str] = {}

    for path in sorted(MIGRATIONS.glob("*.sql")):
        raw = path.read_text(encoding="utf-8", errors="replace")
        for m in re.finditer(
            r"CREATE\s+(?:UNIQUE\s+)?INDEX\s+(?:IF NOT EXISTS\s+)?(\w+)\s+ON\s+app\.(\w+)\s*\(([^)]+)\)",
            raw,
            re.IGNORECASE,
        ):
            indexes.append(
                DbIndex(
                    name=m.group(1),
                    table=m.group(2),
                    columns=m.group(3).strip(),
                    definition=m.group(0).split("\n")[0][:200],
                )
            )
        for m in re.finditer(
            r"COMMENT ON (?:TABLE|COLUMN)\s+app\.(\w+(?:\.\w+)?)\s+IS\s+'([^']*(?:''[^']*)*)'",
            raw,
            re.IGNORECASE | re.DOTALL,
        ):
            text = m.group(2).replace("''", "'")
            comments[m.group(1).lower()] = text

    return indexes, comments


SETTING_KEYS = [
    ("app.name", "branding", "ชื่อแอป"),
    ("app.logo_bytes", "branding", "โลโก้ WebP base64"),
    ("app.locale", "system", "Locale BCP-47"),
    ("app.timezone", "system", "Asia/Bangkok"),
    ("app.year_format", "system", "BE / AD"),
    ("maintenance.enabled", "system", "โหมดบำรุงรักษา"),
    ("backup.schedule_cron", "backup", "Cron สำรองอัตโนมัติ"),
    ("backup.retention_days", "backup", "เก็บ backup (วัน)"),
    ("security.max_login_attempts", "system", "จำกัด login ผิด"),
    ("app.password_min_length", "system", "ความยาวรหัสผ่านขั้นต่ำ"),
]

PERM_GROUPS = [
    ("dashboard", "หน้าแรก KPI"),
    ("calendar", "ปฏิทิน / backlog"),
    ("planning", "แผน PM/CM · assign"),
    ("work-orders", "ใบงาน WO"),
    ("confirmation", "ปิดงาน · QC · export"),
    ("iw37n", "นำเข้า IW37N"),
    ("manhours", "ชั่วโมงงาน"),
    ("personnel", "บุคลากร · confirm %"),
    ("master-data", "ข้อมูลหลัก"),
    ("reports", "รายงาน"),
    ("admin.users", "จัดการผู้ใช้"),
    ("admin.roles", "บทบาท & สิทธิ์"),
    ("admin.backup", "สำรองข้อมูล"),
    ("admin.health", "สุขภาพระบบ"),
    ("portal", "โมดูล Portal"),
]


def add_table_dict(doc: Document, tname: str, tbl: Table, comments: dict[str, str]) -> None:
    _, desc = module_for(tname, tbl.is_view)
    doc.add_heading(f"app.{tname}", level=3)
    doc.add_paragraph(desc)
    tcomment = comments.get(tname.lower())
    if tcomment:
        doc.add_paragraph(f"หมายเหตุ: {tcomment}")

    if tbl.is_view:
        if tbl.view_sql:
            p = doc.add_paragraph()
            r = p.add_run(f"Definition: {tbl.view_sql[:400]}…")
            r.font.size = Pt(8)
            r.italic = True
        return

    if not tbl.columns:
        return

    rows = []
    for cname, col in sorted(tbl.columns.items(), key=lambda x: (not x[1].is_pk, x[0])):
        flags = []
        if col.is_pk:
            flags.append("PK")
        if col.fk_ref:
            flags.append(f"FK→{col.fk_ref}")
        if not col.nullable:
            flags.append("NOT NULL")
        ckey = f"{tname}.{cname}".lower()
        note = comments.get(ckey, "")[:80]
        rows.append([
            cname,
            col.dtype,
            ", ".join(flags) or "—",
            col.default or "—",
            note or "—",
        ])
    add_table(doc, ["คอลัมน์", "ชนิดข้อมูล", "ข้อจำกัด", "Default", "คำอธิบาย"], rows, font_size=7)


def main() -> None:
    today = date.today().isoformat()
    tables = parse_migrations()
    indexes, comments = parse_indexes_and_comments()
    fk_rels = collect_fk_relationships(tables)
    table_count = sum(1 for t in tables.values() if not t.is_view)
    view_count = sum(1 for t in tables.values() if t.is_view)
    mig_count = len(list(MIGRATIONS.glob("*.sql")))

    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\nเอกสาร Database Design ฉบับละเอียด\n(PostgreSQL · Schema · Dictionary)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายการออกแบบฐานข้อมูล PM-Pepsi-App ครบวงจร "
        "รวมหลักการออกแบบ มาตรฐานชนิดข้อมูล โมเดลความปลอดภัย "
        "พจนานุกรมตาราง/คอลัมน์ Index View และแนวทาง migration — "
        "สร้างอัตโนมัติจาก database/migrations/"
    )
    doc.add_paragraph("เอกสารคู่กัน: ER-DIAGRAM-DETAILED-TH.docx (แผนภาพความสัมพันธ์)")

    doc.add_page_break()

    # 1 Overview
    doc.add_heading("1. ภาพรวมระบบฐานข้อมูล", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า / รายละเอียด"],
        [
            ["DBMS", "PostgreSQL 11+ (production) / 14+ (dev แนะนำ)"],
            ["Schema", "app (เดียว — ไม่ใช้ public สำหรับข้อมูลแอป)"],
            ["Connection", "DATABASE_URL — ผ่าน pg pool ใน Express backend"],
            ["ตาราง (Tables)", str(table_count)],
            ["View", str(view_count)],
            ["Foreign Key", str(len(fk_rels))],
            ["Index", str(len(indexes))],
            ["Migration files", str(mig_count)],
            ["ศูนย์กลางข้อมูล", "app.tbiw37n (Work Order / IW37N)"],
            ["Legacy source", "MySQL sap_lay — ชื่อตาราง/คอลัมน์คงเดิม"],
        ],
    )

    doc.add_heading("1.1 เป้าหมายการออกแบบ", level=2)
    add_bullets(
        doc,
        [
            "รองรับ workflow PM จริงของโรงงาน Pepsi — Import SAP → Planning → Calendar → Confirmation → Export",
            "สืบทอด schema legacy PHP เพื่อลดความเสี่ยง migration ข้อมูล",
            "Normalize ข้อมูล transactional รอบ tbiw37n — master data แยกตาราง lookup",
            "RBAC แยก role/permission ชัดเจน — ไม่พึ่ง userst อย่างเดียว",
            "เก็บรูป confirmation ใน DB (BYTEA WebP) — ไม่พึ่งโฟลเดอร์ imgMember",
            "Audit trail สำหรับ admin และ master data changes",
            "รองรับ automation (integration_job, import batch) และ Portal หลาย module",
        ],
    )

    doc.add_page_break()

    # 2 Conventions
    doc.add_heading("2. หลักการและมาตรฐานการออกแบบ", level=1)

    doc.add_heading("2.1 การตั้งชื่อ (Naming)", level=2)
    add_table(
        doc,
        ["รูปแบบ", "ตัวอย่าง", "หมายเหตุ"],
        [
            ["ตาราง legacy", "tbworkcenter, tbiw37n, tbplangingwork", "prefix tb + ชื่อเดิม MySQL"],
            ["ตารางใหม่ admin", "tbl_role, tbl_setting, tbl_audit_log", "prefix tbl_"],
            ["Primary key legacy", "idwkctr (varchar), idiw37 (serial)", "คงรูปแบบเดิม"],
            ["View", "view_order, view_exportconfirm", "prefix view_"],
            ["Index", "idx_{table}_{columns}", "สร้างใน migration"],
            ["Permission", "work-orders.read, planning.assign", "group.action"],
        ],
    )

    doc.add_heading("2.2 มาตรฐานชนิดข้อมูล", level=2)
    add_table(
        doc,
        ["ชนิด PostgreSQL", "ใช้กับ", "หมายเหตุ"],
        [
            ["varchar(n)", "รหัส SAP, wkorder, wkctr, status", "ความยาวตาม legacy"],
            ["text", "คำอธิบาย, หมายเหตุ, path", "ไม่จำกัดความยาว"],
            ["serial / bigserial", "idiw37, idmenu, log id", "auto increment"],
            ["bigint", "bscstart, cday, wkctrdate", "Unix epoch จาก legacy PHP"],
            ["timestamptz", "created_at, confirm_qc_at", "timezone Asia/Bangkok"],
            ["numeric(12,2)", "work, actwork, untime", "ชั่วโมงงาน WO"],
            ["boolean", "is_secret, menu flags", "ค่า true/false"],
            ["jsonb", "tbl_setting.setting_value", "flexible config"],
            ["bytea", "tbconfirmimg.img_data", "รูป WebP ใน DB"],
            ["char(1)", "userst (A/U/W legacy)", "ยังใช้ fallback RBAC"],
        ],
    )

    doc.add_heading("2.3 วันที่และเวลา", level=2)
    doc.add_paragraph(
        "ฟิลด์ legacy หลายตัวเก็บวันที่เป็น BIGINT (unix timestamp วินาที) เช่น tbiw37n.bscstart, "
        "tbmoveplan.cday — backend แปลงเป็น Date ตอน API/UI. "
        "ฟิลด์ใหม่ใช้ timestamptz เป็นหลัก (confirm_qc_at, created_at)"
    )

    doc.add_heading("2.4 การเก็บไฟล์และรูปภาพ", level=2)
    add_bullets(
        doc,
        [
            "รูป confirmation: tbconfirmimg.img_data (BYTEA) + mime (image/webp)",
            "img_phase: before | after — แยกรูปก่อน/หลัง PM",
            "โลโก้แอป: tbl_setting app.logo_bytes (jsonb base64) หรือ API branding",
            "Backup: ไฟล์ .tar.gz บนดิสก์ — metadata ใน tbl_backup_history",
            "Import SAP: ไฟล์ชั่วคราวใน inbound/ — ไม่เก็บถาวรใน DB",
        ],
    )

    doc.add_page_break()

    # 3 Security
    doc.add_heading("3. โมเดลความปลอดภัยและ RBAC", level=1)
    doc.add_paragraph("การควบคุมสิทธิ์ใช้ตาราง RBAC หลักดังนี้:")

    add_table(
        doc,
        ["ตาราง", "บทบาท"],
        [
            ["tbl_role", "บทบาทระบบ (Admin, Planner, Technician…) — role_code"],
            ["tbl_permission", "สิทธิ์ย่อย perm_code (group.action)"],
            ["tbl_role_permission", "Many-to-many role ↔ permission"],
            ["tbworkcenter.userst", "Legacy A/U/W — fallback ถ้ายังไม่ migrate RBAC"],
            ["tbmenu.menuright", "Legacy กรองเมนู — แทนที่ด้วย permissions"],
            ["tbl_blocked_ip", "บล็อก IP หลัง login ผิดเกินกำหนด"],
        ],
    )

    doc.add_heading("3.1 กลุ่ม Permission (perm_group)", level=2)
    add_table(doc, ["perm_group", "โมดูลแอป"], PERM_GROUPS)

    doc.add_heading("3.2 ข้อมูลอ่อนไหว", level=2)
    add_table(
        doc,
        ["ข้อมูล", "การป้องกัน"],
        [
            ["รหัสผ่าน tbworkcenter.pass", "hash ใน production — ไม่ log"],
            ["tbl_setting.is_secret", "ไม่ expose ผ่าน public settings API"],
            ["app_license / keys", "is_secret=true ใน tbl_setting"],
            ["Audit log", "บันทึก actor, action — ไม่เก็บ password"],
        ],
    )

    doc.add_page_break()

    # 4 Module architecture
    doc.add_heading("4. สถาปัตยกรรมข้อมูลแยกโมดูล", level=1)

    backend_groups = group_tables_by_module(tables)
    add_table(
        doc,
        ["โมดูล", "ตาราง", "จำนวนคอลัมน์รวม", "คำอธิบาย"],
        [
            [
                mod,
                str(len(files)),
                str(sum(len(t.columns) for _, t in files)),
                MODULE_MAP.get(files[0][0], ("", files[0][0]))[1] if files else "",
            ]
            for mod, files in sorted(backend_groups.items(), key=lambda x: x[0])
        ],
    )

    doc.add_heading("4.1 Data Flow ระหว่างโมดูล", level=2)
    add_bullets(
        doc,
        [
            "Import: tbiw37n_import_batch → tbiw37n_import_row → tbiw37n",
            "Planning: tbplangingwork (FK idiw37) ← tbiw37n",
            "Calendar: tbmoveplan (FK idiw37) + view_order",
            "Confirmation: tbcofirm, tbconfirmimg, tbwrkclose → confirm_qc_status บน tbiw37n",
            "Export: view_exportconfirm (approved QC เท่านั้น)",
            "PM Execution: tbwo_pm_reading, tbwo_pm_note_entry (FK idiw37)",
            "Manhour: tbmanhours (FK idwkctr → tbworkcenter)",
            "Admin: tbl_setting, tbl_audit_log, tbl_backup_history",
        ],
    )

    doc.add_page_break()

    # 5 Hub table tbiw37n
    doc.add_heading("5. ตารางศูนย์กลาง — app.tbiw37n", level=1)
    doc.add_paragraph(
        "tbiw37n เก็บ Work Order จาก SAP IW37N เป็นจุดเชื่อมโยงหลักของระบบ "
        "ตารางลูกส่วนใหญ่อ้างอิง idiw37 (ON DELETE CASCADE)"
    )

    iw = tables.get("tbiw37n")
    if iw:
        add_table_dict(doc, "tbiw37n", iw, comments)

    doc.add_heading("5.1 ตารางลูกที่อ้างอิง idiw37", level=2)
    hub = sorted({r[0] for r in fk_rels if r[2] == "tbiw37n"})
    add_bullets(doc, [f"app.{t}" for t in hub])

    doc.add_page_break()

    # 6 Data dictionary by module
    doc.add_heading("6. พจนานุกรมข้อมูล (Data Dictionary)", level=1)
    doc.add_paragraph("รายละเอียดทุกตารางและ View แยกตามโมดูล")

    by_module: dict[str, list[tuple[str, Table]]] = defaultdict(list)
    for name, tbl in sorted(tables.items()):
        mod, _ = module_for(name, tbl.is_view)
        by_module[mod].append((name, tbl))

    sec = 0
    for mod in sorted(by_module.keys()):
        sec += 1
        doc.add_heading(f"6.{sec} {mod}", level=2)
        for tname, tbl in by_module[mod]:
            if tname == "tbiw37n":
                doc.add_paragraph("(รายละเอียดเต็มดู §5)")
                continue
            add_table_dict(doc, tname, tbl, comments)
        doc.add_page_break()

    # 7 Views
    doc.add_heading("7. Database Views", level=1)
    view_rows = []
    for name, tbl in sorted(tables.items()):
        if not tbl.is_view:
            continue
        _, desc = module_for(name, True)
        view_rows.append([name, desc, (tbl.view_sql or "")[:150] + "…"])
    add_table(doc, ["View", "วัตถุประสงค์", "SQL (ย่อ)"], view_rows)

    doc.add_page_break()

    # 8 Indexes
    doc.add_heading("8. ดัชนี (Indexes)", level=1)
    doc.add_paragraph(
        "Index สำคัญสำหรับ performance — โดยเฉพาะปฏิทิน (bscstart, syst), "
        "confirmation (idiw37), QC pending partial index"
    )
    idx_by_table: dict[str, list[DbIndex]] = defaultdict(list)
    for idx in indexes:
        idx_by_table[idx.table].append(idx)

    for tname in sorted(idx_by_table.keys()):
        doc.add_heading(f"app.{tname}", level=3)
        add_table(
            doc,
            ["Index name", "คอลัมน์"],
            [[i.name, i.columns] for i in idx_by_table[tname]],
            header_fill="FFF2CC",
            font_size=8,
        )

    doc.add_page_break()

    # 9 FK
    doc.add_heading("9. ความสัมพันธ์ Foreign Key", level=1)
    add_table(
        doc,
        ["ตาราง", "FK Column", "อ้างอิง", "PK"],
        [[a, b, c, d] for a, b, c, d in fk_rels],
    )

    doc.add_page_break()

    # 10 Settings
    doc.add_heading("10. ตารางตั้งค่าระบบ — tbl_setting", level=1)
    doc.add_paragraph("Key-value store แบบ jsonb — แบ่ง category: branding, system, feature, backup")
    add_table(
        doc,
        ["setting_key", "category", "คำอธิบาย"],
        [[k, c, d] for k, c, d in SETTING_KEYS],
    )

    doc.add_page_break()

    # 11 Migration & ops
    doc.add_heading("11. Migration และการดำเนินงาน", level=1)

    doc.add_heading("11.1 วิธีรัน Migration", level=2)
    add_bullets(
        doc,
        [
            "ไฟล์อยู่ที่ database/migrations/ — เรียงตามเลขนำหน้า 001–103+",
            "รันครบ: database/scripts/run-all-migrations.ps1",
            "Admin only: database/scripts/run-admin-migrations.ps1",
            "ตรวจสอบ: database/scripts/verify-admin-environment.ps1",
            "ทุก migration ใช้ CREATE IF NOT EXISTS / ADD COLUMN IF NOT EXISTS — idempotent",
        ],
    )

    doc.add_heading("11.2 Backup & Audit", level=2)
    add_table(
        doc,
        ["ตาราง", "หน้าที่"],
        [
            ["tbl_backup_history", "บันทึกแต่ละครั้ง pg_dump — path, size, status"],
            ["tbl_audit_log", "actor, action, resource, payload JSON"],
            ["tbl_resource_revision", "revision ทรัพยากร master data"],
            ["integration_job", "สถานะ job import/export/watch"],
        ],
    )

    doc.add_heading("11.3 การตรวจสอบ Schema", level=2)
    add_bullets(
        doc,
        [
            "GET /api/v1/admin/health — migration status, DB connectivity",
            "database/scripts/verify_admin_data_tables.sql",
            "Query information_schema.tables WHERE table_schema='app'",
        ],
    )

    doc.add_page_break()

    # 12 ER image reference
    doc.add_heading("12. แผนภาพ ER-Diagram", level=1)
    doc.add_paragraph("อ้างอิงแผนภาพความสัมพันธ์เต็มรูปแบบ:")
    if ER_IMAGE.exists():
        doc.add_picture(str(ER_IMAGE), width=Inches(6.2))
        cap = doc.add_paragraph("รูปที่ 1 — ER-Diagram รวม (ดู ER-DIAGRAM-DETAILED-TH.docx)")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("13. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "docs/customer-requirements/ER-DIAGRAM-DETAILED-TH.docx",
            "docs/customer-requirements/FLOW-DIAGRAM-SWIMLANE-TH.docx",
            "database/migrations/*.sql",
            "docs/USER-MANUAL-TH.md",
            "from customer/tools/generate_database_design_docx.py",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Tables: {table_count}, Views: {view_count}, Indexes: {len(indexes)}, Comments: {len(comments)}")


if __name__ == "__main__":
    main()
