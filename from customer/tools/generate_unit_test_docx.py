#!/usr/bin/env python3
"""Generate PM Pepsi App — detailed Unit Test documentation (Word .docx)."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "PM-Pepsi-App" / "backend" / "src"
FRONTEND = ROOT / "PM-Pepsi-App" / "frontend" / "src"
OUT = ROOT / "docs" / "customer-requirements" / "UNIT-TEST-DETAILED-TH.docx"

# Vitest results (re-run npm test to refresh)
BACKEND_FILES = 99
BACKEND_PASSED = 311
BACKEND_SKIPPED = 1
FRONTEND_FILES = 54
FRONTEND_PASSED = 147

TEST_NAME_RE = re.compile(
    r"""(?:^\s*(?:it|test)(?:\.(?:only|skip|todo|concurrent|each))?\s*\(\s*)
        (?:['"`]([^'"`]+)['"`]|`([^`]+)`)
    """,
    re.MULTILINE | re.VERBOSE,
)
DESCRIBE_RE = re.compile(
    r"""^\s*describe(?:\.(?:only|skip|todo))?\s*\(\s*['"`]([^'"`]+)['"`]""",
    re.MULTILINE,
)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "D1E8F5",
    font_size: int = 9,
) -> None:
    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def scan_test_file(path: Path, base: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    tests = []
    for m in TEST_NAME_RE.finditer(text):
        name = m.group(1) or m.group(2) or ""
        if name:
            tests.append(name.strip())
    describes = DESCRIBE_RE.findall(text)
    rel = path.relative_to(base).as_posix()
    parts = rel.split("/")
    if len(parts) >= 2:
        category = parts[0]
        if category == "features" and len(parts) >= 3:
            category = f"features/{parts[1]}"
    else:
        category = parts[0] if parts else "other"
    return {
        "path": rel,
        "category": category,
        "describes": describes,
        "tests": tests,
        "count": len(tests),
    }


def collect_tests(root: Path, patterns: list[str]) -> list[dict]:
    files: list[Path] = []
    for pat in patterns:
        files.extend(sorted(root.rglob(pat)))
    seen: set[Path] = set()
    unique = []
    for f in files:
        if f not in seen:
            seen.add(f)
            unique.append(f)
    return [scan_test_file(f, root) for f in unique]


def group_by_category(items: list[dict]) -> dict[str, list[dict]]:
    groups: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        groups[item["category"]].append(item)
    return dict(sorted(groups.items()))


def format_test_list(tests: list[str], max_show: int = 8) -> str:
    if not tests:
        return "—"
    shown = tests[:max_show]
    lines = "\n".join(f"• {t}" for t in shown)
    if len(tests) > max_show:
        lines += f"\n… (+{len(tests) - max_show} เคส)"
    return lines


def main() -> None:
    today = date.today().isoformat()
    backend_items = collect_tests(BACKEND, ["*.test.ts"])
    frontend_items = collect_tests(FRONTEND, ["*.test.ts", "*.test.tsx"])

    backend_scanned = sum(i["count"] for i in backend_items)
    frontend_scanned = sum(i["count"] for i in frontend_items)

    doc = Document()

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\nเอกสาร Unit Test ฉบับละเอียด\n(Backend · Frontend · Integration API)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้แคตตาล็อก unit test ทั้งหมดของ PM-Pepsi-App "
        "รวมชื่อไฟล์ หมวดหมู่ จำนวนเคส และรายชื่อ test case ที่สแกนจากซอร์สโค้ด "
        "ใช้ประกอบการตรวจสอบคุณภาพก่อน UAT และส่งมอบลูกค้า"
    )
    doc.add_paragraph(
        "เอกสารคู่กัน: STRESS-TEST-DETAILED-TH.docx (Stress Test แยกฉบับ)"
    )

    doc.add_page_break()

    # --- 1 Overview ---
    doc.add_heading("1. ภาพรวม", level=1)
    doc.add_paragraph(
        "Unit test ใช้ Vitest รันแยก backend และ frontend ไม่ต้องเปิด browser "
        "(frontend ใช้ jsdom สำหรับ component test บางส่วน) "
        "Integration API test ใช้ Supertest เรียก Express app จริง"
    )

    doc.add_heading("1.1 สรุปผลการรันล่าสุด", level=2)
    add_table(
        doc,
        ["ส่วน", "ไฟล์ทดสอบ", "เคสผ่าน", "ข้าม", "เคสสแกนจากซอร์ส", "คำสั่ง"],
        [
            [
                "Backend",
                str(BACKEND_FILES),
                str(BACKEND_PASSED),
                f"{BACKEND_SKIPPED} skipped",
                str(backend_scanned),
                "cd PM-Pepsi-App/backend && npm test",
            ],
            [
                "Frontend",
                str(FRONTEND_FILES),
                str(FRONTEND_PASSED),
                "0",
                str(frontend_scanned),
                "cd PM-Pepsi-App/frontend && npm test",
            ],
            [
                "รวม",
                str(BACKEND_FILES + FRONTEND_FILES),
                str(BACKEND_PASSED + FRONTEND_PASSED),
                f"{BACKEND_SKIPPED} skipped",
                str(backend_scanned + frontend_scanned),
                "—",
            ],
        ],
    )
    doc.add_paragraph(
        "หมายเหตุ: จำนวนเคสสแกนจากซอร์สอาจต่างจาก Vitest เล็กน้อย "
        "(เช่น test.each, dynamic name) — เกณฑ์ผ่านอ้างอิงผล Vitest เป็นหลัก"
    )

    doc.add_heading("1.2 เครื่องมือและโครงสร้าง", level=2)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Test runner", "Vitest v4"],
            ["Backend runtime", "Node.js (native, ไม่ใช้ jsdom)"],
            ["Frontend runtime", "jsdom + @testing-library/react"],
            ["Integration", "Supertest + Express app instance"],
            ["Assertion", "Vitest expect (compatible Jest API)"],
            ["Backend path", "PM-Pepsi-App/backend/src/**/*.test.ts"],
            ["Frontend path", "PM-Pepsi-App/frontend/src/**/*.test.{ts,tsx}"],
        ],
    )

    doc.add_heading("1.3 สภาพแวดล้อม", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า", "หมายเหตุ"],
        [
            ["Node.js", "20+", "ตาม package.json engines"],
            ["DATABASE_URL", "ไม่จำเป็นสำหรับ unit ส่วนใหญ่", "Supertest ใช้ mock/stub"],
            ["เวลารัน backend", "~18 วินาที", "99 ไฟล์"],
            ["เวลารัน frontend", "~31 วินาที", "54 ไฟล์"],
        ],
    )

    doc.add_page_break()

    # --- 2 Backend summary by category ---
    doc.add_heading("2. Backend Unit Test — สรุปตามหมวด", level=1)
    backend_groups = group_by_category(backend_items)
    cat_rows = []
    for cat, files in backend_groups.items():
        cat_rows.append([cat, str(len(files)), str(sum(f["count"] for f in files))])
    add_table(doc, ["หมวด (โฟลเดอร์)", "จำนวนไฟล์", "เคส (สแกน)"], cat_rows)

    doc.add_heading("2.1 หมวด lib/ — Business logic & utilities", level=2)
    doc.add_paragraph(
        "ทดสอบ pure function: parser, filter, คำนวณชั่วโมง, RBAC helper, "
        "calendar policy, confirmation export, image processing"
    )

    doc.add_heading("2.2 หมวด services/ — Service layer", level=2)
    doc.add_paragraph(
        "ทดสอบ business service: IW37N import, plan calendar, confirmation, "
        "manhour parse, board feed, admin health, telegram hooks"
    )

    doc.add_heading("2.3 หมวด schemas/ — Zod validation", level=2)
    doc.add_paragraph("ทดสอบ request/response schema ของ API endpoints")

    doc.add_heading("2.4 หมวด routes/ — Integration API (Supertest)", level=2)
    doc.add_paragraph(
        "ทดสอบ HTTP จริง: admin API, roles, menu, board kiosk — ตรวจ status code, "
        "RBAC 403, response shape"
    )

    doc.add_heading("2.5 หมวด middleware/ — Middleware", level=2)
    doc.add_paragraph("rate-limit, audit-master-data")

    doc.add_page_break()

    # --- 3 Backend full catalog ---
    doc.add_heading("3. Backend — รายการไฟล์และ Test Case ทั้งหมด", level=1)

    for cat, files in backend_groups.items():
        doc.add_heading(f"3.{list(backend_groups.keys()).index(cat) + 1} {cat}/", level=2)
        for f in sorted(files, key=lambda x: x["path"]):
            doc.add_heading(f["path"], level=3)
            if f["describes"]:
                doc.add_paragraph(f"describe: {', '.join(f['describes'][:5])}")
            add_table(
                doc,
                ["#", "ชื่อ test case"],
                [[str(i + 1), t] for i, t in enumerate(f["tests"])],
                font_size=8,
            )
            if not f["tests"]:
                doc.add_paragraph("(ไม่พบ it/test แบบ static — อาจใช้ test.each)")

    doc.add_page_break()

    # --- 4 Frontend ---
    doc.add_heading("4. Frontend Unit Test — สรุปตามหมวด", level=1)
    frontend_groups = group_by_category(frontend_items)
    fe_cat_rows = []
    for cat, files in frontend_groups.items():
        fe_cat_rows.append([cat, str(len(files)), str(sum(f["count"] for f in files))])
    add_table(doc, ["หมวด", "จำนวนไฟล์", "เคส (สแกน)"], fe_cat_rows)

    doc.add_heading("4.1 หมวด lib/ — Utilities & RBAC", level=2)
    doc.add_paragraph("navigation, permissions, API client, theme, locale, chart helpers")

    doc.add_heading("4.2 หมวด components/ — UI components", level=2)
    doc.add_paragraph("scheduling dialogs, admin tour, nav config, report export button")

    doc.add_heading("4.3 หมวด features/ — Feature schemas & pages", level=2)
    doc.add_paragraph("Zod schemas, date filters, utilization charts, master entities")

    doc.add_page_break()

    doc.add_heading("5. Frontend — รายการไฟล์และ Test Case ทั้งหมด", level=1)

    for cat, files in frontend_groups.items():
        doc.add_heading(f"5.{list(frontend_groups.keys()).index(cat) + 1} {cat}/", level=2)
        for f in sorted(files, key=lambda x: x["path"]):
            doc.add_heading(f["path"], level=3)
            if f["describes"]:
                doc.add_paragraph(f"describe: {', '.join(f['describes'][:5])}")
            add_table(
                doc,
                ["#", "ชื่อ test case"],
                [[str(i + 1), t] for i, t in enumerate(f["tests"])],
                font_size=8,
            )
            if not f["tests"]:
                doc.add_paragraph("(ไม่พบ it/test แบบ static)")

    doc.add_page_break()

    # --- 6 Domain mapping ---
    doc.add_heading("6. แมปโดเมนธุรกิจ → Unit Test", level=1)
    doc.add_paragraph("ตารางด้านล่างเชื่อมโมดูลแอปกับไฟล์ทดสอบหลัก (ไม่ครบทุกไฟล์)")

    domain_map = [
        ["IW37N Import", "iw37n-parser, iw37n-factory-scope, confirmation-import, iw37n-import-summary"],
        ["Planning / Calendar", "plan-calendar, calendar-move-policy, planning-group, planning-available-hours"],
        ["Work Order / PM", "work-order-workflow, wo-pm-phase, wo-pm-execution, work-order-close-guard"],
        ["Confirmation / Export SAP", "confirm-qc-status, confirmation-export-scope, confirmation-mass-export"],
        ["Manhour / HR", "manhour-minutes, manhours-parse, manhours-hr-utilization, week-to-week"],
        ["Reports / KPI", "reports.test, reports-range, reports-wktype, eng-utilization-chart (FE)"],
        ["RBAC / Security", "has-permission, rbac-role-code, password-policy, rate-limit, login-lockout"],
        ["Admin Console", "admin-health, admin-settings, admin-api.supertest, admin-tour (FE)"],
        ["Board / Kiosk", "board-activity, board-period, board-kiosk.supertest"],
        ["Personnel", "personnel-admin.test, personnel-schemas (FE)"],
        ["Telegram", "telegram-confirm, telegram-close"],
        ["Images / Media", "confirm-image, image-magic, remove-light-background, svg-sanitize"],
        ["Portal", "portal-enabled, portal-rbac-preview (FE)"],
        ["Module Handoff", "module-handoff.test"],
    ]
    add_table(doc, ["โมดูลธุรกิจ", "ไฟล์ทดสอบหลัก"], domain_map)

    doc.add_page_break()

    # --- 7 Commands ---
    doc.add_heading("7. คำสั่งรันและเกณฑ์ผ่าน", level=1)

    doc.add_heading("7.1 คำสั่งหลัก", level=2)
    add_table(
        doc,
        ["คำสั่ง", "คำอธิบาย"],
        [
            ["cd PM-Pepsi-App/backend && npm test", "รัน unit ทั้งหมด backend (99 ไฟล์)"],
            ["cd PM-Pepsi-App/backend && npm run test:integration", "Supertest admin API เท่านั้น"],
            ["cd PM-Pepsi-App/frontend && npm test", "รัน unit ทั้งหมด frontend (54 ไฟล์)"],
            ["cd PM-Pepsi-App/frontend && npm run test:qa", "ชุด QA nav/shell แบบย่อ"],
            ["npx vitest run path/to/file.test.ts", "รันไฟล์เดียว"],
            ["npx vitest watch path/to/file.test.ts", "โหมด watch ขณะพัฒนา"],
        ],
    )

    doc.add_heading("7.2 เกณฑ์ผ่าน", level=2)
    add_bullets(
        doc,
        [
            "ทุกเคสต้อง pass — exit code 0",
            "ไม่มี test fail หรือ timeout",
            "การ skip ต้องมีเหตุผล (เช่น OS-specific) และบันทึกใน release note",
            "Regression: หลังแก้ bug ต้องมี test case ครอบคลุม (ถ้าเป็น logic ที่ unit test ได้)",
        ],
    )

    doc.add_heading("7.3 Checklist ก่อน merge / UAT", level=2)
    add_table(
        doc,
        ["ลำดับ", "รายการ", "คำสั่ง", "ผ่าน"],
        [
            ["1", "Backend unit test", "cd backend && npm test", "☐"],
            ["2", "Frontend unit test", "cd frontend && npm test", "☐"],
            ["3", "Integration API (ถ้าแก้ route)", "npm run test:integration", "☐"],
            ["4", "Build frontend", "cd frontend && npm run build", "☐"],
            ["5", "Build backend", "cd backend && npm run build", "☐"],
        ],
    )

    doc.add_heading("8. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "docs/customer-requirements/STRESS-TEST-DETAILED-TH.docx — Stress Test (เอกสารคู่)",
            "docs/customer-requirements/PRE-UAT-UI-PHASES.md — เกณฑ์ UI U4",
            "docs/PRE-UAT-MASTER-PHASES.md — แผน UAT รวม",
            "AGENTS.md — คำสั่ง dev ท้องถิ่น",
            "from customer/tools/generate_unit_test_docx.py — สคริปต์สร้างเอกสารฉบับนี้",
            "from customer/tools/generate_testing_docs.py — สร้างทั้ง 2 เอกสารพร้อมกัน",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Backend: {len(backend_items)} files, {backend_scanned} tests scanned")
    print(f"Frontend: {len(frontend_items)} files, {frontend_scanned} tests scanned")


if __name__ == "__main__":
    main()
