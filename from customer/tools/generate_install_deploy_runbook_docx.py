#!/usr/bin/env python3
"""Generate PM Pepsi App — Installation / Deploy + Operations Runbook (Word .docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "INSTALL-DEPLOY-RUNBOOK-TH.docx"


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "2F5496",
    font_size: int = 8,
    header_font_white: bool = True,
) -> None:
    from docx.shared import RGBColor

    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                if header_font_white:
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def main() -> None:
    today = date.today().isoformat()
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Dashboard & Monitoring System\n"
        "Installation, Deployment & Operations Runbook\n"
        "(คู่มือติดตั้ง · Deploy · ปฏิบัติการ)"
    )
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"ลูกค้า: Pepsi-Cola (Thailand) Trading Co., Ltd.\n"
        f"ผู้พัฒนา: S.Y. Interactive Development Limited\n"
        f"รหัสเอกสาร: PM-PEPSI-OPS-001\n"
        f"เวอร์ชัน: 1.0 · วันที่: {today}\n"
        f"แพลตฟอร์มเป้าหมาย: Windows Server 2019 · PostgreSQL · Node.js LTS"
    ).font.size = Pt(10)

    doc.add_page_break()

    # Document control
    doc.add_heading("การควบคุมเอกสาร", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["ชื่อเอกสาร", "Installation, Deployment & Operations Runbook"],
            ["รหัส", "PM-PEPSI-OPS-001"],
            ["เวอร์ชัน", "1.0"],
            ["สถานะ", "Draft — ส่งมอบ onsite"],
            ["เจ้าของ", "Technical Lead / DevOps"],
            ["ผู้อนุมัติ (ลูกค้า IT)", "___________________  วันที่ _______"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    # 1 Introduction
    doc.add_heading("1. บทนำ", level=1)
    doc.add_paragraph(
        "เอกสารฉบับนี้ครอบคลุมการติดตั้ง PM-Pepsi-App บนเซิร์ฟเวอร์โรงงาน (offline/air-gap), "
        "การ deploy production, และ Runbook สำหรับทีม IT/ผู้ดูแลระบบ — "
        "อ้างอิง skills.md, SETUP-NEW-MACHINE.md, PRE-UAT-MASTER-PHASES §P6"
    )

    doc.add_heading("1.1 ขอบเขต", level=2)
    add_bullets(
        doc,
        [
            "In scope: Windows Server · PostgreSQL · Node.js API · React static (IIS) · PM2",
            "In scope: Migration DB 001–103+ · Backup/Restore · Integration folder SAP",
            "Out of scope: Docker/Tailscale บน production (ใช้ dev ได้ — ภาคผนวก A)",
            "Out of scope: Remote support จากอินเทอร์เน็ต — onsite เท่านั้น (ยกเว้นข้อตกลง SRS)",
        ],
    )

    doc.add_heading("1.2 สถาปัตยกรรม Production (แนะนำ)", level=2)
    add_table(
        doc,
        ["ชั้น", "เทคโนโลยี", "พอร์ต/Path"],
        [
            ["Client", "Chrome/Edge ภายในโรงงาน", "HTTPS → IIS"],
            ["Web", "IIS — serve frontend/dist", "443 หรือ reverse proxy"],
            ["API", "Node.js Express (PM2)", "127.0.0.1:4000"],
            ["DB", "PostgreSQL Windows Service", "127.0.0.1:5432 หรือ 5433"],
            ["Integration", "โฟลเดอร์ inbound/outbound", "D:\\pm-deploy\\integration"],
            ["Backup", "pg_dump + Admin UI", "D:\\pm-deploy\\backups"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )
    doc.add_paragraph(
        "IIS proxy `/api` → `http://127.0.0.1:4000` — ผู้ใช้เข้า URL เดียว (same-origin). "
        "Telegram webhook (ถ้าใช้) ต้อง HTTPS สาธารณะ — ใช้ Cloudflare Tunnel หรือ reverse proxy ที่ IT อนุมัติ"
    )

    doc.add_page_break()

    # 2 Prerequisites
    doc.add_heading("2. ความต้องการเซิร์ฟเวอร์", level=1)
    add_table(
        doc,
        ["รายการ", "ขั้นต่ำ (จากบรีฟลูกค้า)", "หมายเหตุ"],
        [
            ["OS", "Windows Server 2019 Standard 64-bit", "Intel Xeon · 32 GB RAM"],
            ["Drive C:", "≥ 200 GB free", "OS + ซอฟต์แวร์"],
            ["Drive D:", "≥ 1 TB (deploy ~300 GB + backup)", "โค้ด · DB · backup · integration"],
            ["Network", "LAN โรงงาน", "ไม่เปิด VPN remote ตาม default"],
            ["Node.js", "20 LTS ขึ้นไป", "npm ci / build"],
            ["PostgreSQL", "14+ (15/16 แนะนำ)", "schema app"],
            ["Git", "ล่าสุด", "clone/pull onsite"],
            ["PM2", "global npm", "process manager API"],
            ["IIS", "+ URL Rewrite + ARR (proxy)", "static + /api proxy"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("2.1 โครงสร้างโฟลเดอร์บน D:", level=2)
    add_code(
        doc,
        "D:\\pm-deploy\\\n"
        "  repo\\                    ← git clone (sap_241163)\n"
        "  integration\\             ← INTEGRATION_DATA_DIR\n"
        "    inbound\\iw37n\\\n"
        "    inbound\\confirm\\\n"
        "    outbound\\confirm\\\n"
        "    archive\\ · error\\\n"
        "  backups\\                 ← pg_dump รายวัน\n"
        "  logs\\                    ← PM2 logs (optional)\n"
        "  pgdata\\                  ← PostgreSQL data (ถ้า bind ไป D:)",
    )

    doc.add_page_break()

    # 3 Installation
    doc.add_heading("3. การติดตั้งครั้งแรก (Fresh Install)", level=1)

    doc.add_heading("3.1 Checklist ก่อนเริ่ม", level=2)
    add_table(
        doc,
        ["#", "รายการ", "☐"],
        [
            ["1", "สิทธิ์ Administrator บนเซิร์ฟเวอร์", "☐"],
            ["2", "พอร์ต 5432/5433 ไม่ชน SQL Server / PG อื่น", "☐"],
            ["3", "พอร์ต 4000 ว่าง (API internal)", "☐"],
            ["4", "IIS site + certificate (หรือแผน Cloudflare Tunnel)", "☐"],
            ["5", "ไม่ commit .env — เก็บใน D:\\pm-deploy\\repo\\PM-Pepsi-App\\backend\\.env", "☐"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("3.2 ติดตั้ง PostgreSQL", level=2)
    add_bullets(
        doc,
        [
            "ติดตั้ง PostgreSQL บน Windows — บันทึกพอร์ต (5432 หรือ 5433 ถ้าชน)",
            "สร้าง user/database/schema app (ดู SQL ด้านล่าง)",
            "ตั้ง Windows Service ให้ start อัตโนมัติหลัง reboot",
        ],
    )
    add_code(
        doc,
        "CREATE USER pepsipm WITH PASSWORD '***';\n"
        "CREATE DATABASE pepsi_pm OWNER pepsipm ENCODING 'UTF8';\n"
        "GRANT ALL PRIVILEGES ON DATABASE pepsi_pm TO pepsipm;\n"
        "\\c pepsi_pm\n"
        "CREATE SCHEMA IF NOT EXISTS app AUTHORIZATION pepsipm;\n"
        "GRANT ALL ON SCHEMA app TO pepsipm;\n"
        "ALTER ROLE pepsipm SET search_path TO app, public;",
    )

    doc.add_heading("3.3 ดึงโค้ดและ build", level=2)
    add_code(
        doc,
        "cd D:\\pm-deploy\n"
        "git clone <repo-url> repo\n"
        "cd repo\\PM-Pepsi-App\\backend\n"
        "copy .env.example .env\n"
        "npm ci\n"
        "npm run build\n"
        "cd ..\\frontend\n"
        "npm ci\n"
        "npm run build",
    )

    doc.add_heading("3.4 ตั้งค่า Environment (.env)", level=2)
    add_table(
        doc,
        ["ตัวแปร", "บังคับ", "ตัวอย่าง / คำอธิบาย"],
        [
            ["DATABASE_URL", "ใช่", "postgresql://pepsipm:***@127.0.0.1:5433/pepsi_pm"],
            ["PORT", "ใช่", "4000"],
            ["SESSION_SECRET", "ใช่", "สตริงสุ่ม ≥32 ตัวอักษร"],
            ["CORS_ORIGIN", "ใช่", "https://pm.factory.local (URL ที่ผู้ใช้เปิด)"],
            ["APP_PUBLIC_URL", "ใช่", "URL สาธารณะ API (Telegram webhook)"],
            ["INTEGRATION_DATA_DIR", "แนะนำ", "D:\\pm-deploy\\integration"],
            ["INTEGRATION_WATCH_SCHEDULER", "ไม่", "1 (default) — 0 ปิด scan อัตโนมัติ"],
            ["TELEGRAM_BOT_TOKEN", "ถ้าใช้ TG", "จาก @BotFather"],
            ["TELEGRAM_WEBHOOK_SECRET", "ถ้าใช้ TG", "secret สำหรับ webhook header"],
            ["TELEGRAM_NOTIFY_ENABLED", "ไม่", "true/false"],
            ["NODE_ENV", "แนะนำ", "production"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )

    doc.add_heading("3.5 รัน Migration + Seed", level=2)
    add_code(
        doc,
        "cd D:\\pm-deploy\\repo\n"
        "$env:DATABASE_URL = \"postgresql://pepsipm:***@127.0.0.1:5433/pepsi_pm\"\n"
        "pwsh -File database\\scripts\\run-all-migrations.ps1\n"
        "# Production: อย่ารัน seed demo ทับข้อมูลจริง\n"
        "pwsh -File database\\scripts\\verify-admin-environment.ps1\n"
        "psql $env:DATABASE_URL -f database\\scripts\\verify_app_schema.sql",
    )
    doc.add_paragraph(
        "Migration ปัจจุบัน: 001–103+ ใน database/migrations/ — รันตามลำดับเลขไฟล์. "
        "บน production ที่ restore จาก backup ให้รัน migration เฉพาะเลขที่สูงกว่า backup"
    )

    doc.add_page_break()

    # 4 Deploy PM2 + IIS
    doc.add_heading("4. Deploy Production (PM2 + IIS)", level=1)

    doc.add_heading("4.1 PM2 — Backend API", level=2)
    add_code(
        doc,
        "npm install -g pm2\n"
        "cd D:\\pm-deploy\\repo\\PM-Pepsi-App\\backend\n"
        "pm2 start dist/index.js --name pm-api --env production\n"
        "pm2 save\n"
        "pm2 startup",
    )
    doc.add_paragraph(
        "แนะนำสร้าง ecosystem.config.cjs ใน backend — กำหนด cwd, env_file, max_memory_restart. "
        "หลัง reboot ตรวจ: pm2 list → pm-api online"
    )

    doc.add_heading("4.2 IIS — Frontend static + API proxy", level=2)
    add_bullets(
        doc,
        [
            "สร้าง Site ชี้ physical path: PM-Pepsi-App\\frontend\\dist",
            "ติดตั้ง URL Rewrite + Application Request Routing (ARR)",
            "Rule: /api/* → http://127.0.0.1:4000/api/* (preserve host)",
            "Binding HTTPS — certificate ภายใน หรือผ่าน Cloudflare Tunnel",
            "MIME types: .js, .css, .webp, .woff2 (Vite build)",
        ],
    )

    doc.add_heading("4.3 HTTPS / Cloudflare Tunnel (Telegram)", level=2)
    doc.add_paragraph(
        "โรงงาน offline ไม่เปิด inbound 443 — สำหรับ Telegram Bot webhook ใช้ Cloudflare Tunnel "
        "ออกไป HTTPS โดยไม่เปิดพอร์ตจากอินเทอร์เน็ต. ตั้ง APP_PUBLIC_URL ให้ตรง URL tunnel. "
        "ตั้ง webhook: POST https://api.telegram.org/bot<TOKEN>/setWebhook?url=<APP_PUBLIC_URL>/api/v1/telegram/webhook"
    )

    doc.add_heading("4.4 Smoke Test หลัง Deploy", level=2)
    add_table(
        doc,
        ["#", "การทดสอบ", "ผลที่คาด", "☐"],
        [
            ["1", "GET /api/v1/health", "200, db: ok", "☐"],
            ["2", "Login หน้าเว็บ", "เข้า Dashboard", "☐"],
            ["3", "เปิด /calendar, /planning", "ไม่ error 500", "☐"],
            ["4", "Import IW37N ทดสอบ", "inserted > 0", "☐"],
            ["5", "Admin → Backup manual", "ไฟล์ใน backups/", "☐"],
            ["6", "Telegram getWebhookInfo (ถ้ามี)", "url ถูกต้อง", "☐"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    # 5 Runbook
    doc.add_heading("5. Operations Runbook", level=1)

    doc.add_heading("5.1 งานประจำวัน (Daily)", level=2)
    add_table(
        doc,
        ["เวลา", "งาน", "วิธี", "ผู้รับผิดชอบ"],
        [
            ["07:00 / 19:00", "รอบ SAP (อ้างอิง)", "ตรวจ inbound/iw37n มีไฟล์ใหม่", "Planner/IT"],
            ["หลัง SAP export", "IW37N auto-scan", "/integration → Jobs หรือ watch folder", "ระบบ"],
            ["เช้า", "ตรวจ health", "GET /api/v1/health · PM2 online", "IT"],
            ["เย็น", "ตรวจ backup", "/admin/backup → history วันนี้ success", "IT"],
            ["ตามความจำเป็น", "QC + Export Confirm", "/confirmation", "Foreman"],
        ],
        header_fill="2F5496",
    )

    doc.add_heading("5.2 Backup & Restore", level=2)
    doc.add_paragraph("Backup อัตโนมัติ: ตั้ง cron ใน Admin → Settings/Backup (pg_dump → .sql.gz)")
    add_table(
        doc,
        ["สถานการณ์", "ขั้นตอน"],
        [
            ["Backup มือ", "Admin → /admin/backup → Create · เก็บที่ D:\\pm-deploy\\backups"],
            ["Backup CLI", "pg_dump DATABASE_URL -Fc -f pepsi_pm_YYYYMMDD.dump"],
            ["Restore (staging เท่านั้น)", "หยุด pm-api → pg_restore → รัน migration ที่ขาด → start pm-api"],
            ["Restore จาก Admin UI", "Maintenance mode ON → Restore → ตรวจ schema → OFF"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )
    doc.add_paragraph(
        "คำเตือน: Restore ทับ production ต้องมี backup ล่าสุด + แจ้งผู้ใช้ + ทดสอบบน staging ก่อน"
    )

    doc.add_heading("5.3 อัปเดตเวอร์ชัน (Release)", level=2)
    add_bullets(
        doc,
        [
            "1. Backup DB ก่อนทุกครั้ง",
            "2. Maintenance mode ON (Admin → Settings) — mutate API 503",
            "3. git pull / คัดลอก build ใหม่ onsite",
            "4. npm ci + npm run build (backend + frontend)",
            "5. pwsh run-all-migrations.ps1 (migration ใหม่เท่านั้น)",
            "6. pm2 restart pm-api",
            "7. รีเฟรช IIS (หรือ replace dist)",
            "8. Smoke test §4.4",
            "9. Maintenance mode OFF",
            "10. บันทึก commit/tag ใน change log",
        ],
    )

    doc.add_heading("5.4 Rollback", level=2)
    add_bullets(
        doc,
        [
            "Rollback โค้ด: git checkout tag ก่อนหน้า · build · pm2 restart · replace dist",
            "Rollback DB: restore จาก backup ก่อน deploy (เสียข้อมูลหลัง backup)",
            "Rollback migration: ไม่มี down script — ใช้ restore DB เป็นหลัก",
        ],
    )

    doc.add_heading("5.5 Integration SAP (โฟลเดอร์)", level=2)
    add_table(
        doc,
        ["โฟลเดอร์", "用途"],
        [
            ["inbound/iw37n", "วางไฟล์ IW37N .xlsx/.xls — watch scan"],
            ["inbound/confirm", "วาง Confirm IN จาก SAP"],
            ["outbound/confirm", "Export Confirm OUT จากระบบ"],
            ["processing", "ไฟล์กำลังประมวลผล"],
            ["archive/inbound/YYYY-MM", "เก็บไฟล์ที่ import แล้ว"],
            ["error", "ไฟล์ error — ตรวจ log + แก้แล้วย้ายกลับ inbound"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )
    doc.add_paragraph(
        "Runbook ไฟล์ค้าง error/: เปิด /integration → Jobs → ดู error message → "
        "แก้ไฟล์หรือ template → ย้ายกลับ inbound → Retry manual upload"
    )

    doc.add_heading("5.6 Maintenance Mode", level=2)
    doc.add_paragraph(
        "เปิดที่ Admin → Settings → Maintenance. ผล: banner แจ้งผู้ใช้ · API mutate 503 · "
        "Admin backup/settings ยังใช้ได้. ใช้ตอน migrate/restore/deploy"
    )

    doc.add_heading("5.7 Telegram (ถ้าเปิดใช้)", level=2)
    add_table(
        doc,
        ["งาน", "ขั้นตอน"],
        [
            ["ตั้ง Bot", "Admin → /admin/telegram · token ใน .env"],
            ["Webhook", "HTTPS + TELEGRAM_WEBHOOK_SECRET · setWebhook"],
            ["กลุ่ม Planner", "ลูกค้าสร้างกลุ่ม TG · ใส่ Chat ID ใน Admin"],
            ["ผูกช่าง", "Admin → Users → Telegram invite link"],
            ["ตรวจสอบ", "getWebhookInfo · ทดสอบจ่ายงาน 1 WO"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    # 6 Troubleshooting
    doc.add_heading("6. แก้ปัญหาที่พบบ่อย", level=1)
    add_table(
        doc,
        ["อาการ", "สาเหตุที่พบบ่อย", "แก้"],
        [
            ["health db: fail", "PostgreSQL ไม่รัน / DATABASE_URL ผิด", "ตรวจ service · .env · psql"],
            ["502 /api", "PM2 ไม่รัน / พอร์ต 4000", "pm2 restart pm-api · pm2 logs"],
            ["Login ได้แต่เมนูว่าง", "Migration RBAC/menu ไม่ครบ", "run-all-migrations · ตรวจ tbmenu"],
            ["Import skipped ทั้งก้อน", "Parser ไม่ตรงรูปแบบ SAP ALV", "ใช้ preview · ส่งไฟล์ให้ dev"],
            ["ปฏิทินว่างหลัง import", "เดือนปฏิทินไม่ตรง bscstart", "เปลี่ยนปี/เดือนบน /calendar"],
            ["ช่างไม่เห็นงาน", "ไม่ได้จ่ายที่ /planning · migration 098", "จ่าย work center · ตรวจสิทธิ์"],
            ["CORS error", "CORS_ORIGIN ไม่ตรง URL จริง", "แก้ .env · restart API"],
            ["Disk เต็ม D:", "backup/archive สะสม", "ลบ backup เก่า · archive cleanup policy"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("6.1 ตารางพอร์ต", level=2)
    add_table(
        doc,
        ["บริการ", "พอร์ต", "Exposure"],
        [
            ["PostgreSQL", "5432 หรือ 5433", "localhost only"],
            ["PM API (PM2)", "4000", "localhost only (IIS proxy)"],
            ["IIS HTTPS", "443", "LAN / Tunnel"],
            ["Adminer (dev only)", "8080", "อย่าเปิด production"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    # 7 Appendix
    doc.add_heading("7. ภาคผนวก", level=1)

    doc.add_heading("7.1 ภาคผนวก A — Docker (Dev / ทางเลือก)", level=2)
    doc.add_paragraph(
        "Production โรงงานใช้ PM2+IIS ตาม §4. สำหรับเครื่อง dev อาจใช้ Docker บน D: "
        "ดู from customer/server/DOCKER_AND_TAILSCALE.md — ไม่ใช้ Tailscale บน production ตาม skills.md"
    )

    doc.add_heading("7.2 ภาคผนวก B — เอกสารอ้างอิง", level=2)
    add_bullets(
        doc,
        [
            "docs/SETUP-NEW-MACHINE.md",
            "docs/ON-SITE-DATABASE-SETUP.md",
            "docs/PRE-UAT-MASTER-PHASES.md §P6",
            "docs/customer-requirements/UAT-ROUND-3-TH.md",
            "skills.md §4 Deploy offline",
            "PM-Pepsi-App/backend/.env.example",
        ],
    )

    doc.add_heading("7.3 ภาคผนวก C — Change Log (กรอก onsite)", level=2)
    add_table(
        doc,
        ["วันที่", "Commit/Tag", "Migration", "ผู้ดำเนินการ", "หมายเหตุ"],
        [
            ["", "", "", "", ""],
            ["", "", "", "", ""],
            ["", "", "", "", ""],
        ],
        header_fill="FFFFFF",
        header_font_white=False,
    )

    doc.add_heading("7.4 ภาคผนวก D — Escalation", level=2)
    add_table(
        doc,
        ["ระดับ", "ติดต่อ", "เมื่อ"],
        [
            ["L1", "IT โรงงาน", "Backup, restart, user/password"],
            ["L2", "Planner lead / Engineering", "SAP file, workflow"],
            ["L3", "S.Y. Interactive (onsite)", "Bug, migration, release"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร Installation / Deploy / Runbook —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
