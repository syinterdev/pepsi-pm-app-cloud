#!/usr/bin/env python3
"""Generate PM Pepsi App — DBRS per ISO/IEC 29110 (Word .docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "DBRS-ISO29110-DETAILED-TH.docx"

# ISO 29110 aligned requirement levels
REQ_NOTE = (
    "คำศัพท์ระดับความต้องการ (ตาม RFC 2119 / ISO/IEC/IEEE 29148):\n"
    "SHALL = ต้อง (บังคับ) · SHOULD = ควร (แนะนำ) · MAY = อาจ (ทางเลือก)"
)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "2F5496",
    font_size: int = 8,
    header_font_white: bool = True,
) -> None:
    from docx.shared import RGBColor

    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                if header_font_white:
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_req_table(doc: Document, rows: list[list[str]]) -> None:
    add_table(
        doc,
        ["รหัส", "ความต้องการ", "ระดับ", "โมดูล", "เกณฑ์ตรวจสอบ (V&V)"],
        rows,
        header_fill="2F5496",
    )


def main() -> None:
    today = date.today().isoformat()
    doc = Document()

    # === Cover ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Dashboard & Monitoring System\n"
        "Detailed Business Requirements Specification (DBRS)\n"
        "ตามแนวทาง ISO/IEC 29110"
    )
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"ลูกค้า: Pepsi-Cola (Thailand) Trading Co., Ltd.\n"
        f"ผู้พัฒนา: S.Y. Interactive Development Limited\n"
        f"ระบบ: PM-Pepsi-App (Web Application)\n"
        f"รหัสเอกสาร: PM-PEPSI-DBRS-001\n"
        f"เวอร์ชัน: 1.0\nวันที่: {today}\n"
        f"มาตรฐานอ้างอิง: ISO/IEC 29110 (VSE) · ISO/IEC/IEEE 29148"
    ).font.size = Pt(10)

    doc.add_page_break()

    # === Document control (ISO 29110 project artifact) ===
    doc.add_heading("การควบคุมเอกสาร (Document Control)", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["ชื่อเอกสาร", "Detailed Business Requirements Specification (DBRS)"],
            ["รหัส", "PM-PEPSI-DBRS-001"],
            ["เวอร์ชัน", "1.0"],
            ["วันที่มีผล", today],
            ["สถานะ", "Draft for UAT / Customer Review"],
            ["เจ้าของเอกสาร", "Project Manager / Business Analyst"],
            ["ผู้อนุมัติ (ลูกค้า)", "___________________  วันที่ _______"],
            ["ผู้อนุมัติ (ผู้พัฒนา)", "___________________  วันที่ _______"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("ประวัติการแก้ไข (Revision History)", level=2)
    add_table(
        doc,
        ["เวอร์ชัน", "วันที่", "ผู้แก้ไข", "รายละเอียด"],
        [
            ["1.0", today, "S.Y. Interactive", "ฉบับแรก — รวมความต้องการจากประชุม 1–2 และ UAT prep"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 1 Introduction ===
    doc.add_heading("1. บทนำ (Introduction)", level=1)

    doc.add_heading("1.1 วัตถุประสงค์ของเอกสาร (Purpose)", level=2)
    doc.add_paragraph(
        "เอกสาร DBRS ฉบับนี้กำหนดความต้องการทางธุรกิจ (Business Requirements) อย่างละเอียด "
        "สำหรับระบบ PM Dashboard & Monitoring — Pepsi-Cola (Thailand) "
        "เป็นเอาต์พุตของกิจกรรม Requirements Analysis ตาม ISO/IEC 29110-5-1-2 "
        "และเป็นแหล่งอ้างอิงสำหรับ Software Requirements Specification (SRS), "
        "การออกแบบ, การทดสอบ (V&V) และ UAT"
    )

    doc.add_heading("1.2 ขอบเขต (Scope)", level=2)
    doc.add_paragraph(
        "ครอบคลุมความต้องการธุรกิจของแอปพลิเคชัน PM-Pepsi-App ทั้งหมด ได้แก่ "
        "การบูรณาการ SAP (IW37N/Confirm IN/OUT), การจัดตารางงาน PM/CM, "
        "การปิดงานและ QC, รายงาน Engineering, Admin Console, Engineering Board, "
        "Telegram (ถ้ารวม go-live) และ Portal หลายโมดูล"
    )
    add_bullets(
        doc,
        [
            "In scope: Web app React + API Express + PostgreSQL",
            "In scope: การนำเข้า/ส่งออกไฟล์ Excel/CSV จาก SAP",
            "Out of scope: การเชื่อม SAP real-time API (ใช้ไฟล์เป็นสัญญา)",
            "Out of scope: แอป mobile native (ใช้ responsive web)",
            "Out of scope: โมดูลสโตร์อะไหล่/แจ้งซ่อม (Portal handoff เท่านั้น)",
        ],
    )

    doc.add_heading("1.3 คำจำกัดความและตัวย่อ (Definitions & Acronyms)", level=2)
    add_table(
        doc,
        ["คำย่อ", "ความหมาย"],
        [
            ["DBRS", "Detailed Business Requirements Specification"],
            ["SRS", "Software Requirements Specification"],
            ["VSE", "Very Small Entity — องค์กรขนาดเล็กมาก (ISO 29110)"],
            ["WO", "Work Order — ใบสั่งงาน"],
            ["IW37N", "รายงาน SAP สำหรับ Work Order list"],
            ["PM/CM", "Preventive / Corrective Maintenance"],
            ["QC", "Quality Check ก่อนนับ KPI และ export SAP"],
            ["RBAC", "Role-Based Access Control"],
            ["UAT", "User Acceptance Testing"],
            ["FL", "Functional Location"],
            ["ZB02/ZD02", "รหัสประเภทงาน PM ใน SAP"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("1.4 เอกสารอ้างอิง (References)", level=2)
    add_table(
        doc,
        ["ประเภท", "เอกสาร"],
        [
            ["มาตรฐาน", "ISO/IEC 29110-5-1-2 — Management and engineering guide for VSE"],
            ["มาตรฐาน", "ISO/IEC/IEEE 29148 — Systems and software engineering — Life cycle processes — Requirements engineering"],
            ["ภายใน", "docs/customer-requirements/MEETING-SUMMARY-REQUIREMENTS.md"],
            ["ภายใน", "docs/USER-MANUAL-TH.md"],
            ["ภายใน", "docs/PRE-UAT-MASTER-PHASES.md"],
            ["ภายใน", "docs/customer-requirements/UAT-ROUND-2-TH.md"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("1.5 ระดับความต้องการ (Requirement Conventions)", level=2)
    p = doc.add_paragraph(REQ_NOTE)
    for r in p.runs:
        r.font.size = Pt(9)

    doc.add_page_break()

    # === 2 Overall description ===
    doc.add_heading("2. ภาพรวมระบบ (Overall Description)", level=1)

    doc.add_heading("2.1 บริบทธุรกิจ (Business Context)", level=2)
    doc.add_paragraph(
        "โรงงาน Pepsi-Cola (Thailand) ใช้ SAP PM เป็นหลักสำหรับ Work Order "
        "ทีม Engineering ต้องการระบบ PM Dashboard ที่รวมการนำเข้าข้อมูล SAP, "
        "จัดตารางงาน, มอบหมายช่าง, บันทึกปิดงานพร้อมรูป, ส่งกลับ SAP "
        "และรายงาน utilization รายสัปดาห์ — แทนระบบ PHP legacy"
    )

    doc.add_heading("2.2 วัตถุประสงค์ทางธุรกิจ (Business Objectives)", level=2)
    add_table(
        doc,
        ["รหัส", "วัตถุประสงค์", "ตัวชี้วัดความสำเร็จ"],
        [
            ["BO-01", "ลดเวลา key-in ปิดงานซ้ำ", "Mass confirm ≤44 WO/ครั้ง"],
            ["BO-02", "จัดลำดับความสำคัญงาน PM ได้", "Filter/filter ปฏิทิน 5 ชนิด"],
            ["BO-03", "ข้อมูล SAP เข้าระบบได้แม่นยำ", "Import IW37N inserted > 0, error ชัด"],
            ["BO-04", "ติดตามแผน PM รายสัปดาห์", "Eng Utilization report เทียบ Excel"],
            ["BO-05", "ควบคุมคุณภาพก่อนส่ง SAP", "QC approve ก่อน export"],
            ["BO-06", "ช่างทำงานนอกรอบ SAP 07:00/19:00", "อัปโหลดมือ 24 ชม."],
            ["BO-07", "ตรวจสอบย้อนหลังได้", "Activity log + Audit trail"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("2.3 ผู้มีส่วนได้ส่วนเสีย (Stakeholders)", level=2)
    add_table(
        doc,
        ["Stakeholder", "บทบาท", "ความสนใจหลัก"],
        [
            ["Planner / Engineering", "จัดแผน PM/CM", "ปฏิทิน, assign, import SAP"],
            ["ช่าง (Technician)", "ปฏิบัติงาน", "ดูงาน, ปิดงาน, รูป, PM readings"],
            ["หัวหน้างาน / Foreman", "Supervisor close, QC", "อนุมัติ QC, mass confirm"],
            ["Admin", "ดูแลระบบ", "RBAC, backup, settings, Telegram"],
            ["HR", "รายงานชั่วโมง", "Manhour import, utilization"],
            ["Auditor", "ตรวจสอบ", "Activity log, revision"],
            ["SAP Team", "รับ/ส่งไฟล์", "CONFIRM_OUT format ตรง SAP"],
            ["S.Y. Interactive", "พัฒนาและส่งมอบ", "UAT pass, maintainability"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("2.4 ลักษณะผู้ใช้ (User Characteristics)", level=2)
    add_table(
        doc,
        ["บทบาท", "ทักษะ", "อุปกรณ์", "ภาษา UI"],
        [
            ["Planner", "คุ้น SAP PM, Excel", "PC desktop", "EN/TH"],
            ["ช่าง", "ใช้งานพื้นฐาน", "PC / tablet / Telegram", "EN/TH"],
            ["Admin", "IT พื้นฐาน", "PC", "EN/TH"],
            ["Board viewer", "ไม่ login", "TV kiosk", "TH/EN"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 3 Constraints ===
    doc.add_heading("3. ข้อจำกัด สมมติฐาน และการพึ่งพา", level=1)

    doc.add_heading("3.1 ข้อจำกัด (Constraints)", level=2)
    add_req_table(
        doc,
        [
            ["CON-01", "SHALL ใช้ไฟล์ Excel/CSV เป็นสัญญากับ SAP ไม่ใช่ API real-time", "SHALL", "Integration", "ไม่มี SAP RFC call"],
            ["CON-02", "SHALL Mass Confirm ไม่เกิน 44 WO ต่อชุด (ข้อจำกัด SAP)", "SHALL", "Confirmation", "API reject >44"],
            ["CON-03", "SHALL รูปเก็บใน PostgreSQL ไม่ใช่โฟลเดอร์ disk", "SHALL", "Images", "BYTEA ใน DB"],
            ["CON-04", "SHALL UI default ภาษาอังกฤษ สลับ TH ได้", "SHALL", "i18n", "locale toggle"],
            ["CON-05", "SHALL RBAC ควบคุมทุก action สำคัญ", "SHALL", "Security", "403 without perm"],
            ["CON-06", "SHOULD รองรับ factory FL 7151 เป็นหลัก", "SHOULD", "Import", "factory scope filter"],
        ],
    )

    doc.add_heading("3.2 สมมติฐาน (Assumptions)", level=2)
    add_bullets(
        doc,
        [
            "ลูกค้าส่งไฟล์ SAP export ตาม template ที่ตกลง (ZB02All, IW37N ALV)",
            "PostgreSQL และ Windows Server พร้อมสำหรับ production",
            "ผู้ใช้มีบัญชี work center ใน tbworkcenter",
            "รอบ SAP 07:00/19:00 ไม่เท่ากับเวลาเข้างานช่าง 7:00–15:30",
            "Telegram Bot token จัดหาโดยลูกค้า (ถ้า scope P3)",
        ],
    )

    doc.add_heading("3.3 การพึ่งพา (Dependencies)", level=2)
    add_table(
        doc,
        ["การพึ่งพา", "รายละเอียด"],
        [
            ["SAP PM", "Export IW37N, Confirm IN; รับ Confirm OUT"],
            ["PostgreSQL", "Schema app, migrations 001–103+"],
            ["Node.js / PM2", "Backend API runtime"],
            ["IIS / reverse proxy", "HTTPS, static frontend"],
            ["Telegram Bot API", "แจ้งเตือน assign (optional)"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 4 Functional requirements by domain ===
    doc.add_heading("4. ความต้องการเชิงฟังก์ชัน (Functional Requirements)", level=1)
    doc.add_paragraph(
        "จัดกลุ่มตามหัวข้อประชุมลูกค้า A–F และ ISO 29110 traceability ID: BR-FR-xxx"
    )

    doc.add_heading("4.1 กลุ่ม A — การบูรณาการข้อมูลและการจัดตาราง (Data Integration & Scheduling)", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-A01", "SHALL นำเข้า Work Order จากไฟล์ Excel IW37N (SAP ALV + legacy)", "SHALL", "/iw37n, /integration", "Import inserted > 0"],
            ["BR-FR-A02", "SHALL แสดง preview และ error ก่อน commit import", "SHALL", "IW37N", "Preview table ก่อน confirm"],
            ["BR-FR-A03", "SHALL แสดง WO บนปฏิทิน Work Scheduling หลัง import", "SHALL", "/calendar", "Events > 0 หลัง import"],
            ["BR-FR-A04", "SHALL กรองปฏิทิน: กิจกรรม, wktype, syst, wkctr, team, priority", "SHALL", "/calendar", "Filter 5 ชนิดทำงาน"],
            ["BR-FR-A05", "SHALL ย้ายวันแผน (move plan) บนปฏิทินด้วย drag-drop", "SHALL", "/calendar", "tbmoveplan อัปเดต"],
            ["BR-FR-A06", "SHOULD สแกน inbound folder อัตโนมัติ (integration watch)", "SHOULD", "/integration", "Job status done"],
            ["BR-FR-A07", "SHALL ไม่จำกัดผู้ใช้แค่รอบ SAP 07:00/19:00 — อัปโหลดมือได้", "SHALL", "/iw37n", "Upload นอกรอบสำเร็จ"],
        ],
    )

    doc.add_heading("4.2 กลุ่ม B — Import/Export และ Validation", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-B01", "SHALL นำเข้า Confirm IN จาก SAP Excel แยกจาก IW37N", "SHALL", "/integration", "Confirm IN tab"],
            ["BR-FR-B02", "SHALL ส่งออก Confirm OUT เป็น CSV/XLSX ตาม SAP", "SHALL", "/confirmation", "Export เปิดใน Excel"],
            ["BR-FR-B03", "SHALL ป้องกัน duplicate import (SHA256/idempotent)", "SHALL", "Integration", "Skip duplicate batch"],
            ["BR-FR-B04", "SHALL แสดง integration job log (running/done/failed)", "SHALL", "/integration", "Job list UI"],
            ["BR-FR-B05", "SHOULD แสดงแบนเนอร์อธิบายรอบ SAP vs เวลาทำงานช่าง", "SHOULD", "/integration", "Banner ภาษาไทย"],
        ],
    )

    doc.add_heading("4.3 กลุ่ม C — Dashboard & รายงานรายสัปดาห์", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-C01", "SHALL แสดง KPI Dashboard หน้าแรก", "SHALL", "/", "KPI cards โหลด"],
            ["BR-FR-C02", "SHALL รายงาน Eng Utilization รายสัปดาห์", "SHALL", "/summary-weekly", "เทียบ Excel ลูกค้า"],
            ["BR-FR-C03", "SHALL Engineering Board kiosk ไม่ต้อง login", "SHALL", "/board", "เปิด TV ได้"],
            ["BR-FR-C04", "SHALL แสดง PM readings บน Board โซน C", "SHALL", "/board", "กราฟ refresh 60s"],
            ["BR-FR-C05", "SHALL KPI ปิดเดือนนับเฉพาะ QC approved", "SHALL", "Dashboard API", "confirm_qc_status=approved"],
        ],
    )

    doc.add_heading("4.4 กลุ่ม D — บริหาร WO, สถานะ, ทรัพยากร", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-D01", "SHALL แยกประเภทงาน ZB01/ZB02/ZB05 และป้าย ZD01/ZD02/ZD05", "SHALL", "WO, Calendar", "Label แสดงชัด"],
            ["BR-FR-D02", "SHALL แสดงสถานะ REL / Confirm / Create (CRTD) ใน UI", "SHALL", "/calendar", "สีและป้ายสถานะ"],
            ["BR-FR-D03", "SHALL จ่ายงานให้ work center / ทีม (multi-assign)", "SHALL", "/planning", "tbplangingwork"],
            ["BR-FR-D04", "SHALL batch assign ทีม A/B/EE/UT ทั้งหน้า WO", "SHALL", "/work-orders", "PATCH team/batch"],
            ["BR-FR-D05", "SHALL แสดง Available hour ต่อช่างใน dialog จ่ายงาน", "SHALL", "Planning dialog", "HR prorate − assigned"],
            ["BR-FR-D06", "SHALL WO modal 3 แท็บ: Task / Planning / Close", "SHALL", "WO modal", "แท็บครบ"],
        ],
    )

    doc.add_heading("4.5 กลุ่ม E — Mass Confirm, Logging, Analytics", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-E01", "SHALL Mass Confirm สูงสุด 44 WO ต่อชุด", "SHALL", "/confirmation", "API limit 44"],
            ["BR-FR-E02", "SHALL บันทึก Activity Log การกระทำสำคัญ", "SHALL", "/activity-log", "Audit entries"],
            ["BR-FR-E03", "SHALL Admin QC อนุมัติ/ปฏิเสธก่อน export และ KPI", "SHALL", "/confirmation", "QC queue"],
            ["BR-FR-E04", "SHALL บันทึกรูปหลัง PM ใน DB (WebP BYTEA)", "SHALL", "WO Confirm", "tbconfirmimg"],
            ["BR-FR-E05", "SHOULD Mass Confirm Export Panel หลัง batch", "SHOULD", "/confirmation", "Export ชุดเดียวกัน"],
        ],
    )

    doc.add_heading("4.6 กลุ่ม F — Audit & ปรับปรุงแผน", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-F01", "SHALL Auditor Hub สำหรับตรวจสอบ", "SHALL", "/reports/audit", "หน้าโหลด"],
            ["BR-FR-F02", "SHALL Resource revision log สำหรับ master data", "SHALL", "Admin", "tbl_resource_revision"],
            ["BR-FR-F03", "SHOULD Comment ชดเชย SAP บน WO", "SHOULD", "WO modal", "PM comment thread"],
        ],
    )

    doc.add_heading("4.7 PM Execution & ค่าวัด", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-PM01", "SHALL บันทึก PM comment ต่อ WO", "SHALL", "WO Task tab", "tbwo_pm_note"],
            ["BR-FR-PM02", "SHALL บันทึกค่าวัดกระแส 3 เฟส R/S/T", "SHALL", "/pm-vibration", "tbwo_pm_reading"],
            ["BR-FR-PM03", "SHALL แสดงกราฟ trend ค่าวัดใน WO modal", "SHALL", "WO Task", "Chart render"],
            ["BR-FR-PM04", "SHOULD Import/Export Excel ค่าวัด PM", "SHOULD", "/pm-vibration", "xlsx download"],
            ["BR-FR-PM05", "SHALL ฟอร์ม PM สอดคล้องกระดาษ WO 4001565681", "SHALL", "/pm-vibration", "UAT field-by-field"],
        ],
    )

    doc.add_heading("4.8 Admin, Security, Portal", level=2)
    add_req_table(
        doc,
        [
            ["BR-FR-AD01", "SHALL Admin Console จัดการ users, roles, menu, settings", "SHALL", "/admin/*", "13 modules"],
            ["BR-FR-AD02", "SHALL Backup/Restore PostgreSQL ผ่าน UI", "SHALL", "/admin/backup", "pg_dump"],
            ["BR-FR-AD03", "SHALL RBAC permission matrix แก้ได้", "SHALL", "/admin/roles", "Matrix UI"],
            ["BR-FR-AD04", "SHOULD Portal เลือก module หลัง login", "SHOULD", "/portal", "Module cards"],
            ["BR-FR-AD05", "SHOULD Telegram แจ้ง assign + รับทราบ", "SHOULD", "Telegram", "Webhook flow"],
        ],
    )

    doc.add_page_break()

    # === 5 Non-functional ===
    doc.add_heading("5. ความต้องการเชิงคุณลักษณะ (Non-Functional Requirements)", level=1)
    add_table(
        doc,
        ["รหัส", "หมวด", "ความต้องการ", "ระดับ", "เกณฑ์วัด"],
        [
            ["NFR-01", "Performance", "API p95 calendar/events < 800ms (local)", "SHOULD", "stress-api.ts"],
            ["NFR-02", "Performance", "หน้าหลักโหลด < 3s บน LAN", "SHOULD", "Lighthouse/manual"],
            ["NFR-03", "Security", "Session JWT httpOnly + RBAC ทุก route", "SHALL", "Security review"],
            ["NFR-04", "Security", "Login rate limit + IP block", "SHALL", "tbl_blocked_ip"],
            ["NFR-05", "Availability", "Maintenance mode แสดง banner + 503 mutate", "SHALL", "Maintenance test"],
            ["NFR-06", "Usability", "Responsive sidebar drawer บน mobile", "SHALL", "Viewport E2E"],
            ["NFR-07", "Usability", "UI EN default + TH toggle", "SHALL", "Locale E2E"],
            ["NFR-08", "Reliability", "Backup อัตโนมัติตาม cron", "SHOULD", "tbl_backup_history"],
            ["NFR-09", "Maintainability", "Unit test backend 311+ / frontend 147+", "SHOULD", "npm test pass"],
            ["NFR-10", "Compatibility", "Chrome/Edge latest 2 versions", "SHALL", "UAT browser matrix"],
            ["NFR-11", "Data", "รูป WebP ใน DB ไม่พึ่ง path ไฟล์", "SHALL", "No disk img path prod"],
            ["NFR-12", "Auditability", "Audit log ทุก admin write", "SHALL", "tbl_audit_log"],
        ],
        header_fill="7030A0",
    )

    doc.add_page_break()

    # === 6 Interface requirements ===
    doc.add_heading("6. ความต้องการอินเทอร์เฟซ (Interface Requirements)", level=1)

    doc.add_heading("6.1 User Interface", level=2)
    add_table(
        doc,
        ["หน้าหลัก", "Route", "ผู้ใช้หลัก"],
        [
            ["Dashboard", "/", "Planner"],
            ["ปฏิทิน WO", "/calendar", "Planner, ช่าง"],
            ["Work Orders", "/work-orders", "ทุก role"],
            ["Planning", "/planning", "Planner"],
            ["Confirmation", "/confirmation", "Foreman, Admin"],
            ["Integration", "/integration", "Admin"],
            ["Eng Utilization", "/summary-weekly", "HR, Planner"],
            ["Admin Console", "/admin", "Admin"],
            ["Engineering Board", "/board", "Public kiosk"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )

    doc.add_heading("6.2 Software Interfaces", level=2)
    add_table(
        doc,
        ["อินเทอร์เฟซ", "โปรโตคอล", "รายละเอียด"],
        [
            ["PM-Pepsi-App API", "HTTPS REST JSON", "/api/v1/* — ดู BACKEND-DESIGN doc"],
            ["SAP IW37N", "Excel .xlsx ALV", "Import inbound"],
            ["SAP Confirm IN", "Excel .xls ALV", "Import inbound/confirm"],
            ["SAP Confirm OUT", "CSV/XLSX", "Export outbound"],
            ["Telegram Bot", "HTTPS webhook", "Assign notify, ack"],
            ["PostgreSQL", "TCP 5432", "DATABASE_URL"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 7 Data requirements ===
    doc.add_heading("7. ความต้องการข้อมูล (Data Requirements)", level=1)
    doc.add_paragraph(
        "ข้อมูลหลักอยู่ใน PostgreSQL schema app — ศูนย์กลาง tbiw37n (Work Order). "
        "รายละเอียดเต็มดู DATABASE-DESIGN-DETAILED-TH.docx และ ER-DIAGRAM-DETAILED-TH.docx"
    )
    add_table(
        doc,
        ["หมวดข้อมูล", "ตารางหลัก", "Retention"],
        [
            ["Work Order", "tbiw37n, tbmoveplan", "ตามนโยบายลูกค้า"],
            ["Planning", "tbplangingwork", "ตาม WO"],
            ["Confirmation", "tbcofirm, tbconfirmimg, tbwrkclose", "รูปใน DB"],
            ["Manhour", "tbmanhours", "รายปี"],
            ["Master", "tbzone, tbworkcenter, …", "Admin แก้ได้"],
            ["Audit", "tbl_audit_log", "ตาม audit retention setting"],
            ["Settings", "tbl_setting", "Persistent"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 8 Business rules ===
    doc.add_heading("8. กฎธุรกิจ (Business Rules)", level=1)
    add_table(
        doc,
        ["รหัส", "กฎ", "ผลเมื่อฝ่าฝืน"],
        [
            ["BR-BUS-01", "WO สถานะ TECO/ปิดแล้ว ห้าม move ปฏิทิน", "API 409"],
            ["BR-BUS-02", "Export SAP เฉพาะ WO ที่ QC approved", "ไม่มีแถวใน export"],
            ["BR-BUS-03", "Mass confirm สูงสุด 44 รายการ", "Validation error"],
            ["BR-BUS-04", "รูป confirm รับเฉพาะ phase=after (หลัง PM)", "Reject before"],
            ["BR-BUS-05", "ทีม A/B/EE/UT แยกจากการจ่ายงาน individual", "UI แสดงถูกทีม"],
            ["BR-BUS-06", "Factory scope FL 7151 สำหรับ import default", "Skip แถวนอก scope"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 9 Traceability (ISO 29110 key) ===
    doc.add_heading("9. การติดตามความต้องการ (Requirements Traceability)", level=1)
    doc.add_paragraph(
        "ตาม ISO/IEC 29110 ความต้องการธุรกิจ (BR-FR) ต้อง trace ไปยัง "
        "Software Requirements, Design, และ Test — ตารางด้านล่างเป็นระดับ Business → Module → UAT"
    )
    add_table(
        doc,
        ["Business Need", "BR-FR", "Module/Route", "UAT Flow"],
        [
            ["Import SAP WO", "BR-FR-A01", "/integration, /iw37n", "P0, P1-A2"],
            ["ปฏิทิน + filter", "BR-FR-A03,A04", "/calendar", "P1-A2"],
            ["จ่ายงาน", "BR-FR-D03", "/planning", "P1-A4, A5"],
            ["ปิดงาน + รูป", "BR-FR-E04", "WO modal", "P1-A8"],
            ["QC + Export", "BR-FR-E03, B02", "/confirmation", "P1-A9"],
            ["Mass 44", "BR-FR-E01", "/personnel/confirm", "P1-A7"],
            ["Eng Report", "BR-FR-C02", "/summary-weekly", "P1-A10"],
            ["PM ค่าวัด", "BR-FR-PM01-05", "/pm-vibration", "P2"],
            ["Admin RBAC", "BR-FR-AD01-03", "/admin", "P5"],
            ["Telegram", "BR-FR-AD05", "Telegram", "P3"],
        ],
        header_fill="2F5496",
    )

    doc.add_page_break()

    # === 10 V&V and Acceptance ===
    doc.add_heading("10. การตรวจสอบและการยอมรับ (Verification & Validation)", level=1)

    doc.add_heading("10.1 วิธีการตรวจสอบ (ISO 29110 V&V Methods)", level=2)
    add_table(
        doc,
        ["วิธี", "ใช้กับ", "เครื่องมือ"],
        [
            ["Review", "DBRS, SRS, Design docs", "Customer sign-off"],
            ["Unit Test", "Business logic, parsers", "Vitest 458+ cases"],
            ["Integration Test", "API routes", "Supertest"],
            ["System Test", "E2E flows", "Playwright 76+ routes"],
            ["Performance Test", "API load", "stress-api.ts"],
            ["UAT", "Business acceptance", "UAT-ROUND-2-TH.md"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("10.2 เกณฑ์การยอมรับระบบ (Acceptance Criteria — Master Gate)", level=2)
    add_bullets(
        doc,
        [
            "P0: Import IW37N + Confirm IN คู่กันได้ · ปฏิทิน/WO ไม่ว่างเมื่อมีข้อมูล",
            "P1: Flow A1–A10 ทีม dev ทดสอบผ่านบน env เดียวกับลูกค้า",
            "P2: PM ฟอร์ม WO 4001565681 บันทึก/แสดงครบ (ถ้ารวม scope)",
            "P4: UI Gate U4 ผ่าน (console, viewport, locale, portal)",
            "P6: HTTPS production · login · health · backup",
            "ลูกค้าลงนาม UAT Round ผ่านตามชีต",
        ],
    )

    doc.add_heading("10.3 รายการ UAT Flow อ้างอิง (P1)", level=2)
    add_table(
        doc,
        ["Flow", "รายละเอียด", "Route"],
        [
            ["A1", "Login · EN/TH · RBAC menu", "/login"],
            ["A2", "Import IW37N → ปฏิทินมี event", "/integration, /calendar"],
            ["A3", "ค้นหา WO → modal Task", "/work-orders"],
            ["A4", "จ่ายงาน manual + batch", "/planning"],
            ["A5", "ช่างเห็นงาน", "/planning, /plan-calendar"],
            ["A6", "Bulk team A/B", "/work-orders"],
            ["A7", "Mass confirm ≤44", "/personnel/confirm"],
            ["A8", "รูปหลัง PM + ปิดงาน", "WO modal Confirm"],
            ["A9", "QC approve → dashboard", "/confirmation"],
            ["A10", "Eng Utilization / Board", "/summary-weekly, /board"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    # === 11 Appendices ===
    doc.add_heading("11. ภาคผนวก (Appendices)", level=1)

    doc.add_heading("11.1 ภาคผนวก A — เอกสารที่เกี่ยวข้อง", level=2)
    add_bullets(
        doc,
        [
            "DBRS-ISO29110-DETAILED-TH.docx (เอกสารนี้)",
            "DATABASE-DESIGN-DETAILED-TH.docx",
            "BACKEND-DESIGN-DETAILED-TH.docx",
            "FRONTEND-DESIGN-DETAILED-TH.docx",
            "FLOW-DIAGRAM-SWIMLANE-TH.docx",
            "UNIT-TEST-DETAILED-TH.docx / STRESS-TEST-DETAILED-TH.docx",
            "USER-MANUAL-TH.md",
            "UAT-ROUND-2-TH.md",
        ],
    )

    doc.add_heading("11.2 ภาคผนวก B — ISO/IEC 29110 Mapping", level=2)
    doc.add_paragraph(
        "เอกสาร DBRS นี้สอดคล้องกับกิจกรรม ISO/IEC 29110-5-1-2 ดังนี้:"
    )
    add_table(
        doc,
        ["กิจกรรม ISO 29110", "เอาต์พุตในโปรเจกต์"],
        [
            ["Requirements elicitation", "MEETING-SUMMARY, ประชุมลูกค้า 1–2"],
            ["Customer requirements specification", "DBRS (เอกสารนี้)"],
            ["Software requirements specification", "SRS แยกใน Design docs + API catalog"],
            ["Traceability", "§9 Traceability matrix"],
            ["Verification", "Unit/E2E/Stress test docs"],
            ["Validation", "UAT-ROUND-2-TH.md"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("11.3 ภาคผนวก C — ลายเซ็นอนุมัติ", level=2)
    add_table(
        doc,
        ["บทบาท", "ชื่อ", "ลายเซ็น", "วันที่"],
        [
            ["ลูกค้า — Engineering Manager", "", "", ""],
            ["ลูกค้า — IT / SAP", "", "", ""],
            ["ผู้พัฒนา — Project Manager", "", "", ""],
            ["ผู้พัฒนา — Technical Lead", "", "", ""],
        ],
        header_fill="FFFFFF",
        header_font_white=False,
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร DBRS —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
