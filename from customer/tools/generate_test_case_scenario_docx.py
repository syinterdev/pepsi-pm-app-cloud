#!/usr/bin/env python3
"""Generate PM Pepsi App — Test Case Scenario document (Word .docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "TEST-CASE-SCENARIO-DETAILED-TH.docx"

# (id, module, title, priority, actor, precond, steps, expected, type, trace)
SCENARIOS: list[tuple[str, ...]] = [
    # AUTH
    (
        "TC-AUTH-001",
        "Authentication",
        "Login สำเร็จด้วย work center ที่มีสิทธิ์",
        "P0",
        "Admin / Planner / ช่าง",
        "มีบัญชีใน tbworkcenter · backend + DB พร้อม",
        "1) เปิด /login\n2) กรอก username/password\n3) กด Login",
        "เข้า Dashboard หรือ /portal · ไม่มี error · session cookie ตั้งค่า",
        "Manual / E2E",
        "UAT-A1 · BR-FR-AD01",
    ),
    (
        "TC-AUTH-002",
        "Authentication",
        "Login ล้มเหลว — รหัสผ่านผิด",
        "P1",
        "Guest",
        "—",
        "1) กรอก password ผิด\n2) กด Login",
        "แสดงข้อความ error · ไม่เข้าระบบ · audit failed login (ถ้ามี)",
        "Manual / API",
        "Security · api-smoke",
    ),
    (
        "TC-AUTH-003",
        "Authentication",
        "ผู้ไม่ login เข้า /work-orders ถูก redirect",
        "P0",
        "Guest",
        "—",
        "1) เปิด /work-orders โดยไม่ login",
        "redirect ไป /login",
        "E2E",
        "app-smoke",
    ),
    (
        "TC-AUTH-004",
        "RBAC",
        "ช่างไม่เห็นเมนู Admin",
        "P0",
        "ช่าง (W)",
        "Login ช่างที่ไม่มี admin.*",
        "1) Login ช่าง\n2) ดู sidebar",
        "ไม่มีลิงก์ /admin · เรียก API admin ได้ 403",
        "Manual / E2E",
        "UAT-I8 · BR-FR-AD03",
    ),
    (
        "TC-AUTH-005",
        "i18n",
        "สลับภาษา EN/TH หลัง login",
        "P1",
        "Planner",
        "Login แล้ว",
        "1) กด Globe แถบบน → EN\n2) เปิด /planning\n3) สลับ TH",
        "ข้อความ UI เปลี่ยนตาม locale · ไม่ crash",
        "E2E",
        "UAT-A1 · u4e-locale",
    ),
    # INTEGRATION
    (
        "TC-INT-001",
        "Integration SAP",
        "Import IW37N ไฟล์ล่าสุด — preview แล้ว commit",
        "P0",
        "Planner / Admin",
        "มีไฟล์ IW37N .xlsx จาก SAP · migration ครบ",
        "1) เปิด /integration หรือ /iw37n\n2) อัปโหลดไฟล์\n3) ดู preview OK/error\n4) Confirm import",
        "inserted > 0 · แสดง batch id · ไม่ skipped ทั้งก้อน",
        "Manual",
        "UAT-A2 · BR-FR-A01 · P0-I6",
    ),
    (
        "TC-INT-002",
        "Integration SAP",
        "Import ไฟล์ซ้ำ (SHA256 เดิม) — skip",
        "P1",
        "Admin",
        "เคย import ไฟล์เดิมแล้ว",
        "1) อัปโหลดไฟล์เดิมซ้ำ",
        "ระบบแจ้ง duplicate/skip · ไม่ทำข้อมูลซ้ำ",
        "Manual / Unit",
        "BR-FR-B03",
    ),
    (
        "TC-INT-003",
        "Integration SAP",
        "Confirm IN preview + import",
        "P0",
        "Admin",
        "มีไฟล์ Confirm IN คู่ batch ZB02",
        "1) /integration แท็บ Confirm IN\n2) Preview\n3) Import",
        "inserted > 0 · ข้อมูลผูก WO",
        "Manual",
        "UAT-E4 · P0-I7",
    ),
    (
        "TC-INT-004",
        "Integration SAP",
        "อัปโหลด IW37N นอกรอบ SAP 07:00/19:00",
        "P1",
        "Planner",
        "—",
        "1) อัปโหลด manual นอกรอบ\n2) ตรวจ job log",
        "สำเร็จ · แบนเนอร์อธิบายรอบ SAP vs เวลาทำงานช่าง",
        "Manual",
        "UAT-E1,E2 · BR-FR-A07",
    ),
    (
        "TC-INT-005",
        "Integration SAP",
        "Export Confirm OUT หลัง QC",
        "P0",
        "Admin / Foreman",
        "มี WO ที่ QC approved",
        "1) /confirmation\n2) Export CSV/XLSX",
        "ไฟล์เปิด Excel · คอลัมน์ตรง SAP · เฉพาะ approved",
        "Manual",
        "UAT-A9 · BR-FR-B02",
    ),
    # CALENDAR & WO
    (
        "TC-CAL-001",
        "Calendar",
        "หลัง import — ปฏิทินแสดง event WO",
        "P0",
        "Planner",
        "Import IW37N สำเร็จ",
        "1) เปิด /calendar\n2) เลือกปี/เดือนตรง bscstart",
        "มี event บนปฏิทิน · คลิกดูรายละเอียดได้",
        "Manual / E2E",
        "UAT-A2 · BR-FR-A03",
    ),
    (
        "TC-CAL-002",
        "Calendar",
        "กรอง Z1/Z2/Z5 · ทีม · สถานะ — ข้อมูลไม่หาย",
        "P0",
        "Planner",
        "มี WO หลายประเภท",
        "1) ตั้ง filter หลายค่า\n2) ล้าง filter",
        "จำนวน event สอดคล้อง · ไม่หายทั้งก้อนเมื่อ filter",
        "Manual",
        "UAT-G7 · I1 regression",
    ),
    (
        "TC-CAL-003",
        "Calendar",
        "ลากย้ายวันแผน (move plan)",
        "P1",
        "Planner",
        "WO สถานะ REL · มีสิทธิ calendar write",
        "1) ลาก event ไปวันอื่น\n2) Reload",
        "วันแผนเปลี่ยน · tbmoveplan อัปเดต",
        "Manual",
        "BR-FR-A05",
    ),
    (
        "TC-CAL-004",
        "Calendar",
        "งาน TECO/ปิดแล้ว — ห้าม move",
        "P0",
        "Planner",
        "WO สถานะปิด/TECO",
        "1) พยายามลาก event",
        "ลากไม่ได้หรือ API 409 · แผนไม่เปลี่ยน",
        "Manual",
        "UAT-C6 · I2 regression",
    ),
    (
        "TC-WO-001",
        "Work Orders",
        "ค้นหา WO และเปิด modal 3 แท็บ",
        "P0",
        "Planner",
        "มี WO ใน DB",
        "1) /work-orders ค้นหาเลข WO\n2) เปิด modal\n3) สลับ Task / Planning / Close",
        "แท็บครบ · ข้อมูลโหลด · ไม่ crash",
        "Manual / E2E",
        "UAT-A3",
    ),
    (
        "TC-WO-002",
        "Work Orders",
        "Bulk ตั้งทีม A/B/EE/UT ทั้งหน้า",
        "P1",
        "Planner",
        "มีหลาย WO ในตาราง",
        "1) เลือกหลายแถว\n2) ตั้งทีม batch\n3) Save",
        "ทีมอัปเดตทุกแถว · toast สรุป",
        "Manual",
        "UAT-A6 · BR-FR-D04",
    ),
    (
        "TC-WO-003",
        "Work Orders",
        "แสดงป้าย ZB02/ZD02 ชัดเจน",
        "P1",
        "Planner",
        "มี WO ZB02",
        "1) ดูคอลัมน์/ป้าย wktype",
        "แสดง ZB02 · ZD02 ตาม mapping",
        "Manual",
        "BR-FR-D01",
    ),
    # PLANNING
    (
        "TC-PLN-001",
        "Planning",
        "จ่ายงานรายคน — ช่างเห็นบน plan-calendar",
        "P0",
        "Planner + ช่าง",
        "migration 038,098 · WO มี bscstart",
        "1) Planner จ่ายงาน Manual ที่ /planning\n2) ช่าง login ดู /plan-calendar (เดือนถูก)",
        "ช่างเห็น WO ที่จ่าย",
        "Manual",
        "UAT-A4,A5 · TC-PLN",
    ),
    (
        "TC-PLN-002",
        "Planning",
        "จ่ายงาน batch หลายคนครั้งเดียว",
        "P1",
        "Planner",
        "WO เปิดอยู่",
        "1) เปิด dialog จ่ายงาน\n2) ติ๊กหลายช่าง\n3) เพิ่ม Assignee (N)",
        "ทุกคนใน tbplangingwork · dialog ไม่ปิดทันที",
        "Manual",
        "UAT-A4 · BR-FR-D03",
    ),
    (
        "TC-PLN-003",
        "Planning",
        "Dialog แสดง Available hour ต่อช่าง",
        "P2",
        "Planner",
        "มี manhour HR data",
        "1) เปิด dialog จ่ายงาน\n2) ดูคอลัมน์ Available hour",
        "แสดงชม.คงเหลือหลังหัก assigned",
        "Manual",
        "UAT-C4 · BR-FR-D05",
    ),
    (
        "TC-PLN-004",
        "Planning",
        "ทีม A/B แยกจากการจ่าย work center",
        "P1",
        "Planner",
        "—",
        "1) ตั้งทีม A ที่ /work-orders\n2) จ่ายช่างที่ /planning",
        "ทีมกับ assignee แสดงถูกต้อง ไม่สับสน",
        "Manual",
        "UAT-C5",
    ),
    # CONFIRMATION
    (
        "TC-CFM-001",
        "Confirmation",
        "อัปโหลดรูปหลัง PM เท่านั้น",
        "P0",
        "ช่าง",
        "WO เปิด confirm · มีสิทธิ confirmation.write",
        "1) WO modal แท็บ Close\n2) อัปโหลดรูปหลัง PM\n3) ไม่มีตัวเลือก Before",
        "รูปบันทึกใน DB · แสดงหลัง reload",
        "Manual",
        "UAT-A8 · BR-FR-E04",
    ),
    (
        "TC-CFM-002",
        "Confirmation",
        "Mass confirm ≤44 รายการ",
        "P0",
        "Foreman",
        "มี WO พร้อม confirm",
        "1) /personnel/confirm เลือก ≤44\n2) Mass confirm",
        "สำเร็จ · toast เดียว · ถ้า >44 reject",
        "Manual",
        "UAT-A7 · BR-FR-E01",
    ),
    (
        "TC-CFM-003",
        "Confirmation",
        "QC approve ก่อน export และ KPI",
        "P0",
        "Admin / Foreman",
        "มี WO รอ QC",
        "1) /confirmation แท็บรอ QC\n2) Approve\n3) ตรวจ Dashboard KPI",
        "สถานะ approved · export มีแถว · KPI นับ",
        "Manual",
        "UAT-A9 · BR-FR-E03",
    ),
    (
        "TC-CFM-004",
        "Confirmation",
        "WO ยังไม่ QC — ไม่อยู่ใน export",
        "P0",
        "Admin",
        "มี WO confirm แต่ยัง pending QC",
        "1) Export Confirm OUT",
        "แถว WO นั้นไม่อยู่ในไฟล์",
        "Manual",
        "UAT-D7",
    ),
    (
        "TC-CFM-005",
        "Confirmation",
        "หลัง confirm เห็นรูป/เวลาใน WO modal",
        "P0",
        "Planner",
        "ปิดงานแล้ว",
        "1) เปิด WO modal แท็บ Close",
        "แสดงรูป · เวลา · comment",
        "Manual",
        "I3 regression",
    ),
    # PM MANUAL
    (
        "TC-PM-001",
        "PM Manual",
        "กรอกกระแส 3 เฟส R/S/T — WO 4001565681",
        "P0",
        "ช่าง",
        "WO 4001565681 ในระบบ · confirmation.write",
        "1) /pm-vibration เลือก WO\n2) กรอก R/S/T 3 จุด\n3) บันทึก\n4) Reload",
        "ค่าคงอยู่ · WO modal Task แสดงตรงกัน",
        "Manual",
        "UAT-B2 · BR-FR-PM02",
    ),
    (
        "TC-PM-002",
        "PM Manual",
        "กราฟ trend กระแสอัปเดตหลังบันทึกตาราง",
        "P1",
        "ช่าง",
        "WO เลือกแล้ว",
        "1) เพิ่มแถว Time+R/S/T\n2) บันทึก",
        "กราฟ 3 สีเปลี่ยน",
        "Manual",
        "UAT-B3",
    ),
    (
        "TC-PM-003",
        "PM Manual",
        "Comments หน้า 2 ตรง WO modal",
        "P1",
        "ช่าง",
        "—",
        "1) กรอก Comments หน้า 2\n2) บันทึก\n3) เปิด WO modal Task",
        "ข้อความเดียวกัน",
        "Manual",
        "UAT-B5",
    ),
    (
        "TC-PM-004",
        "PM Manual",
        "Import/Export Excel ค่าวัด",
        "P2",
        "Planner",
        "มี template PM",
        "1) Download template\n2) กรอก · upload\n3) Export xlsx",
        "ค่าใน WO · กราฟแสดง",
        "Manual",
        "UAT-B6 · BR-FR-PM04",
    ),
    (
        "TC-PM-005",
        "PM Manual",
        "Layout หน้า 1–2 เทียบกระดาษ SAP",
        "P1",
        "Planner + ช่าง",
        "มี print WO 4001565681",
        "1) เปรียบฟิลด์ทีละรายการกับกระดาษ",
        "ฟิลด์สำคัญครบ · อ่านอย่างเดียวตรง SAP",
        "Manual UAT",
        "UAT-B1 · BR-FR-PM05",
    ),
    # REPORTS
    (
        "TC-RPT-001",
        "Reports",
        "Eng Utilization รายสัปดาห์",
        "P1",
        "Planner / HR",
        "มีข้อมูล manhour + confirm",
        "1) /summary-weekly เลือกสัปดาห์\n2) เทียบ Excel ลูกค้า",
        "ตัวเลขสมเหตุสมผล · export ได้",
        "Manual",
        "UAT-A10 · BR-FR-C02",
    ),
    (
        "TC-RPT-002",
        "Engineering Board",
        "Kiosk /board ไม่ต้อง login",
        "P1",
        "Guest",
        "—",
        "1) เปิด /board ใน browser/TV",
        "โหลด KPI + PM readings · รีเฟรช",
        "E2E",
        "UAT-A10 · BR-FR-C03",
    ),
    (
        "TC-RPT-003",
        "Dashboard",
        "KPI นับเฉพาะ QC approved",
        "P1",
        "Planner",
        "มี WO approved/pending",
        "1) ดู KPI ปิดเดือน",
        "ไม่นับ pending QC",
        "Manual",
        "BR-FR-C05",
    ),
    (
        "TC-RPT-004",
        "Activity Log",
        "บันทึกการกระทำสำคัญ",
        "P2",
        "Admin",
        "ทำ action import/assign",
        "1) /activity-log หรือ /admin/audit",
        "มีแถว audit ตรง action",
        "Manual",
        "BR-FR-E02",
    ),
    # ADMIN
    (
        "TC-ADM-001",
        "Admin",
        "จัดการ users + อัปโหลดรูปช่าง",
        "P1",
        "Admin",
        "admin.users.write",
        "1) /admin/users\n2) อัปโหลดรูป\n3) ดูรายงาน",
        "รูปใน DB · แสดงบน utilization",
        "Manual",
        "UAT-H7",
    ),
    (
        "TC-ADM-002",
        "Admin",
        "Backup manual + ประวัติ",
        "P1",
        "Admin",
        "—",
        "1) /admin/backup Create\n2) ดู history",
        "ไฟล์ .sql.gz · status success",
        "Manual",
        "Runbook · BR-FR-AD02",
    ),
    (
        "TC-ADM-003",
        "Admin",
        "RBAC matrix แก้ permission",
        "P2",
        "Admin",
        "—",
        "1) /admin/roles แก้ matrix\n2) Login role ที่เกี่ยวข้อง",
        "เมนู/ปุ่มตรงสิทธิ์ใหม่",
        "Manual",
        "BR-FR-AD03",
    ),
    (
        "TC-ADM-004",
        "Admin Security",
        "Block IP หลัง brute-force",
        "P2",
        "Admin",
        "มี failed login",
        "1) /admin/security\n2) Block IP\n3) ลอง login จาก IP นั้น",
        "ถูกบล็อก · audit บันทึก",
        "Manual",
        "Security doc",
    ),
    # TELEGRAM
    (
        "TC-TG-001",
        "Telegram",
        "Admin สร้างกลุ่ม + ทดสอบส่ง",
        "P2",
        "Admin",
        "scope Telegram รวม go-live",
        "1) /admin/telegram สร้างกลุ่ม\n2) Test send",
        "ข้อความถึงกลุ่ม TG",
        "Manual",
        "UAT-F1",
    ),
    (
        "TC-TG-002",
        "Telegram",
        "จ่ายงาน → ช่างได้ DM → รับทราบ",
        "P2",
        "Planner + ช่าง",
        "ช่างผูก Telegram แล้ว",
        "1) Planner จ่ายงาน\n2) ช่างกดรับทราบใน TG\n3) Planner ได้แจ้ง",
        "ack บันทึก · แจ้งกลุ่ม Planner",
        "Manual",
        "UAT-F3,F4",
    ),
    (
        "TC-TG-003",
        "Telegram",
        "ปิดงานในแชท — รูปหลัง PM + comment",
        "P2",
        "ช่าง",
        "WO จ่ายแล้ว",
        "1) ส่งรูปหลัง PM ใน TG\n2) ส่ง comment\n3) ดู /confirmation",
        "ข้อมูลปรากฏบนเว็บ",
        "Manual",
        "UAT-F5,F6",
    ),
    # PORTAL & UI
    (
        "TC-PRT-001",
        "Portal",
        "Login → เลือกโมดูล PM",
        "P1",
        "Admin",
        "PORTAL_ENABLED",
        "1) Login\n2) /portal กดการ์ด PM",
        "เข้าแอป PM · เมนูตาม role",
        "E2E",
        "UAT-G1 · u4f-portal",
    ),
    (
        "TC-PRT-002",
        "Portal",
        "โมดูลยังไม่พร้อม — ไม่ crash",
        "P2",
        "Admin",
        "Store/Repair ready=false",
        "1) คลิกการ์ด coming soon",
        "แสดงสถานะชัด · ไม่พาไปหน้าว่าง",
        "Manual / E2E",
        "UAT-G2",
    ),
    (
        "TC-UI-001",
        "UI Regression",
        "ทุก route ไม่ crash / ไม่ console error",
        "P0",
        "Admin",
        "E2E env",
        "1) npm run test:e2e:all-routes",
        "76/76 passed (EN+TH)",
        "E2E Auto",
        "U4 Gate · E2E doc",
    ),
    (
        "TC-UI-002",
        "UI Regression",
        "ไม่มีข้อความ (mock) บน production UI",
        "P1",
        "QA",
        "build production",
        "1) npm run audit:ui",
        "ผ่าน · ไม่มี (mock)",
        "Auto",
        "P4 Gate",
    ),
    (
        "TC-UI-003",
        "UI Regression",
        "ปฏิทินไม่แสดงสรุป WO/Hrs ใต้ช่องวัน",
        "P2",
        "Planner",
        "—",
        "1) เปิด /calendar month view",
        "ไม่มี {n} WO · {h} Hrs ใต้วัน",
        "Manual",
        "UAT-G6",
    ),
    # REGRESSION
    (
        "TC-REG-001",
        "Regression",
        "แก้ Task list แล้วไม่เด้งหน้าแรก",
        "P1",
        "Planner",
        "WO มี task list",
        "1) แก้ task ใน modal\n2) Save",
        "ยังอยู่หน้าเดิม · ข้อมูลบันทึก",
        "Manual",
        "I4 regression",
    ),
    (
        "TC-REG-002",
        "Regression",
        "Factory FL 7151 — WO แสดงบนปฏิทิน",
        "P0",
        "Planner",
        "Import ไฟล์โรงงาน",
        "1) Import IW37N\n2) ตรวจปฏิทิน",
        "WO 7151 แสดง",
        "Manual",
        "P0-I7",
    ),
]

# End-to-end scenario chains (business flows)
E2E_CHAINS = [
    (
        "SC-E2E-01",
        "SAP → ปฏิทิน → จ่ายงาน → ปิดงาน → QC → Export",
        "P0",
        "Planner + ช่าง + Foreman",
        "TC-INT-001 → TC-CAL-001 → TC-PLN-001 → TC-CFM-001 → TC-CFM-003 → TC-INT-005",
        "Flow ธุรกิจครบวงจร go-live",
        "UAT A1–A10",
    ),
    (
        "SC-E2E-02",
        "PM Manual กระดาษ → WO modal → Board",
        "P1",
        "ช่าง + Planner",
        "TC-PM-001 → TC-PM-003 → TC-RPT-002",
        "ค่าวัด PM ปรากฏทุกจุด",
        "UAT §B",
    ),
    (
        "SC-E2E-03",
        "Telegram assign → ack → close → Confirmation",
        "P2",
        "Planner + ช่าง",
        "TC-TG-002 → TC-TG-003 → TC-CFM-003",
        "ถ้า scope Telegram รวม",
        "UAT §F",
    ),
    (
        "SC-E2E-04",
        "Portal → PM app → Admin backup",
        "P2",
        "Admin",
        "TC-PRT-001 → TC-ADM-002",
        "ส่งมอบ IT",
        "P6/P7",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "7B3F00",
    font_size: int = 8,
    header_font_white: bool = True,
) -> None:
    from docx.shared import RGBColor

    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                if header_font_white:
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_scenario_detail(doc: Document, s: tuple[str, ...]) -> None:
    tc_id, module, title, priority, actor, precond, steps, expected, ttype, trace = s
    doc.add_heading(f"{tc_id} — {title}", level=3)
    add_table(
        doc,
        ["ฟิลด์", "ค่า"],
        [
            ["รหัส", tc_id],
            ["โมดูล", module],
            ["ความสำคัญ", priority],
            ["ผู้ทดสอบ", actor],
            ["ประเภททดสอบ", ttype],
            ["Traceability", trace],
            ["ผลทดสอบ", "☐ ผ่าน  ☐ ไม่ผ่าน  ☐ N/A"],
            ["ผู้ทดสอบ (ชื่อ)", ""],
            ["วันที่", ""],
            ["หมายเหตุ", ""],
        ],
        header_fill="F2F2F2",
        header_font_white=False,
    )
    doc.add_paragraph("Preconditions:")
    doc.add_paragraph(precond)
    doc.add_paragraph("Test Steps:")
    for line in steps.split("\n"):
        doc.add_paragraph(line, style="List Number")
    doc.add_paragraph("Expected Results:")
    doc.add_paragraph(expected)


def main() -> None:
    today = date.today().isoformat()
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\n"
        "Test Case Scenario ฉบับละเอียด\n"
        "(Manual · UAT · Traceability)"
    )
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"รหัสเอกสาร: PM-PEPSI-TCS-001\n"
        f"วันที่: {today} · เวอร์ชัน: 1.0\n"
        f"จำนวน Test Case: {len(SCENARIOS)} · E2E Scenario Chain: {len(E2E_CHAINS)}"
    ).font.size = Pt(10)

    doc.add_page_break()

    doc.add_heading("การควบคุมเอกสาร", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["ชื่อเอกสาร", "Test Case Scenario Specification"],
            ["รหัส", "PM-PEPSI-TCS-001"],
            ["คู่กับ", "UAT-ROUND-3-TH · E2E-TEST · UNIT-TEST · DBRS"],
            ["วัตถุประสงค์", "คู่มือทดสอบทีม QA/UAT — ขั้นตอนและผลที่คาด"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("1. วิธีใช้เอกสาร", level=1)
    add_bullets = [
        "ใช้ร่วมกับ UAT-ROUND-3-TH.md — ชีต UAT เป็น checklist สรุป · เอกสารนี้เป็นรายละเอียดขั้นตอน",
        "P0 = ต้องผ่านก่อน go-live · P1 = ควรผ่าน · P2 = ตาม scope/Telegram",
        "กรอกผลทดสอบในแต่ละ TC: ผ่าน/ไม่ผ่าน/N/A + หมายเหตุ",
        "E2E Scenario Chain (§3) — รันต่อเนื่องเป็น story ธุรกิจ",
        "Automated cases อ้างอิง E2E-TEST-DETAILED-TH.docx",
    ]
    for item in add_bullets:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.1 ความสำคัญ (Priority)", level=2)
    add_table(
        doc,
        ["ระดับ", "ความหมาย"],
        [
            ["P0", "Blocker go-live — SAP · ปฏิทิน · confirm · RBAC"],
            ["P1", "สำคัญ — PM manual · รายงาน · planning"],
            ["P2", "ตาม scope — Telegram · portal handoff · admin เสริม"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("2. สารบัญ Test Case (สรุป)", level=1)
    modules: dict[str, list[tuple[str, ...]]] = {}
    for s in SCENARIOS:
        modules.setdefault(s[1], []).append(s)

    summary_rows = []
    for s in SCENARIOS:
        summary_rows.append([s[0], s[1], s[2][:50], s[3], s[4], s[8], s[9]])

    add_table(
        doc,
        ["รหัส", "โมดูล", "ชื่อ Scenario", "P", "Actor", "Type", "Trace"],
        summary_rows,
        header_fill="7B3F00",
    )

    doc.add_page_break()

    doc.add_heading("3. E2E Scenario Chains (Business Flow)", level=1)
    add_table(
        doc,
        ["รหัส", "ชื่อ Flow", "P", "Actor", "Test Cases ที่เชื่อม", "หมายเหตุ", "UAT"],
        [list(c) for c in E2E_CHAINS],
        header_fill="548235",
    )

    doc.add_page_break()

    doc.add_heading("4. รายละเอียด Test Case ตามโมดูล", level=1)
    module_order = [
        "Authentication",
        "Integration SAP",
        "Calendar",
        "Work Orders",
        "Planning",
        "Confirmation",
        "PM Manual",
        "Reports",
        "Admin",
        "Admin Security",
        "Telegram",
        "Portal",
        "UI Regression",
        "Regression",
    ]
    for mod in module_order:
        items = modules.get(mod, [])
        if not items:
            continue
        doc.add_heading(mod, level=2)
        for s in items:
            add_scenario_detail(doc, s)
            doc.add_paragraph()

    doc.add_page_break()

    doc.add_heading("5. Traceability Matrix", level=1)
    add_table(
        doc,
        ["UAT Flow", "Test Cases"],
        [
            ["A1 Login/EN/TH", "TC-AUTH-001, TC-AUTH-005"],
            ["A2 Import→Calendar", "TC-INT-001, TC-CAL-001"],
            ["A3 WO modal", "TC-WO-001"],
            ["A4–A5 Planning", "TC-PLN-001, TC-PLN-002"],
            ["A6 Bulk team", "TC-WO-002"],
            ["A7 Mass 44", "TC-CFM-002"],
            ["A8 รูปหลัง PM", "TC-CFM-001"],
            ["A9 QC+Export", "TC-CFM-003, TC-INT-005"],
            ["A10 Reports/Board", "TC-RPT-001, TC-RPT-002"],
            ["§F Telegram", "TC-TG-001–003"],
            ["§G Portal/UI", "TC-PRT-001, TC-UI-001"],
            ["§I Regression", "TC-REG-001, TC-CAL-004, TC-CFM-005"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("6. สรุปผลทดสอบ (กรอกหลังรัน)", level=1)
    add_table(
        doc,
        ["Priority", "จำนวน TC", "ผ่าน", "ไม่ผ่าน", "N/A", "% ผ่าน"],
        [
            ["P0", str(sum(1 for s in SCENARIOS if s[3] == "P0")), "", "", "", ""],
            ["P1", str(sum(1 for s in SCENARIOS if s[3] == "P1")), "", "", "", ""],
            ["P2", str(sum(1 for s in SCENARIOS if s[3] == "P2")), "", "", "", ""],
            ["รวม", str(len(SCENARIOS)), "", "", "", ""],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("7. เอกสารอ้างอิง", level=1)
    refs = [
        "docs/customer-requirements/UAT-ROUND-3-TH.md",
        "docs/customer-requirements/E2E-TEST-DETAILED-TH.docx",
        "docs/customer-requirements/DBRS-ISO29110-DETAILED-TH.docx",
        "docs/customer-requirements/MEETING-SUMMARY-REQUIREMENTS.md",
        "docs/USER-MANUAL-TH.md",
        "docs/customer-requirements/FLOW-DIAGRAM-SWIMLANE-TH.docx",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร Test Case Scenario —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({len(SCENARIOS)} test cases, {len(E2E_CHAINS)} chains)")


if __name__ == "__main__":
    main()
