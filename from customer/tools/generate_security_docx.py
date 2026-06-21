#!/usr/bin/env python3
"""Generate PM Pepsi App — Security & Cybersecurity document (Word .docx)."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "SECURITY-DETAILED-TH.docx"


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "2F5496",
    font_size: int = 8,
    header_font_white: bool = True,
) -> None:
    from docx.shared import RGBColor

    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                if header_font_white:
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    today = date.today().isoformat()
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Dashboard & Monitoring System\n"
        "Security & Cybersecurity Specification\n"
        "(เอกสารความปลอดภัยระบบ)"
    )
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"ลูกค้า: Pepsi-Cola (Thailand) Trading Co., Ltd.\n"
        f"ผู้พัฒนา: S.Y. Interactive Development Limited\n"
        f"รหัสเอกสาร: PM-PEPSI-SEC-001\n"
        f"เวอร์ชัน: 1.0 · วันที่: {today}\n"
        f"อ้างอิง: skills.md §3 · ISO/IEC 27001 controls (mapping แนวทาง)"
    ).font.size = Pt(10)

    doc.add_page_break()

    doc.add_heading("การควบคุมเอกสาร", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["ชื่อเอกสาร", "Security & Cybersecurity Specification"],
            ["รหัส", "PM-PEPSI-SEC-001"],
            ["เวอร์ชัน", "1.0"],
            ["สถานะ", "Draft — ส่งมอบ onsite / UAT"],
            ["เจ้าของ", "Technical Lead / Security"],
            ["ผู้อนุมัติ (ลูกค้า IT)", "___________________  วันที่ _______"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("1. บทนำ", level=1)
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายมาตรการความปลอดภัยของ PM-Pepsi-App ตามหลัก "
        "defense in depth — ยืนยันตัวตนและสิทธิ์ที่เซิร์ฟเวอร์, ควบคุมข้อมูลอ่อนไหว, "
        "รองรับสภาพแวดล้อม offline/air-gap ของโรงงาน Pepsi"
    )

    doc.add_heading("1.1 ขอบเขต", level=2)
    add_bullets(
        doc,
        [
            "ครอบคลุม: Web app (React) · API (Express) · PostgreSQL · Integration SAP · Admin",
            "ครอบคลุม: Telegram Bot (ถ้าเปิดใช้) · Portal module handoff",
            "ไม่ครอบคลุม: SAP ERP ภายใน · SQL Server/XAMPP อื่นบนเครื่องเดียวกัน",
            "ไม่ครอบคลุม: นโยบายเครือข่ายโรงงานทั้งหมด — IT ลูกค้าร่วมกำหนด perimeter",
        ],
    )

    doc.add_heading("1.2 หลักการสำคัญ", level=2)
    add_table(
        doc,
        ["หลัก", "การนำไปใช้ใน PM-Pepsi-App"],
        [
            ["Least privilege", "DB user ไม่ใช้ superuser · RBAC ต่อ permission code"],
            ["Fail secure", "ไม่มี session → 401 · ไม่มีสิทธิ์ → 403 · audit status denied"],
            ["Server-side enforcement", "UI ซ่อนปุ่ม + API ตรวจ permission ทุก route สำคัญ"],
            ["Air-gap / offline", "ไม่พึ่ง cloud SaaS · onsite update · ไม่ remote โดย default"],
            ["No trust in client", "ไม่เก็บ secret ใน frontend · validate ทุก input ที่ API"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("2. สถาปัตยกรรมความปลอดภัย", level=1)
    add_table(
        doc,
        ["ชั้น", "ควบคุม", "รายละเอียด"],
        [
            ["Perimeter", "Firewall LAN", "ไม่เปิด VPN/Tailscale remote — onsite support"],
            ["Transport", "HTTPS (IIS/Tunnel)", "TLS client ↔ reverse proxy"],
            ["Application", "Express middleware stack", "Helmet · CORS · rate limit · RBAC"],
            ["Session", "Signed cookie pm_session", "HMAC-SHA256 · TTL · HttpOnly"],
            ["Data", "PostgreSQL schema app", "BYTEA รูป · audit log · blocked IP"],
            ["Files", "Integration dirs", "SHA256 dedupe · archive · error quarantine"],
        ],
        header_fill="2F5496",
    )

    doc.add_page_break()

    doc.add_heading("3. การยืนยันตัวตน (Authentication)", level=1)

    doc.add_heading("3.1 รูปแบบ Session", level=2)
    add_table(
        doc,
        ["รายการ", "ค่า/พฤติกรรม"],
        [
            ["Mechanism", "Signed token ใน cookie `pm_session`"],
            ["Algorithm", "HMAC-SHA256 (session-token.ts)"],
            ["Secret", "SESSION_SECRET ใน .env (≥16 chars)"],
            ["TTL default", "8 ชม. (480 นาที — ปรับได้ app.session_ttl_min)"],
            ["Impersonation", "Admin สวมสิทธิ์ — TTL 30 นาที อัตโนมัติ"],
            ["Verification", "timingSafeEqual ป้องกัน timing attack"],
            ["Logout", "ลบ cookie · audit auth.logout"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )
    doc.add_paragraph(
        "Production: ตั้ง cookie Secure + HttpOnly ผ่าน reverse proxy HTTPS. "
        "ไม่เก็บรหัสผ่าน plain text — ใช้ hash ตาม tbworkcenter/tbmember"
    )

    doc.add_heading("3.2 Login & Rate Limit", level=2)
    add_table(
        doc,
        ["การควบคุม", "รายละเอียด"],
        [
            ["Rate limit /api/v1/auth", "100 req/IP/60s (default, ปรับ env)"],
            ["Rate limit /api/v1/admin", "100 req/IP/60s"],
            ["429 response", "JSON RATE_LIMIT + audit security.rate_limit"],
            ["Failed login", "บันทึก audit · ดูที่ /admin/security"],
            ["Block IP", "tbl_blocked_ip — Admin block manual"],
            ["Env", "RATE_LIMIT_ENABLED=0 ปิด (test only)"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("4. การควบคุมการเข้าถึง (RBAC)", level=1)
    doc.add_paragraph(
        "ทุก API ที่แก้ข้อมูลหรืออ่านข้อมูลสำคัญ ใช้ middleware requirePermission(code) "
        "อ้างอิง tbl_role · tbl_permission · tbl_role_permission"
    )
    add_table(
        doc,
        ["บทบาท (userst)", "role_code", "ตัวอย่างสิทธิ์"],
        [
            ["A — Admin", "admin", "admin.* · backup · settings · security"],
            ["H — Head/Foreman", "head", "confirmation.import (QC) · mass confirm"],
            ["U — User/Planner", "planner", "planning.write · iw37n.write · calendar.write"],
            ["W — Worker/ช่าง", "worker", "planning.read · confirmation.write"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )
    add_table(
        doc,
        ["หลักการ", "รายละเอียด"],
        [
            ["UI", "CanPermission / usePermission ซ่อนปุ่ม — ไม่แทน API check"],
            ["API deny", "403 + audit status denied (rbac.deny)"],
            ["Menu", "tbmenu + permission ต่อ route — Admin แก้ matrix ได้"],
            ["Board kiosk", "Public read-only routes แยก — ไม่ expose write"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("5. ความปลอดภัยเครือข่าย", level=1)
    add_table(
        doc,
        ["หัวข้อ", "นโยบาย PM Pepsi App"],
        [
            ["Remote access", "ไม่ใช้ VPN/Tailscale เข้าเซิร์ฟเวอร์ (skills.md)"],
            ["User access", "LAN โรงงานเท่านั้น — ตาม IT ลูกค้า"],
            ["Support", "On-site ที่โรงงาน — ไม่ remote admin จากอินเทอร์เน็ต"],
            ["API exposure", "127.0.0.1:4000 — IIS proxy /api เท่านั้น"],
            ["PostgreSQL", "localhost only — ไม่ expose 5432 สู่ LAN ถ้าไม่จำเป็น"],
            ["Telegram webhook", "HTTPS สาธารณะเฉพาะ endpoint webhook (Cloudflare Tunnel)"],
            ["Trust proxy", "app.set('trust proxy', 1) — X-Forwarded-For สำหรับ IP audit"],
        ],
        header_fill="7030A0",
    )

    doc.add_heading("5.1 TLS / HTTPS", level=2)
    add_bullets(
        doc,
        [
            "Production SHALL ใช้ HTTPS สำหรับผู้ใช้ web",
            "Terminate TLS ที่ IIS หรือ reverse proxy ที่ IT อนุมัติ",
            "SESSION cookie ต้อง Secure ใน production",
            "Telegram setWebhook ต้อง URL https://",
        ],
    )

    doc.add_page_break()

    doc.add_heading("6. การเสริมความแข็งแกร่ง API (API Hardening)", level=1)
    add_table(
        doc,
        ["มาตรการ", "Implementation", "ไฟล์/หมายเหตุ"],
        [
            ["Security headers", "helmet()", "app.ts"],
            ["CORS แคบ", "origin = CORS_ORIGIN + credentials", ".env"],
            ["JSON body limit", "1mb default", "express.json"],
            ["Upload size guard", "dynamic จาก tbl_setting upload_max_mb", "enforce-upload-size.ts"],
            ["Blocked IP", "middleware ก่อน routes", "blocked-ip.ts · migration 072"],
            ["Maintenance mode", "503 mutate ยกเว้น admin backup/settings", "maintenance-mode.ts"],
            ["API metrics", "request timing (ops)", "api-metrics.ts"],
        ],
        header_fill="2F5496",
    )

    doc.add_heading("6.1 Environment Variables (Secrets)", level=2)
    add_table(
        doc,
        ["ตัวแปร", "ความลับ", "ข้อกำหนด"],
        [
            ["SESSION_SECRET", "สูง", "สุ่มยาว · ไม่ commit · rotate ตามนโยบาย IT"],
            ["DATABASE_URL", "สูง", "รหัส DB แยก · least privilege user"],
            ["TELEGRAM_BOT_TOKEN", "สูง", "เก็บ .env เท่านั้น"],
            ["TELEGRAM_WEBHOOK_SECRET", "สูง", "ตรวจ header webhook"],
            ["MODULE_HANDOFF_CLIENTS", "สูง", "secret ต่อ module (portal)"],
            ["CORS_ORIGIN", "กลาง", "URL เดียว — ป้องกัน CSRF cross-origin"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )
    doc.add_paragraph("tbl_setting.is_secret = true — mask ค่าใน UI/audit export")

    doc.add_page_break()

    doc.add_heading("7. Input Validation & ไฟล์อัปโหลด", level=1)
    add_table(
        doc,
        ["ประเภท", "การควบคุม"],
        [
            ["API body/query", "Zod schema ทุก route — reject 400"],
            ["IW37N/Confirm import", "Parser whitelist · preview ก่อน commit"],
            ["Duplicate file", "SHA256 batch — skip ไม่ upsert ซ้ำ"],
            ["Confirm image", "image/* → WebP max width 1600px → BYTEA DB"],
            ["Confirm phase", "รับเฉพาะ phase=after (หลัง PM) — reject before"],
            ["PM import Excel", "Schema validate · จำกัดขนาด upload"],
            ["Integration inbound", "Extensions .xlsx/.xls/.csv only"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )
    doc.add_paragraph(
        "แนะนำ IT ลูกค้า: สแกนมัลแวร์ไฟล์ SAP ก่อนวาง inbound (ตามนโยบายโรงงาน)"
    )

    doc.add_heading("7.1 การเก็บรูป", level=2)
    add_bullets(
        doc,
        [
            "รูป confirm → tbconfirmimg.img_data (BYTEA WebP)",
            "รูปช่าง → tbworkcenter.imgmember_data",
            "โลโก้ → tbl_setting app.logo_bytes",
            "ไม่เก็บ path บน disk ใน production flow",
            "Backup DB ครอบคลุมรูป — ไม่แยกไฟล์ imgMember/",
        ],
    )

    doc.add_page_break()

    doc.add_heading("8. Audit Trail & Logging", level=1)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["ตาราง", "app.tbl_audit_log"],
            ["ฟิลด์", "actor_id, action, resource, before/after json, ip, status"],
            ["Actions ตัวอย่าง", "auth.login · auth.logout · rbac.deny · planning.assign · admin.user.create"],
            ["Denied", "status=denied — failed login, RBAC, rate limit"],
            ["UI", "/admin/audit · /admin/security · /activity-log"],
            ["Retention", "ตั้งใน Admin (แนะนำ 365 วัน — ตาม IT)"],
            ["Export", "Admin export CSV — ไม่รวม secret fields"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("8.1 Admin Security Console", level=2)
    add_table(
        doc,
        ["ฟีเจอร์", "Route", "Permission"],
        [
            ["Security dashboard", "/admin/security", "admin.security.read"],
            ["Failed logins", "GET .../security/failed-login", "admin.security.read"],
            ["RBAC denied", "GET .../security/denied", "admin.security.read"],
            ["Block IP", "POST .../security/blocked-ips", "admin.security.write"],
            ["Unblock IP", "DELETE .../blocked-ips/:id", "admin.security.write"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("9. Frontend Security", level=1)
    add_table(
        doc,
        ["หัวข้อ", "มาตรการ"],
        [
            ["XSS", "React default escape · หลีกเลี่ยง dangerouslySetInnerHTML"],
            ["CSP", "Helmet headers จาก API proxy layer / IIS"],
            ["Credentials", "fetch credentials:include กับ same-origin / CORS ที่กำหนด"],
            ["Secrets", "ไม่มี API key ใน Vite bundle — VITE_* เฉพาะ public config"],
            ["IndexedDB", "feature flag off default — PDPA/cache policy"],
            ["i18n", "ไม่ render raw HTML จาก user input"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("10. Database Security", level=1)
    add_bullets(
        doc,
        [
            "ใช้ role pepsipm — ไม่ใช้ postgres superuser ใน DATABASE_URL",
            "Schema app แยกจาก public",
            "PostgreSQL listen localhost / firewall",
            "รหัสผ่าน strong · เปลี่ยนตามนโยบาย IT",
            "Backup encrypted at rest — ตาม disk encryption โรงงาน",
            "Restore ต้อง maintenance mode + audit",
        ],
    )

    doc.add_page_break()

    doc.add_heading("11. Telegram & Portal (ถ้าเปิดใช้)", level=1)

    doc.add_heading("11.1 Telegram Bot", level=2)
    add_table(
        doc,
        ["การควบคุม", "รายละเอียด"],
        [
            ["Webhook auth", "Header X-Telegram-Bot-Api-Secret-Token = TELEGRAM_WEBHOOK_SECRET"],
            ["Link account", "One-time token · migration 100"],
            ["Chat session", "State machine ปิดงาน — phase after only"],
            ["Notify", "TELEGRAM_NOTIFY_ENABLED=false ปิดส่ง"],
            ["Admin", "กลุ่ม Chat ID จัดการใน /admin/telegram"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("11.2 Portal Module Handoff", level=2)
    add_table(
        doc,
        ["การควบคุม", "รายละเอียด"],
        [
            ["Feature flag", "PORTAL_ENABLED · MODULE_HANDOFF_ENABLED"],
            ["TTL token", "MODULE_HANDOFF_TTL_SEC (default 60s)"],
            ["Rate limit", "MODULE_HANDOFF_RATE_PER_MIN"],
            ["Client secret", "MODULE_HANDOFF_CLIENTS=module:secret"],
            ["Allowed origins", "MODULE_HANDOFF_ALLOWED_ORIGINS"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("12. Supply Chain & Delivery Integrity", level=1)
    add_table(
        doc,
        ["หัวข้อ", "แนวทาง"],
        [
            ["Dependencies", "package-lock.json · npm audit ก่อน release"],
            ["Offline delivery", "SHA256 manifest ก่อน docker load / copy build (skills.md §4.1)"],
            ["Git", "ไม่ commit .env · .gitignore secrets"],
            ["CI", "lint + typecheck + unit test ก่อน merge"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("13. การตอบสนองเหตุการณ์ (Incident Response)", level=1)
    add_table(
        doc,
        ["เหตุการณ์", "ขั้นตอนเบื้องต้น"],
        [
            ["Brute-force login", "ดู /admin/security → Block IP → แจ้ง IT"],
            ["Account compromise", "Reset password user · revoke session · audit log"],
            ["Data leak suspicion", "Maintenance ON · ตรวจ audit · restore ถ้าจำเป็น"],
            ["Ransomware/malware", "Isolated server · restore จาก backup offline · onsite"],
            ["Webhook abuse", "Rotate TELEGRAM_WEBHOOK_SECRET · setWebhook ใหม่"],
        ],
        header_fill="FCE4D6",
        header_font_white=False,
    )

    doc.add_heading("14. Security Checklist ก่อน Go-live", level=1)
    add_table(
        doc,
        ["#", "รายการ", "☐"],
        [
            ["1", "HTTPS เปิดใช้ · cookie Secure", "☐"],
            ["2", "SESSION_SECRET เปลี่ยนจาก default", "☐"],
            ["3", "DATABASE_URL ใช้ user least privilege", "☐"],
            ["4", "CORS_ORIGIN ตรง URL production เท่านั้น", "☐"],
            ["5", ".env ไม่อยู่ใน git / สิทธิ์ไฟล์จำกัด", "☐"],
            ["6", "PostgreSQL ไม่ expose สาธารณะ", "☐"],
            ["7", "Backup schedule + ทดสอบ restore staging", "☐"],
            ["8", "Admin security page ดู failed login ได้", "☐"],
            ["9", "RBAC ทดสอบช่างไม่เห็นปุ่ม admin", "☐"],
            ["10", "Telegram webhook secret ตั้งแล้ว (ถ้าใช้)", "☐"],
            ["11", "npm audit ไม่มี critical ค้าง (หรือบันทึกยอมรับ)", "☐"],
            ["12", "IT ลูกค้าลงนาม checklist นี้", "☐"],
        ],
        header_fill="2F5496",
    )

    doc.add_heading("15. การแมป ISO / มาตรฐาน (แนวทาง)", level=1)
    add_table(
        doc,
        ["Control area", "PM-Pepsi-App"],
        [
            ["A.5 Organizational", "On-site support · escalation ใน Runbook"],
            ["A.8 Asset management", "D: drive layout · backup retention"],
            ["A.9 Access control", "RBAC · session · block IP"],
            ["A.10 Cryptography", "TLS · HMAC session · hashed passwords"],
            ["A.12 Operations", "Audit · maintenance · backup SHA256"],
            ["A.14 System acquisition", "Zod validation · secure SDLC"],
            ["A.17 Communications", "HTTPS · CORS · no VPN remote"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("16. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "skills.md §3 Security & Cybersecurity",
            "docs/customer-requirements/INSTALL-DEPLOY-RUNBOOK-TH.docx",
            "docs/customer-requirements/DBRS-ISO29110-DETAILED-TH.docx",
            "docs/parity-pending/14-administrator.md §3.4 audit",
            "PM-Pepsi-App/backend/src/app.ts",
            "PM-Pepsi-App/backend/.env.example",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร Security —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
