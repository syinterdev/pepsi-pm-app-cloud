#!/usr/bin/env python3
"""Generate PM Pepsi App — Backend Design documentation (Word .docx)."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
BACKEND = ROOT / "PM-Pepsi-App" / "backend"
ROUTES = BACKEND / "src" / "routes"
OUT = ROOT / "docs" / "customer-requirements" / "BACKEND-DESIGN-DETAILED-TH.docx"

ROUTE_MODULE_DESC: dict[str, str] = {
    "health": "Health check — ไม่ต้อง auth",
    "auth": "Login / logout / session / user log",
    "portal": "โมดูล Portal หลัง login",
    "settings": "Public settings (branding, locale)",
    "dashboard": "KPI หน้าแรก",
    "calendar": "ปฏิทิน WO — events, move plan",
    "scheduling": "Plan calendar + line schedule",
    "backlog": "WO ค้างแผน",
    "planning": "แผน PM/CM — assign งาน",
    "work-orders": "ค้นหา WO · batch team · PM execution",
    "iw37n": "นำเข้า/แก้ไข IW37N",
    "integration": "Integration jobs · watch inbound",
    "confirmation": "(ใน work-orders/personnel routes)",
    "manhours": "ชั่วโมงงาน · worktime · import",
    "personnel": "บุคลากร · confirm % · รูป",
    "reports": "KPI · activity log · weekly summary",
    "master-data": "CRUD master 17 entities",
    "nav": "เมนู sidebar ตาม RBAC",
    "profile": "โปรไฟล์ผู้ใช้",
    "users": "รายชื่อ users สำหรับ picker",
    "user-pref": "ค่าตั้งผู้ใช้ (sidebar, tour)",
    "pm-readings": "PM vibration readings",
    "board-kiosk": "Engineering board kiosk",
    "board-activity": "Board activity feed",
    "board-pm-readings": "Board PM readings",
    "telegram": "Telegram webhook (Bot)",
    "announcements": "ประกาศ active banner",
    "admin-branding": "โลโก้ · สี · favicon",
    "admin-settings": "ตั้งค่าระบบ tbl_setting",
    "admin-audit": "Audit log",
    "admin-health": "สุขภาพระบบ · migration",
    "admin-users": "จัดการผู้ใช้ workcenter",
    "admin-roles": "บทบาท & permission matrix",
    "admin-menu": "Menu builder tbmenu",
    "admin-backup": "Backup / restore pg_dump",
    "admin-announcement": "ประกาศ admin",
    "admin-security": "ความปลอดภัย · blocked IP",
    "admin-about": "About · license · version",
    "admin-telegram": "ตั้งค่า Telegram groups",
}


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "BDD7EE",
    font_size: int = 8,
) -> None:
    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def scan_endpoints() -> list[tuple[str, str, str]]:
    pat = re.compile(r"""app\.(get|post|put|patch|delete)\(\s*['"]([^'"]+)['"]""", re.IGNORECASE)
    eps: list[tuple[str, str, str]] = []
    for f in sorted(ROUTES.glob("*.ts")):
        if ".test." in f.name:
            continue
        text = f.read_text(encoding="utf-8", errors="replace")
        stem = f.stem
        for m in pat.finditer(text):
            path = m.group(2)
            if path.startswith("/api"):
                eps.append((stem, m.group(1).upper(), path))
    return eps


def count_files(folder: Path, pattern: str) -> int:
    return len([f for f in folder.glob(pattern) if ".test." not in f.name])


def main() -> None:
    today = date.today().isoformat()
    endpoints = scan_endpoints()
    src = BACKEND / "src"

    route_count = count_files(ROUTES, "*.ts")
    service_count = count_files(src / "services", "*.ts")
    schema_count = count_files(src / "schemas", "*.ts")
    middleware_count = count_files(src / "middleware", "*.ts")
    lib_count = count_files(src / "lib", "*.ts")

    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\nเอกสาร Backend Design ฉบับละเอียด\n(Express API · PostgreSQL · RBAC)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0\n"
        f"แพ็กเกจ: PM-Pepsi-App/backend (pm-api)"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายสถาปัตยกรรม backend ของ PM-Pepsi-App "
        "รวมเทคโนโลยี ชั้นโค้ด middleware pipeline การยืนยันตัวตน RBAC "
        "รายการ API endpoints โมดูล services และแนวทาง deploy — "
        "สแกนจากซอร์สโค้ดจริง"
    )

    doc.add_page_break()

    # 1 Overview
    doc.add_heading("1. ภาพรวมสถาปัตยกรรม", level=1)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Runtime", "Node.js 20+ · ESM (type: module)"],
            ["Framework", "Express 4.x"],
            ["ภาษา", "TypeScript 5.x"],
            ["Database", "PostgreSQL ผ่าน pg (connection pool)"],
            ["Validation", "Zod schemas"],
            ["Auth", "HMAC JWT — Cookie pm_session + Bearer header"],
            ["API prefix", "/api/v1/*"],
            ["Default port", "4000"],
            ["Route modules", str(route_count)],
            ["API endpoints (สแกน)", str(len(endpoints))],
            ["Services", str(service_count)],
            ["Schemas (Zod)", str(schema_count)],
            ["Middleware", str(middleware_count)],
            ["Lib utilities", str(lib_count)],
        ],
    )

    doc.add_heading("1.1 Layered Architecture", level=2)
    doc.add_paragraph(
        "Backend ใช้โครงสร้างชั้นแบบ Route → Service → Lib/DB แยกความรับผิดชอบชัดเจน"
    )
    add_table(
        doc,
        ["ชั้น", "โฟลเดอร์", "หน้าที่"],
        [
            ["Entry", "src/index.ts", "โหลด env, สร้าง pool, start server, schedulers"],
            ["App factory", "src/app.ts", "ประกอบ middleware + register routes"],
            ["Routes", "src/routes/*.ts", "HTTP handlers — thin, เรียก service"],
            ["Services", "src/services/*.ts", "Business logic, SQL queries, orchestration"],
            ["Schemas", "src/schemas/*.ts", "Zod — validate request/response"],
            ["Middleware", "src/middleware/*.ts", "Auth, RBAC, rate limit, maintenance"],
            ["Lib", "src/lib/*.ts", "Pure helpers, audit, permissions, parsers"],
            ["Config", "src/config/env.ts", "Environment validation (Zod)"],
            ["DB", "src/db/pool.ts", "PostgreSQL connection pool"],
        ],
    )

    doc.add_page_break()

    # 2 Request pipeline
    doc.add_heading("2. Request Pipeline (Middleware Chain)", level=1)
    doc.add_paragraph("ลำดับ middleware ใน createApp() — ทุก request ผ่านตามนี้:")

    add_table(
        doc,
        ["ลำดับ", "Middleware", "หน้าที่"],
        [
            ["1", "helmet()", "Security headers"],
            ["2", "express.json({ limit: '1mb' })", "Parse JSON body"],
            ["3", "cors (ถ้ามี CORS_ORIGIN)", "Cross-origin + credentials"],
            ["4", "blocked-ip guard", "ตรวจ tbl_blocked_ip"],
            ["5", "rate-limiters", "จำกัด /auth/* และ /admin/*"],
            ["6", "api-metrics", "บันทึก latency สำหรับ slow API report"],
            ["7", "upload-size guard", "จำกัดขนาด upload จาก tbl_setting"],
            ["8", "maintenance mode", "503 สำหรับ mutating methods"],
            ["9", "Route handlers", "Auth + RBAC ต่อ route"],
        ],
    )

    doc.add_heading("2.1 แผนภาพ Request Flow", level=2)
    flow_lines = [
        "Client (React / Board / Telegram)",
        "    ↓ HTTP",
        "Express Middleware Chain",
        "    ↓",
        "requireApiAuth (JWT cookie / Bearer)",
        "    ↓",
        "requirePermission('xxx.yyy')",
        "    ↓",
        "validateBody(Zod schema)",
        "    ↓",
        "Service layer → pg Pool → PostgreSQL app.*",
        "    ↓",
        "JSON response { data } หรือ { error, message }",
    ]
    p = doc.add_paragraph()
    r = p.add_run("\n".join(flow_lines))
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.add_page_break()

    # 3 Auth & RBAC
    doc.add_heading("3. Authentication & Authorization", level=1)

    doc.add_heading("3.1 Session / JWT", level=2)
    add_bullets(
        doc,
        [
            "POST /api/v1/auth/login — ตรวจ tbworkcenter หรือ tbl_member",
            "สร้าง HMAC-signed JWT (SESSION_SECRET) — เก็บใน cookie pm_session (httpOnly)",
            "Client ส่ง Bearer token หรือ cookie — requireApiAuth ถอดรหัส",
            "GET /api/v1/auth/me — คืน user + permissions[]",
            "Session TTL จาก tbl_setting app.session_ttl_min (default 8 ชม.)",
            "Login lockout: security.max_login_attempts + tbl_blocked_ip",
        ],
    )

    doc.add_heading("3.2 RBAC", level=2)
    add_bullets(
        doc,
        [
            "createRequirePermission(pool, secret)('perm.code') — คืน [requireAuth, checkPerm]",
            "createRequireAnyPermission — ผ่านถ้ามีสิทธิ์ใดสิทธิ์หนึ่ง",
            "hasPermission() — อ่าน tbl_role_permission + cache 60s",
            "Legacy fallback: userst A/U/W ถ้ายังไม่ migrate RBAC",
            "403 FORBIDDEN + audit rbac.deny",
        ],
    )

    doc.add_heading("3.3 Permission กลุ่มหลัก", level=2)
    add_table(
        doc,
        ["perm_group", "ตัวอย่าง perm_code"],
        [
            ["dashboard", "dashboard.read"],
            ["calendar", "calendar.read, calendar.write"],
            ["planning", "planning.read, planning.assign"],
            ["work-orders", "work-orders.read, work-orders.write"],
            ["confirmation", "confirmation.read, .import, .export"],
            ["iw37n", "iw37n.read, iw37n.import, iw37n.write"],
            ["manhours", "manhours.read, manhours.admin, manhours.import"],
            ["admin.*", "admin.users.read, admin.backup.write, …"],
        ],
    )

    doc.add_page_break()

    # 4 Error handling
    doc.add_heading("4. รูปแบบ Response และ Error", level=1)
    add_table(
        doc,
        ["HTTP", "error code", "ความหมาย"],
        [
            ["200", "—", "สำเร็จ"],
            ["400", "VALIDATION_ERROR", "Zod validate ไม่ผ่าน"],
            ["401", "UNAUTHORIZED", "ไม่ได้ login / session หมดอายุ"],
            ["403", "FORBIDDEN", "ไม่มีสิทธิ์ RBAC"],
            ["404", "NOT_FOUND", "ไม่พบทรัพยากร"],
            ["409", "CONFLICT", "สถานะไม่อนุญาต (WO ปิดแล้ว ฯลฯ)"],
            ["413", "PAYLOAD_TOO_LARGE", "ไฟล์/upload ใหญ่เกิน"],
            ["503", "MAINTENANCE", "โหมดบำรุงรักษา"],
            ["503", "SCHEMA_NOT_READY", "Migration ยังไม่ครบ"],
        ],
    )

    doc.add_paragraph("รูปแบบ error JSON:")
    p = doc.add_paragraph()
    r = p.add_run('{\n  "error": "FORBIDDEN",\n  "message": "ไม่มีสิทธิ์ (planning.assign)"\n}')
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.add_page_break()

    # 5 Environment
    doc.add_heading("5. Environment Variables", level=1)
    add_table(
        doc,
        ["ตัวแปร", "บังคับ", "คำอธิบาย"],
        [
            ["DATABASE_URL", "ใช่", "postgresql://user:pass@host:port/db"],
            ["SESSION_SECRET", "ใช่", "อย่างน้อย 16 ตัวอักษร — sign JWT"],
            ["PORT", "ไม่", "default 4000"],
            ["NODE_ENV", "ไม่", "development | production | test"],
            ["CORS_ORIGIN", "ไม่", "เช่น http://127.0.0.1:5173"],
            ["BACKUP_SCHEDULER", "ไม่", "0 = ปิด auto backup"],
            ["INTEGRATION_WATCH_SCHEDULER", "ไม่", "0 = ปิด inbound watch"],
            ["RATE_LIMIT_MAX", "ไม่", "จำกัด request ต่อ IP/นาที"],
            ["STRESS_*", "ไม่", "ใช้กับ scripts/stress-api.ts เท่านั้น"],
        ],
    )

    doc.add_page_break()

    # 6 Background jobs
    doc.add_heading("6. Background Jobs & Schedulers", level=1)
    add_table(
        doc,
        ["Job", "ไฟล์", "เปิดเมื่อ", "หน้าที่"],
        [
            ["Backup scheduler", "services/admin-backup.ts", "BACKUP_SCHEDULER≠0", "pg_dump ตาม cron ใน tbl_setting"],
            ["Integration watch", "services/integration-scheduler.ts", "INTEGRATION_WATCH_SCHEDULER≠0", "สแกน inbound/ ทุก 1 นาที"],
            ["Upload limit refresh", "index.ts setInterval 60s", "always", "อ่าน app.upload_max_mb จาก DB"],
        ],
    )

    doc.add_page_break()

    # 7 API catalog
    doc.add_heading("7. รายการ API Endpoints", level=1)
    doc.add_paragraph(f"สแกนจาก src/routes/ — รวม {len(endpoints)} endpoints")

    by_module: dict[str, list[tuple[str, str]]] = defaultdict(list)
    for stem, method, path in endpoints:
        by_module[stem].append((method, path))

    sec = 0
    for stem in sorted(by_module.keys()):
        sec += 1
        desc = ROUTE_MODULE_DESC.get(stem, stem)
        doc.add_heading(f"7.{sec} {stem} — {desc}", level=2)
        rows = [[m, p] for m, p in sorted(by_module[stem], key=lambda x: (x[1], x[0]))]
        add_table(doc, ["Method", "Path"], rows)
        if sec % 4 == 0 and sec < len(by_module):
            doc.add_page_break()

    doc.add_page_break()

    # 8 Services catalog
    doc.add_heading("8. Services Catalog", level=1)
    doc.add_paragraph("Business logic หลักอยู่ใน src/services/ — แยกตามโดเมน")

    service_groups = [
        ("Auth & Admin", ["auth.ts", "admin-users.ts", "admin-roles.ts", "admin-settings.ts", "admin-audit.ts", "admin-health.ts", "admin-backup.ts", "admin-branding.ts", "admin-menu.ts", "admin-security.ts", "admin-about.ts", "admin-telegram.ts", "setting-store.ts", "maintenance-mode.ts"]),
        ("Work Order & Calendar", ["work-orders.ts", "work-order-workflow.ts", "calendar.ts", "planning.ts", "plan-calendar.ts", "scheduling-shared.ts", "backlog (via routes)"]),
        ("Confirmation & Export", ["confirmation.ts", "confirm-image.ts", "confirmation-export-csv.ts", "confirmation-import.ts", "personnel-close.ts", "personnel-confirm.ts"]),
        ("Import & Integration", ["iw37n.ts", "iw37n-parser", "integration-watch.ts", "integration-scheduler.ts", "integration-job.ts"]),
        ("Manhour & Reports", ["manhours-parse.ts", "reports.ts", "eng-utilization-db.ts", "eng-utilization-summary.ts", "activity-log.ts"]),
        ("PM & Board", ["pm-readings-query.ts", "board-activity.ts", "board-kiosk.ts"]),
        ("Telegram", ["telegram-webhook.ts", "telegram-assignment-notify.ts", "telegram-notify-groups.ts", "telegram-link.ts"]),
        ("Portal", ["portal-modules.ts", "module-handoff.ts"]),
    ]

    for group_name, files in service_groups:
        doc.add_heading(group_name, level=3)
        existing = []
        for name in files:
            matches = list((src / "services").glob(f"{name}*"))
            existing.extend(m.stem for m in matches if ".test." not in m.name)
        add_bullets(doc, sorted(set(existing)) or ["—"])

    doc.add_page_break()

    # 9 Key design patterns
    doc.add_heading("9. รูปแบบการออกแบบ (Design Patterns)", level=1)
    add_bullets(
        doc,
        [
            "Thin routes — route เรียก service + validateBody เท่านั้น",
            "Zod schemas แยกไฟล์ — ใช้ซ้ำใน test และ route",
            "registerXxxRoutes(app, pool, sessionSecret) — ลงทะเบียน modular",
            "Pool per app instance — ไม่สร้าง connection ต่อ request",
            "Audit voidAudit / auditLog — บันทึก mutation แบบ fire-and-forget",
            "Advisory lock — backup/restore ไม่ซ้อนกัน",
            "BYTEA images — sharp แปลง WebP ก่อนเก็บ tbconfirmimg",
            "Idempotent import — SHA256 batch dedup ใน integration",
        ],
    )

    doc.add_heading("9.1 การทดสอบ Backend", level=2)
    add_table(
        doc,
        ["ชั้น", "เครื่องมือ", "คำสั่ง"],
        [
            ["Unit", "Vitest", "cd backend && npm test"],
            ["Integration API", "Supertest", "npm run test:integration"],
            ["Stress", "stress-api.ts", "npm run stress:api"],
            ["UAT scripts", "phase2-uat.ts", "npm run uat:phase2"],
        ],
    )

    doc.add_heading("9.2 Deploy", level=2)
    add_bullets(
        doc,
        [
            "Build: npm run build → dist/",
            "Start: npm start หรือ PM2 process",
            "Reverse proxy: IIS / nginx → :4000",
            "รัน migration ก่อน deploy: database/scripts/run-all-migrations.ps1",
            "ตรวจ health: GET /api/v1/health และ /api/v1/admin/health",
        ],
    )

    doc.add_page_break()

    # 10 References
    doc.add_heading("10. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "PM-Pepsi-App/backend/src/app.ts — route registration",
            "PM-Pepsi-App/backend/package.json — dependencies",
            "docs/customer-requirements/DATABASE-DESIGN-DETAILED-TH.docx",
            "docs/customer-requirements/FLOW-DIAGRAM-SWIMLANE-TH.docx",
            "docs/customer-requirements/UNIT-TEST-DETAILED-TH.docx",
            "from customer/tools/generate_backend_design_docx.py",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Endpoints: {len(endpoints)}, Routes: {route_count}, Services: {service_count}")


if __name__ == "__main__":
    main()
