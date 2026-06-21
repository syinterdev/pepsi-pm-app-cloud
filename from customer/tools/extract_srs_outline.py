# -*- coding: utf-8 -*-
"""Dump SRS docx outline for mapping to PROJECT_PLAN."""
from pathlib import Path

from docx import Document

DOC = Path(__file__).resolve().parents[1] / "docs" / (
    "Software Requirement Specification Pepsi Cola PM Project.docx"
)
OUT = Path(__file__).resolve().parents[1] / "docs" / "_srs_outline_extract.txt"


def main() -> None:
    doc = Document(str(DOC))
    lines: list[str] = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if not t:
            continue
        st = p.style.name if p.style else ""
        lines.append(f"{i:5d} | {st:30s} | {t[:200]}")
    # Tables: first row of each table (often headers)
    lines.append("\n--- TABLES (first 2 rows each) ---\n")
    for ti, table in enumerate(doc.tables):
        lines.append(f"\nTABLE {ti} rows={len(table.rows)} cols={len(table.columns)}")
        for ri in range(min(2, len(table.rows))):
            row = table.rows[ri]
            cells = [c.text.strip().replace("\n", " ")[:80] for c in row.cells]
            lines.append("  | " + " :: ".join(cells))
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print("Wrote", OUT, "lines", len(lines))


if __name__ == "__main__":
    main()
