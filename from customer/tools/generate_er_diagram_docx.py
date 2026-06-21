#!/usr/bin/env python3
"""Generate PM Pepsi App — detailed ER-Diagram documentation (Word .docx)."""
from __future__ import annotations

import re
from collections import defaultdict
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
MIGRATIONS = ROOT / "database" / "migrations"
OUT = ROOT / "docs" / "customer-requirements" / "ER-DIAGRAM-DETAILED-TH.docx"
ER_IMAGE = ROOT / "docs" / "customer-requirements" / "diagrams" / "er-diagram-full.png"

MODULE_MAP: dict[str, tuple[str, str]] = {
    "tbworkcenter": ("1. Auth & Personnel", "ผู้ใช้ / Work Center / ช่าง"),
    "tbworkcenter_userlog": ("1. Auth & Personnel", "ประวัติ login/logout"),
    "tbl_member": ("1. Auth & Personnel", "สมาชิกระบบ (legacy)"),
    "tbl_user_pref": ("1. Auth & Personnel", "ค่าตั้งผู้ใช้ (sidebar, locale)"),
    "tbl_system_userlog": ("1. Auth & Personnel", "System user log"),
    "tbl_role": ("2. RBAC & Menu", "บทบาท (EE, ME, Admin…)"),
    "tbl_permission": ("2. RBAC & Menu", "สิทธิ์ API/UI"),
    "tbl_role_permission": ("2. RBAC & Menu", "เชื่อม role ↔ permission"),
    "tbmenu": ("2. RBAC & Menu", "เมนูแอป + RBAC"),
    "tbl_app_module": ("2. RBAC & Menu", "โมดูล Portal (PM, …)"),
    "tbl_module_handoff_code": ("2. RBAC & Menu", "รหัส handoff ข้ามโมดูล"),
    "tbzone": ("3. Master Data", "Zone / พื้นที่"),
    "tbproductline": ("3. Master Data", "สายการผลิต"),
    "tbwkctrtype": ("3. Master Data", "ประเภท Work Center"),
    "tbwkctrgroup": ("3. Master Data", "กลุ่ม Work Center"),
    "tbdepartment": ("3. Master Data", "แผนก"),
    "tbposition": ("3. Master Data", "ตำแหน่ง"),
    "tbwklevel": ("3. Master Data", "ระดับช่าง"),
    "tbequipment": ("3. Master Data", "อุปกรณ์ (Equipment)"),
    "tbmainteanance": ("3. Master Data", "เครื่องจักร / Machine"),
    "tbfunctional": ("3. Master Data", "Functional Location"),
    "tbwkzb": ("3. Master Data", "Work Center ZB mapping"),
    "tbmaterial": ("3. Master Data", "วัสดุ / spare part"),
    "tbactivitytype": ("3. Master Data", "ประเภทกิจกรรม PM"),
    "tbreason": ("3. Master Data", "เหตุผลย้ายแผน"),
    "tbwkctrstatus": ("3. Master Data", "สถานะ Work Center"),
    "tbiw37n": ("4. Work Order (IW37N)", "Work Order หลัก — ศูนย์กลาง ERD"),
    "tbwkstatus": ("4. Work Order (IW37N)", "สีสถานะ WO (CRTD, REL…)"),
    "tbmoveplan": ("4. Work Order (IW37N)", "ย้ายวันแผน / calendar day"),
    "tbplangingwork": ("5. Planning & Scheduling", "มอบหมายงาน / planning assign"),
    "tblineschdul": ("5. Planning & Scheduling", "ตาราง Line schedule"),
    "tbcofirm": ("6. Confirmation & Export", "Confirmation header"),
    "tbconfirmcom": ("6. Confirmation & Export", "ความคิดเห็น confirmation"),
    "tbconfirmimg": ("6. Confirmation & Export", "รูป confirmation (BYTEA)"),
    "tbwrkclose": ("6. Confirmation & Export", "ปิดงาน personnel"),
    "tbmanhours": ("7. Manhour & HR", "ชั่วโมงทำงาน"),
    "tbwo_pm_note": ("8. PM Execution", "บันทึก PM note"),
    "tbwo_pm_reading": ("8. PM Execution", "ค่าวัด PM (vibration…)"),
    "tbwo_pm_note_entry": ("8. PM Execution", "รายการ note แยกบรรทัด"),
    "tbiw37n_import_batch": ("9. Import & Integration", "Batch import IW37N"),
    "tbiw37n_import_row": ("9. Import & Integration", "แถว import"),
    "integration_job": ("9. Import & Integration", "คิว integration job"),
    "tbtasklist": ("10. Task List", "General task list SAP"),
    "tbl_setting": ("11. Admin & System", "ตั้งค่าระบบ key-value"),
    "tbl_audit_log": ("11. Admin & System", "Audit log"),
    "tbl_backup_history": ("11. Admin & System", "ประวัติ backup"),
    "tbl_announcement": ("11. Admin & System", "ประกาศแอป"),
    "tbl_blocked_ip": ("11. Admin & System", "IP ที่บล็อก"),
    "tbl_resource_revision": ("11. Admin & System", "revision ทรัพยากร"),
    "tbl_telegram_link_token": ("12. Telegram", "Token เชื่อม Telegram"),
    "tbl_telegram_chat_session": ("12. Telegram", "Session แชท assign WO"),
    "tbl_telegram_notify_group": ("12. Telegram", "กลุ่มแจ้งเตือน"),
    "tbl_telegram_notify_group_member": ("12. Telegram", "สมาชิกกลุ่มแจ้งเตือน"),
}

VIEW_MODULES: dict[str, tuple[str, str]] = {
    "view_order": ("4. Work Order (IW37N)", "ปฏิทิน WO + move plan + สีสถานะ"),
    "view_planwork": ("5. Planning & Scheduling", "แผนงานที่ assign แล้ว"),
    "view_lineschdul": ("5. Planning & Scheduling", "Line schedule view"),
    "view_confirmation": ("6. Confirmation & Export", "Confirmation รวม"),
    "view_confrim": ("6. Confirmation & Export", "Confirmation (legacy spell)"),
    "view_exportconfirm": ("6. Confirmation & Export", "Export SAP confirmation"),
    "view_personelclose": ("6. Confirmation & Export", "ปิดงาน personnel"),
    "view_countpersonelclose": ("7. Manhour & HR", "นับปิดงาน personnel / QC"),
}


@dataclass
class Column:
    name: str
    dtype: str
    nullable: bool = True
    is_pk: bool = False
    fk_ref: str | None = None
    default: str | None = None


@dataclass
class Table:
    schema: str
    name: str
    columns: dict[str, Column] = field(default_factory=dict)
    comment: str | None = None
    is_view: bool = False
    view_sql: str | None = None


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "E2EFDA",
    font_size: int = 8,
) -> None:
    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def strip_sql_comments(sql: str) -> str:
    sql = re.sub(r"--[^\n]*", "", sql)
    sql = re.sub(r"/\*.*?\*/", "", sql, flags=re.DOTALL)
    return sql


def parse_column_line(line: str, pk_cols: set[str]) -> Column | None:
    line = line.strip().rstrip(",")
    if not line or line.upper().startswith(("CONSTRAINT", "PRIMARY KEY", "UNIQUE", "FOREIGN KEY", "CHECK")):
        return None
    m = re.match(
        r"^(\w+)\s+(.+)$",
        line,
        re.IGNORECASE,
    )
    if not m:
        return None
    name, rest = m.group(1), m.group(2).strip()
    if name.upper() in ("PRIMARY", "CONSTRAINT", "UNIQUE", "FOREIGN"):
        return None

    fk_ref = None
    fk_m = re.search(
        r"REFERENCES\s+(?:app\.)?(\w+)\s*\((\w+)\)",
        rest,
        re.IGNORECASE,
    )
    if fk_m:
        fk_ref = f"{fk_m.group(1)}.{fk_m.group(2)}"

    is_pk = "PRIMARY KEY" in rest.upper() or name in pk_cols
    nullable = "NOT NULL" not in rest.upper() or is_pk
    default_m = re.search(r"DEFAULT\s+(.+?)(?:\s+NOT NULL|\s+PRIMARY|\s+REFERENCES|$)", rest, re.IGNORECASE)
    default = default_m.group(1).strip() if default_m else None

    dtype = rest
    for cut in ("PRIMARY KEY", "REFERENCES", "NOT NULL", "DEFAULT", "UNIQUE"):
        idx = dtype.upper().find(cut)
        if idx > 0:
            dtype = dtype[:idx].strip()
            break
    dtype = re.sub(r"\s+", " ", dtype).strip()

    return Column(name=name, dtype=dtype, nullable=nullable, is_pk=is_pk, fk_ref=fk_ref, default=default)


def parse_create_table(sql: str, tables: dict[str, Table]) -> None:
    for m in re.finditer(
        r"CREATE TABLE IF NOT EXISTS\s+app\.(\w+)\s*\((.*?)\);",
        sql,
        re.IGNORECASE | re.DOTALL,
    ):
        tname = m.group(1)
        body = m.group(2)
        pk_cols: set[str] = set()
        for pk_m in re.finditer(r"PRIMARY KEY\s*\(([^)]+)\)", body, re.IGNORECASE):
            pk_cols.update(c.strip() for c in pk_m.group(1).split(","))

        tbl = tables.get(tname) or Table(schema="app", name=tname)
        for raw_line in body.split("\n"):
            col = parse_column_line(raw_line, pk_cols)
            if col:
                tbl.columns[col.name] = col
        tables[tname] = tbl


def parse_alter_table(sql: str, tables: dict[str, Table]) -> None:
    for m in re.finditer(
        r"ALTER TABLE\s+app\.(\w+)\s+(.*?);",
        sql,
        re.IGNORECASE | re.DOTALL,
    ):
        tname, body = m.group(1), m.group(2)
        tbl = tables.get(tname) or Table(schema="app", name=tname)
        for col_m in re.finditer(
            r"ADD\s+COLUMN(?:\s+IF\s+NOT\s+EXISTS)?\s+(\w+)\s+([^,]+?)(?=\s*,\s*ADD\s+COLUMN|\s*$)",
            body,
            re.IGNORECASE | re.DOTALL,
        ):
            col_name, rest = col_m.group(1), col_m.group(2).strip().rstrip(",")
            col = parse_column_line(f"{col_name} {rest}", set())
            if col:
                tbl.columns[col.name] = col
        tables[tname] = tbl

    for m in re.finditer(
        r"ALTER TABLE\s+app\.(\w+)\s+ADD CONSTRAINT\s+\w+\s+FOREIGN KEY\s*\((\w+)\)\s*REFERENCES\s+app\.(\w+)\s*\((\w+)\)",
        sql,
        re.IGNORECASE,
    ):
        tname, col, ref_t, ref_c = m.group(1), m.group(2), m.group(3), m.group(4)
        tbl = tables.get(tname) or Table(schema="app", name=tname)
        if col in tbl.columns:
            tbl.columns[col].fk_ref = f"{ref_t}.{ref_c}"
        else:
            tbl.columns[col] = Column(name=col, dtype="(FK)", fk_ref=f"{ref_t}.{ref_c}")
        tables[tname] = tbl


def parse_views(sql: str, tables: dict[str, Table]) -> None:
    for m in re.finditer(
        r"CREATE(?: OR REPLACE)? VIEW\s+app\.(\w+)\s+AS\s+(.*?);",
        sql,
        re.IGNORECASE | re.DOTALL,
    ):
        vname = m.group(1)
        tables[vname] = Table(
            schema="app",
            name=vname,
            is_view=True,
            view_sql=re.sub(r"\s+", " ", m.group(2).strip())[:500],
        )


def parse_migrations() -> dict[str, Table]:
    tables: dict[str, Table] = {}
    for path in sorted(MIGRATIONS.glob("*.sql")):
        sql = strip_sql_comments(path.read_text(encoding="utf-8", errors="replace"))
        parse_create_table(sql, tables)
        parse_alter_table(sql, tables)
        parse_views(sql, tables)
    return tables


def collect_fk_relationships(tables: dict[str, Table]) -> list[tuple[str, str, str, str]]:
    rels: list[tuple[str, str, str, str]] = []
    for tname, tbl in sorted(tables.items()):
        if tbl.is_view:
            continue
        for col in tbl.columns.values():
            if col.fk_ref:
                ref_t, _, ref_c = col.fk_ref.partition(".")
                rels.append((tname, col.name, ref_t, ref_c))
    return rels


def module_for(name: str, is_view: bool) -> tuple[str, str]:
    if is_view and name in VIEW_MODULES:
        return VIEW_MODULES[name]
    if name in MODULE_MAP:
        return MODULE_MAP[name]
    prefix = "13. Other"
    return (prefix, name)


def main() -> None:
    today = date.today().isoformat()
    tables = parse_migrations()
    table_count = sum(1 for t in tables.values() if not t.is_view)
    view_count = sum(1 for t in tables.values() if t.is_view)
    fk_rels = collect_fk_relationships(tables)

    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PM Pepsi App\nเอกสาร ER-Diagram ฉบับละเอียดและสมบูรณ์\n(PostgreSQL · schema app)")
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0\n"
        f"แหล่งข้อมูล: database/migrations/ ({len(list(MIGRATIONS.glob('*.sql')))} ไฟล์)"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายโครงสร้างฐานข้อมูล PostgreSQL ของ PM-Pepsi-App "
        "รวมแผนภาพ ER โมดูล ความสัมพันธ์ (Foreign Key) รายละเอียดตาราง/คอลัมน์ "
        "และ View ที่ใช้ในแอป — สร้างอัตโนมัติจาก migration SQL"
    )

    doc.add_page_break()

    # 1 Overview
    doc.add_heading("1. ภาพรวมฐานข้อมูล", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["DBMS", "PostgreSQL 14+"],
            ["Schema", "app"],
            ["ตาราง (Tables)", str(table_count)],
            ["View", str(view_count)],
            ["Foreign Key", str(len(fk_rels))],
            ["Migration files", str(len(list(MIGRATIONS.glob('*.sql'))))],
            ["ศูนย์กลาง ERD", "app.tbiw37n (Work Order / IW37N)"],
        ],
    )

    doc.add_heading("1.1 หลักการออกแบบ", level=2)
    add_bullets(
        doc,
        [
            "สืบทอดจากระบบ legacy PHP + MySQL (sap_lay) — ชื่อตาราง/คอลัมน์คงรูปแบบเดิม",
            "Work Order หลักอยู่ที่ tbiw37n — ตารางย่อยส่วนใหญ่อ้างอิง idiw37",
            "รูป confirmation เก็บใน BYTEA (tbconfirmimg.img_data) ไม่ใช่ไฟล์บนดิสก์",
            "RBAC แยก tbl_role, tbl_permission, tbl_role_permission",
            "เมนูและสิทธิ์ UI ผูกกับ tbmenu + permission code",
            "วันที่ legacy บางฟิลด์ใช้ BIGINT (unix epoch) เช่น bscstart, cday",
        ],
    )

    doc.add_heading("1.2 โมดูลหลัก", level=2)
    modules: dict[str, list[str]] = defaultdict(list)
    for name, tbl in sorted(tables.items()):
        mod, _ = module_for(name, tbl.is_view)
        kind = "VIEW" if tbl.is_view else "TABLE"
        modules[mod].append(f"{name} ({kind})")
    add_table(
        doc,
        ["โมดูล", "จำนวนอ็อบเจ็กต์", "ตัวอย่าง"],
        [
            [mod, str(len(items)), ", ".join(items[:4]) + ("…" if len(items) > 4 else "")]
            for mod, items in sorted(modules.items())
        ],
    )

    doc.add_page_break()

    # 2 ER Diagram image
    doc.add_heading("2. แผนภาพ ER-Diagram รวม", level=1)
    doc.add_paragraph(
        "แผนภาพด้านล่างแสดงความสัมพันธ์ระหว่างตารางทั้งระบบ "
        "(อ้างอิงจากโครงสร้าง PM Pepsi / legacy SAP PM)"
    )
    if ER_IMAGE.exists():
        doc.add_picture(str(ER_IMAGE), width=Inches(6.5))
        cap = doc.add_paragraph("รูปที่ 1 — ER-Diagram รวมทั้งระบบ")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("(ไม่พบไฟล์แผนภาพ — วางที่ docs/customer-requirements/diagrams/er-diagram-full.png)")

    doc.add_page_break()

    # 3 Hub diagram text
    doc.add_heading("3. ศูนย์กลางความสัมพันธ์ (tbiw37n)", level=1)
    doc.add_paragraph(
        "ตาราง app.tbiw37n เป็น hub หลัก — ตารางด้านล่างมี Foreign Key ไปยัง idiw37:"
    )
    hub_children = sorted({r[0] for r in fk_rels if r[2] == "tbiw37n"})
    add_bullets(doc, [f"app.{t}.{next(c for a, c, b, d in fk_rels if a == t and b == 'tbiw37n')}" for t in hub_children])

    doc.add_paragraph()
    doc.add_paragraph("ความสัมพันธ์อื่นที่สำคัญ:")
    other_rels = [r for r in fk_rels if r[2] != "tbiw37n"]
    add_table(
        doc,
        ["จากตาราง", "คอลัมน์", "อ้างอิง", "คอลัมน์"],
        [[a, b, c, d] for a, b, c, d in other_rels],
        header_fill="FFF2CC",
    )

    doc.add_page_break()

    # 4 Full FK matrix
    doc.add_heading("4. ตารางความสัมพันธ์ Foreign Key ทั้งหมด", level=1)
    add_table(
        doc,
        ["ลำดับ", "ตารางต้นทาง", "FK Column", "ตารางปลายทาง", "PK Column"],
        [[str(i + 1), a, b, c, d] for i, (a, b, c, d) in enumerate(fk_rels)],
    )

    doc.add_page_break()

    # 5 Per-module table details
    doc.add_heading("5. รายละเอียดตารางและ View แยกตามโมดูล", level=1)

    by_module: dict[str, list[tuple[str, Table]]] = defaultdict(list)
    for name, tbl in sorted(tables.items()):
        mod, _ = module_for(name, tbl.is_view)
        by_module[mod].append((name, tbl))

    section = 0
    for mod in sorted(by_module.keys()):
        section += 1
        doc.add_heading(f"5.{section} {mod}", level=2)
        for tname, tbl in by_module[mod]:
            _, desc = module_for(tname, tbl.is_view)
            kind = "VIEW" if tbl.is_view else "TABLE"
            doc.add_heading(f"app.{tname} ({kind})", level=3)
            doc.add_paragraph(desc)

            if tbl.is_view:
                if tbl.view_sql:
                    p = doc.add_paragraph()
                    r = p.add_run(f"SQL: {tbl.view_sql}…")
                    r.font.size = Pt(8)
                    r.italic = True
                continue

            if not tbl.columns:
                doc.add_paragraph("(ไม่มีคอลัมน์ใน migration — อาจเป็น legacy)")
                continue

            rows = []
            for cname, col in sorted(tbl.columns.items(), key=lambda x: (not x[1].is_pk, x[0])):
                flags = []
                if col.is_pk:
                    flags.append("PK")
                if col.fk_ref:
                    flags.append(f"FK→{col.fk_ref}")
                if not col.nullable:
                    flags.append("NOT NULL")
                rows.append([
                    cname,
                    col.dtype,
                    ", ".join(flags) or "—",
                    col.default or "—",
                ])
            add_table(doc, ["คอลัมน์", "ชนิดข้อมูล", "ข้อจำกัด", "Default"], rows)

        if section < len(by_module):
            doc.add_page_break()

    doc.add_page_break()

    # 6 Views summary
    doc.add_heading("6. สรุป Database Views", level=1)
    view_rows = []
    for name, tbl in sorted(tables.items()):
        if not tbl.is_view:
            continue
        _, desc = module_for(name, True)
        view_rows.append([name, desc, (tbl.view_sql or "")[:120] + "…"])
    add_table(doc, ["View", "วัตถุประสงค์", "SQL (ย่อ)"], view_rows)

    doc.add_page_break()

    # 7 Key tables deep dive
    doc.add_heading("7. ตารางสำคัญ — คำอธิบายเชิงธุรกิจ", level=1)

    key_tables = [
        ("tbiw37n", "Work Order จาก SAP IW37N — wkorder, wktype, systemstatus, wkctr, functionalloc, equipment, bscstart"),
        ("tbmoveplan", "วันที่แสดงบนปฏิทิน (cday) และประวัติการย้ายแผน"),
        ("tbplangingwork", "การมอบหมายช่าง/ทีมต่อ WO ใน Planning"),
        ("tbcofirm", "ข้อมูล Confirmation ก่อน export กลับ SAP"),
        ("tbconfirmimg", "รูปถ่ายก่อน/หลัง PM — เก็บ BYTEA ใน DB"),
        ("tbmanhours", "ชั่วโมงทำงานต่อ work center / วัน"),
        ("tbworkcenter", "บัญชีผู้ใช้ + ข้อมูลช่าง (login idwkctr)"),
        ("tbl_role_permission", "กำหนดสิทธิ์ตามบทบาท"),
        ("tbiw37n_import_batch", "batch นำเข้า Excel IW37N"),
        ("integration_job", "คิวงาน integration (import/export)"),
    ]
    add_table(doc, ["ตาราง", "คำอธิบายธุรกิจ"], key_tables)

    doc.add_heading("7.1 Data flow หลัก", level=2)
    add_bullets(
        doc,
        [
            "Import IW37N (Excel) → tbiw37n_import_batch / tbiw37n_import_row → tbiw37n",
            "Planning assign → tbplangingwork (FK idiw37)",
            "Calendar display → view_order (tbiw37n + tbmoveplan + tbwkstatus)",
            "Confirmation → tbcofirm + tbconfirmcom + tbconfirmimg",
            "Export SAP → view_exportconfirm",
            "PM Execution → tbwo_pm_note / tbwo_pm_reading / tbwo_pm_note_entry",
        ],
    )

    doc.add_page_break()

    # 8 Indexes & maintenance
    doc.add_heading("8. Migration และการบำรุงรักษา", level=1)
    add_bullets(
        doc,
        [
            "รัน migration ตามลำดับเลขใน database/migrations/",
            "ตรวจสอบ schema: GET /api/v1/admin/health หรือ query information_schema",
            "หลังเพิ่มตารางใหม่: อัปเดต migration + รันสคริปต์ generate_er_diagram_docx.py",
            "Backup: tbl_backup_history บันทึกประวัติ pg_dump",
            "Audit: tbl_audit_log บันทึกการเปลี่ยนแปลง master data / admin",
        ],
    )

    doc.add_heading("9. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "database/migrations/*.sql — แหล่ง DDL จริง",
            "docs/customer-requirements/UNIT-TEST-DETAILED-TH.docx",
            "docs/customer-requirements/STRESS-TEST-DETAILED-TH.docx",
            "docs/USER-MANUAL-TH.md — คู่มือผู้ใช้แต่ละโมดูล",
            "from customer/tools/generate_er_diagram_docx.py — สคริปต์สร้างเอกสารฉบับนี้",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Tables: {table_count}, Views: {view_count}, FKs: {len(fk_rels)}")


if __name__ == "__main__":
    main()
