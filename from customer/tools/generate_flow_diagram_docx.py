#!/usr/bin/env python3
"""Generate PM Pepsi App — detailed Flow Diagrams with Swim Lanes (Word .docx)."""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "customer-requirements" / "FLOW-DIAGRAM-SWIMLANE-TH.docx"

LANE_COLORS = {
    "user": "D9E2F3",
    "planner": "D9E2F3",
    "technician": "E2EFDA",
    "admin": "FCE4D6",
    "supervisor": "FFF2CC",
    "frontend": "DDEBF7",
    "backend": "E4DFEC",
    "database": "F2F2F2",
    "sap": "FBE5D6",
    "telegram": "D0CECE",
    "filesystem": "EDEDED",
    "integration": "D6DCE4",
    "external": "F8CBAD",
}


@dataclass
class SwimLane:
    name: str
    role: str
    color_key: str
    steps: list[str]  # one per column; "" = empty


@dataclass
class Flow:
    code: str
    title: str
    purpose: str
    actors: list[str]
    preconditions: list[str]
    lanes: list[SwimLane]
    apis: list[str] = field(default_factory=list)
    tables: list[str] = field(default_factory=list)
    exceptions: list[str] = field(default_factory=list)
    pages: list[str] = field(default_factory=list)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "B4C6E7",
    font_size: int = 9,
) -> None:
    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_swimlane_diagram(doc: Document, flow: Flow) -> None:
    n_steps = max((len(lane.steps) for lane in flow.lanes), default=0)
    if n_steps == 0:
        return

    headers = ["Swim Lane \\ ขั้นตอน"] + [str(i + 1) for i in range(n_steps)]
    table = doc.add_table(rows=1 + len(flow.lanes), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "4472C4")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)

    # Lane rows
    for row_idx, lane in enumerate(flow.lanes):
        row = table.rows[row_idx + 1]
        lane_label = f"{lane.name}\n({lane.role})"
        row.cells[0].text = lane_label
        set_cell_shading(row.cells[0], LANE_COLORS.get(lane.color_key, "FFFFFF"))
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

        steps = lane.steps + [""] * (n_steps - len(lane.steps))
        for col_idx in range(n_steps):
            cell = row.cells[col_idx + 1]
            text = steps[col_idx].strip()
            cell.text = text if text else "—"
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            if text:
                set_cell_shading(cell, "FFFFFF")
            else:
                set_cell_shading(cell, "FAFAFA")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("หมายเหตุ: อ่านจากซ้าย→ขวา แต่ละแถว = 1 swim lane (ผู้รับผิดชอบ)").italic = True
    note.runs[0].font.size = Pt(8)


def build_flows() -> list[Flow]:
    return [
        Flow(
            code="F01",
            title="เข้าสู่ระบบ · Session · RBAC",
            purpose="ผู้ใช้ login ผ่าน work center ได้ JWT/session และโหลด permissions สำหรับเมนู",
            actors=["ผู้ใช้ (ช่าง/Planner/Admin)", "Frontend React", "Backend Express", "PostgreSQL"],
            preconditions=["มีบัญชีใน tbworkcenter", "Backend + DB พร้อม", "ไม่ถูก block IP"],
            pages=["/login", "/portal (ถ้ามีหลาย module)"],
            lanes=[
                SwimLane("ผู้ใช้", "User", "user", [
                    "เปิด /login\nกรอก idwkctr + รหัสผ่าน",
                    "กดเข้าสู่ระบบ",
                    "",
                    "ใช้งานแอปตามเมนูที่แสดง",
                ]),
                SwimLane("Frontend", "React + Vite", "frontend", [
                    "แสดงฟอร์ม login",
                    "POST /api/v1/auth/login",
                    "เก็บ token/session\nGET /auth/me",
                    "filterNavForUser()\nตาม permissions[]",
                ]),
                SwimLane("Backend API", "Express", "backend", [
                    "",
                    "ตรวจรหัสผ่าน\nrate-limit\nlogin-lockout",
                    "คืน JWT + user profile\npermissions จาก RBAC",
                    "middleware requirePermission\nทุก route",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "อ่าน tbworkcenter\nบันทึก tbworkcenter_userlog",
                    "อ่าน tbl_role_permission\ntbl_permission",
                    "",
                ]),
            ],
            apis=["POST /api/v1/auth/login", "GET /api/v1/auth/me", "POST /api/v1/auth/logout"],
            tables=["tbworkcenter", "tbworkcenter_userlog", "tbl_role", "tbl_permission", "tbl_role_permission"],
            exceptions=[
                "รหัสผิด → 401 + rate limit",
                "Maintenance mode → 503 MAINTENANCE",
                "must_change_password → บังคับเปลี่ยนรหัส",
            ],
        ),
        Flow(
            code="F02",
            title="นำเข้า IW37N (Manual Upload)",
            purpose="Planner/Admin อัปโหลด Excel IW37N เข้าระบบเป็น Work Order",
            actors=["Planner/Admin", "Frontend /iw37n", "Backend parser", "PostgreSQL"],
            preconditions=["สิทธิ์ iw37n.import", "ไฟล์ Excel ตาม template ลูกค้า"],
            pages=["/iw37n", "/integration"],
            lanes=[
                SwimLane("Planner/Admin", "User", "planner", [
                    "เลือกไฟล์ Excel IW37N",
                    "ดู Preview แถวที่ parse ได้",
                    "ยืนยัน Commit import",
                    "ตรวจผล batch สำเร็จ/ผิดพลาด",
                ]),
                SwimLane("Frontend", "/iw37n", "frontend", [
                    "อัปโหลด multipart",
                    "แสดงตาราง preview + error",
                    "POST commit batch",
                    "ลิงก์ไป backlog/planning",
                ]),
                SwimLane("Backend API", "Import service", "backend", [
                    "รับไฟล์ → parser",
                    "validate factory scope\nwktype, duplicate",
                    "INSERT tbiw37n_import_batch/row\nUPSERT tbiw37n",
                    "สรุป inserted/updated/skipped",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "",
                    "tbiw37n_import_batch\ntbiw37n_import_row\ntbiw37n",
                    "tbl_audit_log",
                ]),
                SwimLane("SAP (ไฟล์)", "External", "sap", [
                    "Export IW37N\nจาก SAP → Excel",
                    "",
                    "",
                    "",
                ]),
            ],
            apis=["POST /api/v1/iw37n/import/preview", "POST /api/v1/iw37n/import/commit"],
            tables=["tbiw37n", "tbiw37n_import_batch", "tbiw37n_import_row"],
            exceptions=["แถวซ้ำ → skip", "factory ไม่ตรง → reject batch", "schema error → 503"],
        ),
        Flow(
            code="F03",
            title="นำเข้า IW37N (Auto Integration Watch)",
            purpose="Scheduler สแกนโฟลเดอร์ inbound นำเข้า IW37N อัตโนมัติ",
            actors=["SAP/ไฟล์", "Integration Worker", "PostgreSQL", "Admin"],
            preconditions=["INTEGRATION_WATCH_SCHEDULER≠0", "โฟลเดอร์ inbound/iw37n พร้อม"],
            pages=["/integration", "/admin/health"],
            lanes=[
                SwimLane("SAP / ไฟล์", "External", "sap", [
                    "วางไฟล์ CSV/XLSX\nใน inbound/iw37n",
                    "",
                    "",
                ]),
                SwimLane("File System", "inbound/", "filesystem", [
                    "ตรวจจับไฟล์ใหม่",
                    "ย้าย archive/\nหรือ error/",
                    "",
                ]),
                SwimLane("Integration Worker", "Scheduler", "integration", [
                    "integration-scheduler\nทุก 1 นาที",
                    "สร้าง integration_job\nstatus=running",
                    "เรียก parser เดียวกับ manual\nSHA256 idempotent",
                    "job status=done/failed",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "integration_job",
                    "tbiw37n + batch/row",
                    "tbl_audit_log",
                ]),
                SwimLane("Admin", "User", "admin", [
                    "",
                    "",
                    "",
                    "ดูสถานะที่ /integration",
                ]),
            ],
            apis=["GET /api/v1/integration/jobs", "POST /api/v1/integration/scan"],
            tables=["integration_job", "tbiw37n_import_batch", "tbiw37n"],
            exceptions=["ไฟล์ซ้ำ SHA256 → skip", "parse fail → error folder + job failed"],
        ),
        Flow(
            code="F04",
            title="จ่ายงาน Planning (Assign WO)",
            purpose="Planner มอบหมาย WO ให้ work center / ทีม / หลายคน",
            actors=["Planner", "Frontend /planning", "Backend", "PostgreSQL", "Telegram"],
            preconditions=["สิทธิ์ planning.assign", "WO สถานะเปิดอยู่"],
            pages=["/planning", "/plan-calendar"],
            lanes=[
                SwimLane("Planner", "Engineering", "planner", [
                    "เปิด /planning\nเลือก WO",
                    "เลือกช่าง/ทีม\n(multi-assign)",
                    "ยืนยันจ่ายงาน",
                    "ติดตามสถานะบนปฏิทิน",
                ]),
                SwimLane("Frontend", "/planning", "frontend", [
                    "แสดงรายการ open orders",
                    "PlanningMultiAssign dialog",
                    "POST assign API",
                    "invalidate calendar cache",
                ]),
                SwimLane("Backend API", "Planning", "backend", [
                    "",
                    "ตรวจ canAssign\nRBAC",
                    "INSERT/UPDATE tbplangingwork",
                    "ส่ง notify Telegram\n(ถ้าผูกแล้ว)",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "",
                    "tbplangingwork\n(FK idiw37)",
                    "view_planwork",
                ]),
                SwimLane("Telegram Bot", "Notify", "telegram", [
                    "",
                    "",
                    "ข้อความ + ปุ่ม\nรับทราบงาน",
                    "",
                ]),
            ],
            apis=["GET /api/v1/planning/orders", "POST /api/v1/planning/assign"],
            tables=["tbplangingwork", "tbiw37n", "view_planwork"],
            exceptions=["ไม่มีสิทธิ์ → 403", "WO ปิดแล้ว → 400"],
        ),
        Flow(
            code="F05",
            title="ปฏิทิน — ย้ายวันแผน (Move Plan)",
            purpose="Planner ลาก event บนปฏิทินเปลี่ยนวันทำงาน WO",
            actors=["Planner", "Frontend /calendar", "Backend", "PostgreSQL"],
            preconditions=["สิทธิ์ calendar write / planning.assign", "WO ยังไม่ปิด"],
            pages=["/calendar", "/calendar/wc/:code"],
            lanes=[
                SwimLane("Planner", "User", "planner", [
                    "เปิด /calendar\nกรอง Z1/Z2/Z5",
                    "ลาก event ไปวันใหม่",
                    "กรอกเหตุผล (ถ้าต้องการ)",
                    "ตรวจสอบสี/ suffix workflow",
                ]),
                SwimLane("Frontend", "FullCalendar", "frontend", [
                    "POST calendar/events",
                    "eventDrop handler",
                    "POST move-plan API",
                    "รีเฟรช events",
                ]),
                SwimLane("Backend API", "Calendar", "backend", [
                    "calendar-move-policy",
                    "ตรวจ canMovePlan",
                    "UPSERT tbmoveplan\ncday, reasoncode",
                    "คืน event ใหม่",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "view_order",
                    "",
                    "tbmoveplan\n(FK idiw37)",
                    "tbiw37n",
                ]),
            ],
            apis=["POST /api/v1/calendar/events", "POST /api/v1/calendar/move"],
            tables=["tbmoveplan", "tbiw37n", "view_order"],
            exceptions=["นโยบายห้ามย้าย → 409", "มือถือ: long-press 0.4s ก่อนลาก"],
        ),
        Flow(
            code="F06",
            title="ช่างปิดงาน — รูป · เวลา · Supervisor Close",
            purpose="ช่างบันทึกรูป Before/After เวลาทำงาน และปิดงาน supervisor",
            actors=["ช่าง", "หัวหน้างาน", "Frontend WO modal", "Backend", "PostgreSQL"],
            preconditions=["สิทธิ์ work-orders.read", "WO ถูก assign แล้ว"],
            pages=["/work-orders", "/calendar", "/personnel/confirm"],
            lanes=[
                SwimLane("ช่าง", "Technician", "technician", [
                    "เปิด WO modal\nแท็บ Confirm",
                    "อัปโหลดรูป Before/After",
                    "บันทึกเวลาช่าง\n(tbwrkclose)",
                    "หัวหน้าปิดงาน supervisor",
                ]),
                SwimLane("Frontend", "WO Dialog", "frontend", [
                    "โหลดรูปเมื่อเปิดแท็บ",
                    "WebP compress → API",
                    "ฟอร์มเวลา/close",
                    "POST confirm APIs",
                ]),
                SwimLane("Backend API", "Confirmation", "backend", [
                    "",
                    "เก็บ BYTEA\ntbconfirmimg",
                    "INSERT tbwrkclose",
                    "INSERT tbcofirm\nconfirm_qc_status=pending",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "tbconfirmimg",
                    "tbwrkclose",
                    "tbcofirm\ntbiw37n.confirm_qc_status",
                ]),
            ],
            apis=[
                "POST /api/v1/confirmation/:idiw37/images",
                "POST /api/v1/confirmation/closes",
                "POST /api/v1/confirmation/:idiw37/close",
            ],
            tables=["tbconfirmimg", "tbwrkclose", "tbcofirm", "tbiw37n"],
            exceptions=["รูปเกิน limit → 413", "WO ปิดแล้ว → 409"],
        ),
        Flow(
            code="F07",
            title="Admin QC — อนุมัติก่อนนับ KPI / Export",
            purpose="Admin ตรวจคุณภาพ confirmation ก่อนนับ Dashboard และส่ง SAP",
            actors=["Admin", "ช่าง", "Frontend", "Backend", "PostgreSQL"],
            preconditions=["confirm_qc_status=pending", "สิทธิ์ confirmation.import"],
            pages=["/confirmation", "WO modal แท็บ Confirm", "/personnel/confirm"],
            lanes=[
                SwimLane("ช่าง", "Technician", "technician", [
                    "บันทึกปิดงานเสร็จ",
                    "รอ QC",
                    "ถ้า rejected → แก้รูป/เวลา",
                    "ส่ง pending อีกครั้ง",
                ]),
                SwimLane("Admin", "QC", "admin", [
                    "เปิดคิว QC\n/confirmation",
                    "ตรวจรูป/เวลา/close",
                    "กดอนุมัติ หรือส่งกลับ",
                    "",
                ]),
                SwimLane("Frontend", "ConfirmQcPanel", "frontend", [
                    "ConfirmQcPendingQueue",
                    "แสดงรายละเอียด WO",
                    "POST approve/reject",
                    "อัปเดต KPI cards",
                ]),
                SwimLane("Backend API", "QC", "backend", [
                    "GET qc/pending",
                    "",
                    "approve → approved\nreject → rejected + note",
                    "dashboard/summary\nนับเฉพาะ approved",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "confirm_qc_status=pending",
                    "",
                    "tbiw37n.confirm_qc_status",
                    "view_countpersonelclose",
                ]),
            ],
            apis=[
                "GET /api/v1/confirmation/qc/pending",
                "POST /api/v1/confirmation/:idiw37/qc/approve",
                "POST /api/v1/confirmation/:idiw37/qc/reject",
                "POST /api/v1/confirmation/qc/approve-batch",
            ],
            tables=["tbiw37n", "tbcofirm", "view_countpersonelclose"],
            exceptions=["Mass Confirm 44 WO max ต่อชุด SAP"],
        ),
        Flow(
            code="F08",
            title="Mass Confirm → Export CONFIRM_OUT → SAP",
            purpose="ปิดงานหลาย WO พร้อมกัน อนุมัติ QC แล้ว export ไฟล์กลับ SAP",
            actors=["Admin/Planner", "Frontend", "Backend", "SAP"],
            preconditions=["WO พร้อม export", "QC approved", "สิทธิ์ confirmation.export"],
            pages=["/confirmation", "/confirmation/export", "/integration"],
            lanes=[
                SwimLane("Admin/Planner", "User", "admin", [
                    "เลือก WO ≤44 รายการ",
                    "Mass Confirm batch",
                    "อนุมัติ QC ทั้งชุด",
                    "ดาวน์โหลด CSV/XLSX\nส่ง SAP",
                ]),
                SwimLane("Frontend", "/confirmation", "frontend", [
                    "MassConfirmExportPanel",
                    "POST closes/batch",
                    "แสดงสถานะ QC ต่อ WO",
                    "GET export.csv?idiw37n=",
                ]),
                SwimLane("Backend API", "Export", "backend", [
                    "batch close guard",
                    "tbcofirm batch",
                    "approve-batch QC",
                    "view_exportconfirm → CSV",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "tbcofirm\ntbwrkclose",
                    "confirm_qc_status",
                    "view_exportconfirm",
                ]),
                SwimLane("SAP", "External", "sap", [
                    "",
                    "",
                    "",
                    "รับไฟล์ CONFIRM_OUT",
                ]),
            ],
            apis=[
                "POST /api/v1/confirmation/closes/batch",
                "GET /api/v1/confirmation/export.csv",
                "POST /api/v1/confirmation/qc/approve-batch",
            ],
            tables=["tbcofirm", "view_exportconfirm"],
            exceptions=["มี WO ยัง pending QC → ไม่ export", "เกิน 44 → แบ่งชุด"],
        ),
        Flow(
            code="F09",
            title="Telegram — แจ้งจ่ายงาน · รับทราบ",
            purpose="แจ้งช่างผ่าน Telegram เมื่อถูก assign และรับทราบงาน",
            actors=["Planner", "ช่าง", "Telegram Bot", "Web App", "PostgreSQL"],
            preconditions=["ช่างผูก Telegram ที่ /admin/telegram", "Bot token ตั้งค่าแล้ว"],
            pages=["/admin/telegram", "/planning"],
            lanes=[
                SwimLane("Planner", "Engineering", "planner", [
                    "จ่ายงานใน /planning",
                    "",
                    "ได้รับแจ้งว่าช่าง X\nรับทราบ WO Y",
                    "",
                ]),
                SwimLane("Web App", "PM-Pepsi-App", "frontend", [
                    "บันทึก tbplangingwork",
                    "เรียก telegram notify",
                    "",
                    "อัปเดตสถานะใน WO",
                ]),
                SwimLane("Telegram Bot", "Bot API", "telegram", [
                    "",
                    "ส่งข้อความ+ปุ่ม\nรับทราบงาน",
                    "รับ callback\nจากช่าง",
                    "แจ้ง Planner",
                ]),
                SwimLane("ช่าง", "Technician", "technician", [
                    "",
                    "ได้รับข้อความ\nใน Telegram",
                    "กด รับทราบงาน",
                    "ดูงานบนเว็บ",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "tbplangingwork",
                    "tbl_telegram_notify_group",
                    "tbl_telegram_chat_session",
                    "",
                ]),
            ],
            apis=["POST /api/v1/planning/assign", "POST /api/v1/telegram/webhook"],
            tables=["tbl_telegram_notify_group", "tbl_telegram_chat_session", "tbplangingwork"],
            exceptions=["ช่างยังไม่ผูก Telegram → ข้าม notify", "Bot ล่ม → log error ไม่ block assign"],
        ),
        Flow(
            code="F10",
            title="Portal หลายโมดูล (Unified Login)",
            purpose="Login จุดเดียว → เลือก module PM / สโตร์ / แจ้งซ่อม ตามสิทธิ์",
            actors=["ผู้ใช้", "Frontend /portal", "Backend auth", "โมดูลภายนอก"],
            preconditions=["tbl_app_module seed แล้ว", "สิทธิ์ portal ต่อ module"],
            pages=["/login", "/portal"],
            lanes=[
                SwimLane("ผู้ใช้", "User", "user", [
                    "Login /login",
                    "เห็นการ์ด module\nที่มีสิทธิ์",
                    "คลิก PM Pepsi",
                    "ทำงานในแอป PM",
                ]),
                SwimLane("Frontend", "/portal", "frontend", [
                    "auth success",
                    "GET /portal/modules",
                    "นำทาง / หรือ /planning",
                    "App shell PM",
                ]),
                SwimLane("Backend API", "Auth+Portal", "backend", [
                    "JWT + permissions",
                    "กรอง tbl_app_module\nตาม perm_code",
                    "",
                    "RBAC ต่อ API",
                ]),
                SwimLane("โมดูลอื่น", "External", "external", [
                    "",
                    "การ์ดสโตร์/แจ้งซ่อม\n(handoff code)",
                    "redirect URL\nmodule handoff",
                    "",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "tbworkcenter",
                    "tbl_app_module\ntbl_module_handoff_code",
                    "",
                    "",
                ]),
            ],
            apis=["GET /api/v1/portal/modules", "POST /api/v1/portal/handoff"],
            tables=["tbl_app_module", "tbl_module_handoff_code", "tbl_permission"],
            exceptions=["สิทธิ์ module เดียว → ข้าม portal", "handoff code หมดอายุ → 401"],
        ),
        Flow(
            code="F11",
            title="PM Execution — บันทึกค่าวัด / Vibration",
            purpose="ช่างบันทึกค่าวัด PM 3 phase ต่อ WO",
            actors=["ช่าง", "Frontend /pm-vibration", "Backend", "PostgreSQL"],
            preconditions=["WO ประเภท PM", "สิทธิ์ work-orders.write"],
            pages=["/pm-vibration", "/work-orders"],
            lanes=[
                SwimLane("ช่าง", "Technician", "technician", [
                    "เปิด /pm-vibration\nหรือ WO modal",
                    "กรอกค่าวัด\nตามขั้นตอน PM",
                    "บันทึก note รายการ",
                    "ส่งข้อมูล",
                ]),
                SwimLane("Frontend", "PM Forms", "frontend", [
                    "แสดงฟอร์ม PM phase",
                    "validation Zod",
                    "POST readings/notes",
                    "แสดง phase badge",
                ]),
                SwimLane("Backend API", "PM Execution", "backend", [
                    "",
                    "wo-pm-phase rules",
                    "UPSERT tbwo_pm_reading\ntbwo_pm_note_entry",
                    "อัปเดต execution state",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "",
                    "tbwo_pm_reading\ntbwo_pm_note\ntbwo_pm_note_entry",
                    "",
                ]),
            ],
            apis=["GET/POST /api/v1/work-orders/:id/pm-readings", "POST /api/v1/work-orders/:id/pm-notes"],
            tables=["tbwo_pm_reading", "tbwo_pm_note", "tbwo_pm_note_entry", "tbiw37n"],
            exceptions=["WO ไม่ใช่ PM → ซ่อนฟอร์ม", "QC ปิดแล้ว → read-only"],
        ),
        Flow(
            code="F12",
            title="Manhour HR — นำเข้า · Utilization รายสัปดาห์",
            purpose="HR นำเข้าชั่วโมงทำงานและดูรายงาน utilization",
            actors=["HR Admin", "Frontend", "Backend", "PostgreSQL"],
            preconditions=["สิทธิ์ manhours.admin", "ไฟล์ Excel manhour"],
            pages=["/manhours/admin", "/summary-weekly", "/manhours-hr"],
            lanes=[
                SwimLane("HR Admin", "User", "admin", [
                    "อัปโหลด Excel\n/manhours/admin",
                    "ตรวจ parse สำเร็จ",
                    "เปิด /summary-weekly",
                    "พิมพ์กราฟประชุม",
                ]),
                SwimLane("Frontend", "Reports", "frontend", [
                    "upload + preview",
                    "POST import",
                    "Eng utilization chart",
                    "print-friendly view",
                ]),
                SwimLane("Backend API", "Manhour", "backend", [
                    "manhours-parse",
                    "INSERT tbmanhours",
                    "week-to-week KPI",
                    "reports/kpi",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "tbmanhours\n(FK idwkctr)",
                    "aggregates",
                    "",
                ]),
            ],
            apis=["POST /api/v1/manhours/import", "GET /api/v1/reports/kpi", "GET /api/v1/manhours-hr/utilization"],
            tables=["tbmanhours", "tbworkcenter"],
            exceptions=["รูปแบบวันที่ผิด → แถว error", "wkctr ไม่มีในระบบ → skip"],
        ),
        Flow(
            code="F13",
            title="Admin Backup & Restore",
            purpose="Admin สำรองและกู้คืนฐานข้อมูล PostgreSQL",
            actors=["Admin", "Frontend /admin/backup", "Backend", "File system"],
            preconditions=["สิทธิ์ admin.backup.write", "พื้นที่ดิสก์เพียงพอ"],
            pages=["/admin/backup"],
            lanes=[
                SwimLane("Admin", "User", "admin", [
                    "เปิด /admin/backup",
                    "กด Backup now",
                    "ดาวน์โหลด .tar.gz",
                    "Restore (พิมพ์ RESTORE)",
                ]),
                SwimLane("Frontend", "Admin UI", "frontend", [
                    "แสดงประวัติ backup",
                    "POST /admin/backup",
                    "download link",
                    "ConfirmPhraseDialog",
                ]),
                SwimLane("Backend API", "Backup", "backend", [
                    "",
                    "pg_dump + advisory lock",
                    "tbl_backup_history",
                    "restore + confirm phrase",
                ]),
                SwimLane("File System", "Disk", "filesystem", [
                    "",
                    "เขียน backup file",
                    "เก็บ path ใน DB",
                    "อ่านไฟล์ restore",
                ]),
                SwimLane("PostgreSQL", "DB", "database", [
                    "",
                    "dump snapshot",
                    "",
                    "replace data",
                ]),
            ],
            apis=["POST /api/v1/admin/backup", "POST /api/v1/admin/backup/restore", "GET /api/v1/admin/backup/history"],
            tables=["tbl_backup_history", "tbl_setting"],
            exceptions=["backup ซ้อน → 409 busy", "restore ผิด phrase → 400"],
        ),
        Flow(
            code="F14",
            title="Workflow วันทำงานเต็มรูปแบบ (End-to-End)",
            purpose="ภาพรวมวันทำงาน Planner → ช่าง → Admin ครบวงจร",
            actors=["Planner", "ช่าง", "Admin", "ระบบ", "SAP"],
            preconditions=["ระบบพร้อม UAT", "สิทธิ์ครบตามบทบาท"],
            pages=["/iw37n", "/planning", "/calendar", "/work-orders", "/confirmation"],
            lanes=[
                SwimLane("Planner", "Engineering", "planner", [
                    "Import IW37N",
                    "จ่ายงาน Planning",
                    "จัดวันปฏิทิน",
                    "Mass Confirm (ถ้าครบ)",
                ]),
                SwimLane("ช่าง", "Technician", "technician", [
                    "",
                    "รับทราบ Telegram",
                    "ปิดงาน+รูป+เวลา",
                    "แก้ถ้า QC reject",
                ]),
                SwimLane("Admin", "QC", "admin", [
                    "",
                    "",
                    "อนุมัติ QC",
                    "Export SAP",
                ]),
                SwimLane("ระบบ", "PM-Pepsi-App", "backend", [
                    "tbiw37n + batch",
                    "tbplangingwork",
                    "tbcofirm pending",
                    "view_exportconfirm",
                ]),
                SwimLane("SAP", "External", "sap", [
                    "IW37N out",
                    "",
                    "",
                    "CONFIRM_IN",
                ]),
            ],
            apis=["(ดู F02–F08)"],
            tables=["tbiw37n", "tbplangingwork", "tbmoveplan", "tbcofirm", "tbconfirmimg"],
            exceptions=["ทุกขั้นต้องผ่าน RBAC และ QC ก่อน export"],
        ),
    ]


def add_flow_section(doc: Document, flow: Flow, section_no: str) -> None:
    doc.add_heading(f"{section_no} {flow.code} — {flow.title}", level=2)
    doc.add_paragraph(flow.purpose)

    doc.add_heading("วัตถุประสงค์และเงื่อนไข", level=3)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["รหัส Flow", flow.code],
            ["หน้าเว็บ", ", ".join(flow.pages) if flow.pages else "—"],
            ["ผู้เกี่ยวข้อง", ", ".join(flow.actors)],
        ],
        header_fill="B4C6E7",
    )

    doc.add_paragraph("เงื่อนไขก่อนเริ่ม:")
    add_bullets(doc, flow.preconditions)

    doc.add_heading("Swim Lane Diagram", level=3)
    add_swimlane_diagram(doc, flow)

    if flow.apis:
        doc.add_heading("API ที่เกี่ยวข้อง", level=3)
        add_bullets(doc, flow.apis)

    if flow.tables:
        doc.add_heading("ตารางฐานข้อมูล", level=3)
        add_bullets(doc, [f"app.{t}" for t in flow.tables])

    if flow.exceptions:
        doc.add_heading("กรณีผิดปกติ / Exception", level=3)
        add_bullets(doc, flow.exceptions)

    doc.add_paragraph()


def main() -> None:
    today = date.today().isoformat()
    flows = build_flows()
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\nเอกสาร Flow Diagrams ฉบับละเอียด\n(Swim Lane · Business Process)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0\n"
        f"จำนวน Flow: {len(flows)} รายการ"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบาย flow ธุรกิจหลักของ PM-Pepsi-App ในรูปแบบ Swim Lane Diagram "
        "แต่ละแถวแทนผู้รับผิดชอบ (Lane) แต่ละคอลัมน์แทนลำดับขั้นตอน (อ่านซ้าย→ขวา) "
        "ครอบคลุมตั้งแต่ Login · Import SAP · Planning · Calendar · Confirmation · QC · Export · Telegram · Portal"
    )

    doc.add_page_break()

    # 1 Overview
    doc.add_heading("1. คู่มือการอ่าน Swim Lane", level=1)
    add_table(
        doc,
        ["สัญลักษณ์/ส่วน", "ความหมาย"],
        [
            ["แถว (Row)", "Swim Lane = ผู้รับผิดชอบ (User, Frontend, Backend, DB, SAP, Telegram)"],
            ["คอลัมน์ (Column)", "ลำดับขั้นตอนเวลา — อ่านจากซ้ายไปขวา"],
            ["เส้นขอบตาราง", "การส่งมอบงานข้าม lane ในขั้นตอนเดียวกัน"],
            ["— (ขีด)", "Lane นั้นไม่มีกิจกรรมในขั้นตอนนั้น"],
            ["สีแถว Lane", "แยกประเภท: ผู้ใช้=ฟ้า, Frontend=ฟ้าอ่อน, Backend=ม่วงอ่อน, DB=เทา, SAP=ส้ม"],
        ],
    )

    doc.add_heading("1.1 สี Swim Lane", level=2)
    add_table(
        doc,
        ["ประเภท Lane", "สี", "ตัวอย่าง"],
        [
            ["ผู้ใช้ / Planner", "ฟ้า", "Planner, Admin, ช่าง"],
            ["Frontend", "ฟ้าอ่อน", "React pages, modals"],
            ["Backend API", "ม่วงอ่อน", "Express routes, services"],
            ["PostgreSQL", "เทา", "ตาราง app.*"],
            ["SAP / External", "ส้ม", "ไฟล์ IW37N, CONFIRM_OUT"],
            ["Telegram", "เทาเข้ม", "Bot notifications"],
        ],
    )

    doc.add_heading("1.2 สารบัญ Flow", level=2)
    add_table(
        doc,
        ["รหัส", "ชื่อ Flow", "หน้าหลัก"],
        [[f.code, f.title, ", ".join(f.pages[:2])] for f in flows],
    )

    doc.add_page_break()

    # 2 Actor matrix
    doc.add_heading("2. บทบาทผู้ใช้ (Actors)", level=1)
    add_table(
        doc,
        ["บทบาท", "หน้าที่หลัก", "หน้าเว็บ", "Permission ตัวอย่าง"],
        [
            ["Planner / Engineering", "Import · จ่ายงาน · จัดปฏิทิน", "/planning, /calendar", "planning.assign, calendar.read"],
            ["ช่าง (Technician)", "รับงาน · ปิดงาน · รูป · PM readings", "/work-orders, /personnel", "work-orders.read"],
            ["หัวหน้างาน", "Supervisor close", "WO modal Confirm", "confirmation.read"],
            ["Admin", "QC · RBAC · Backup · Telegram", "/admin/*, /confirmation", "confirmation.import, admin.*"],
            ["HR", "Manhour import · รายงาน", "/manhours/admin", "manhours.admin"],
            ["Auditor", "Activity log · รายงาน", "/activity-log, /reports/audit", "reports.read"],
        ],
    )

    doc.add_page_break()

    # 3– Flows
    doc.add_heading("3. รายละเอียด Flow แต่ละรายการ", level=1)
    for i, flow in enumerate(flows):
        add_flow_section(doc, flow, f"3.{i + 1}")
        if i < len(flows) - 1:
            doc.add_page_break()

    # Cross-reference
    doc.add_page_break()
    doc.add_heading("4. ความสัมพันธ์ระหว่าง Flow", level=1)
    add_table(
        doc,
        ["ลำดับงาน", "Flow ที่เชื่อมต่อ"],
        [
            ["1", "F01 Login → F10 Portal (ถ้ามีหลาย module) → โมดูล PM"],
            ["2", "F02/F03 Import IW37N → F04 Assign → F05 Calendar"],
            ["3", "F04 Assign → F09 Telegram แจ้งช่าง"],
            ["4", "F06 ช่างปิดงาน → F07 Admin QC → F08 Export SAP"],
            ["5", "F11 PM Execution คู่ขนานกับ F06 สำหรับ WO ประเภท PM"],
            ["6", "F12 Manhour รายสัปดาห์ — อิสระจาก WO cycle"],
            ["7", "F13 Backup — โดย Admin ตามรอบ"],
            ["8", "F14 สรุป End-to-End ทั้งหมด"],
        ],
    )

    doc.add_heading("5. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "docs/USER-MANUAL-TH.md — คู่มือผู้ใช้ทุกหน้า",
            "docs/customer-requirements/CONFIRM-QC-FLOW.md — QC detail",
            "docs/customer-requirements/TELEGRAM-ASSIGNMENT-FLOW.md — Telegram detail",
            "docs/customer-requirements/AUTOMATION-DESIGN.md — Integration pipelines",
            "docs/customer-requirements/ER-DIAGRAM-DETAILED-TH.docx — โครงสร้าง DB",
            "from customer/tools/generate_flow_diagram_docx.py — สคริปต์สร้างเอกสารฉบับนี้",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Flows: {len(flows)}")


if __name__ == "__main__":
    main()
