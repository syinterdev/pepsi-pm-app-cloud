#!/usr/bin/env python3
"""Generate PM Pepsi App — Frontend Design documentation (Word .docx)."""
from __future__ import annotations

import re
from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
FRONTEND = ROOT / "PM-Pepsi-App" / "frontend"
SRC = FRONTEND / "src"
OUT = ROOT / "docs" / "customer-requirements" / "FRONTEND-DESIGN-DETAILED-TH.docx"

ROUTE_PAGES: list[tuple[str, str, str, str]] = [
    ("/login", "LoginPage", "Guest", "เข้าสู่ระบบ"),
    ("/logout", "LogoutPage", "Auth", "ออกจากระบบ"),
    ("/portal", "PortalPage", "Auth", "Portal เลือก module"),
    ("/board", "EngineeringBoardPage", "Public", "จอ TV kiosk — ไม่มี shell"),
    ("/", "HomePage", "AppShell", "Dashboard KPI"),
    ("/plan-calendar", "PlanCalendarPage", "AppShell", "ปฏิทินจ่ายงาน"),
    ("/calendar", "CalendarPage", "AppShell", "ปฏิทิน WO — drag/drop"),
    ("/calendar/wc/:code", "CalendarPage", "AppShell", "ปฏิทินกรอง work center"),
    ("/backlog", "BacklogPage", "AppShell", "WO ค้างแผน"),
    ("/work-orders", "WorkOrdersPage", "AppShell", "ค้นหา WO · batch team"),
    ("/work-orders/:id", "WorkOrdersPage", "AppShell", "เปิด WO modal จาก URL"),
    ("/pm-vibration", "PmVibrationPage", "AppShell", "PM measurements / 3-phase"),
    ("/confirmation", "ConfirmationPage", "AppShell", "Export confirmation · QC queue"),
    ("/planning", "PlanningPage", "AppShell", "แผน PM/CM · assign"),
    ("/integration", "IntegrationPage", "AppShell", "SAP CSV integration jobs"),
    ("/iw37n", "Iw37nPage", "AppShell", "นำเข้า IW37N Excel"),
    ("/master-data", "MasterDataPage", "AppShell", "ข้อมูลหลัก 17 entities"),
    ("/manhours", "ManhoursPage", "AppShell", "ชั่วโมงงาน"),
    ("/worktime", "WorktimePage", "AppShell", "Summary overall / รูปช่าง"),
    ("/personnel", "PersonnelPage", "AppShell", "Personal dashboard"),
    ("/personnel/confirm", "PersonnelConfirmPage", "AppShell", "Personnel confirm %"),
    ("/reports", "ReportsPage", "AppShell", "รายงาน KPI"),
    ("/reports/audit", "AuditorHubPage", "AppShell", "Auditor hub"),
    ("/activity-log", "ActivityLogPage", "AppShell", "Activity log"),
    ("/manhours-hr", "ManhoursHrPage", "AppShell", "Manhour HR"),
    ("/summary-weekly", "SummaryWeeklyPage", "AppShell", "Eng utilization รายสัปดาห์"),
    ("/summary-weekly/chart/full", "SummaryWeeklyChartFullPage", "Auth", "กราฟเต็มจอพิมพ์"),
    ("/user-log", "UserLogPage", "AppShell", "ประวัติ login"),
    ("/settings", "SettingsPage", "AppShell", "ตั้งค่าผู้ใช้ · Telegram link"),
    ("/admin", "AdminConsolePage", "AdminLayout", "Admin hub KPI"),
    ("/admin/branding", "AdminBrandingPage", "AdminLayout", "โลโก้ · สี · theme"),
    ("/admin/settings", "AdminSettingsPage", "AdminLayout", "ตั้งค่าระบบ"),
    ("/admin/master", "AdminMasterHubPage", "AdminLayout", "Master data hub"),
    ("/admin/audit", "AdminAuditPage", "AdminLayout", "Audit log"),
    ("/admin/health", "AdminHealthPage", "AdminLayout", "สุขภาพระบบ"),
    ("/admin/backup", "AdminBackupPage", "AdminLayout", "Backup / restore"),
    ("/admin/announcements", "AdminAnnouncementsPage", "AdminLayout", "ประกาศ"),
    ("/admin/telegram", "AdminTelegramPage", "AdminLayout", "Telegram groups"),
    ("/admin/security", "AdminSecurityPage", "AdminLayout", "ความปลอดภัย"),
    ("/admin/about", "AdminAboutPage", "AdminLayout", "About · license"),
    ("/admin/users", "AdminUsersPage", "AdminLayout", "จัดการผู้ใช้"),
    ("/admin/roles", "AdminRolesPage", "AdminLayout", "บทบาท & สิทธิ์"),
    ("/admin/menu", "AdminMenuPage", "AdminLayout", "Menu builder"),
    ("/dev/ui", "UiPlaygroundPage", "Dev", "UI playground (dev only)"),
    ("/error/:code", "HttpErrorPage", "Public", "หน้า error 403/404/500"),
]

FEATURE_MODULES = [
    ("auth", "Login, logout, AuthGuards, session"),
    ("portal", "Portal การ์ด module หลัง login"),
    ("home", "Dashboard KPI cards"),
    ("calendar", "Work scheduling FullCalendar"),
    ("plan-calendar", "Plan calendar events"),
    ("backlog", "WO backlog filters"),
    ("work-orders", "WO search, filter, batch team"),
    ("confirmation", "Mass confirm, QC queue, export"),
    ("planning", "PM/CM planning assign"),
    ("integration", "Integration job monitor"),
    ("iw37n", "IW37N import preview/commit"),
    ("master-data", "CRUD master entities"),
    ("manhours", "Manhours, worktime, HR charts"),
    ("personnel", "Personnel dashboard & confirm"),
    ("reports", "Reports, activity log, utilization"),
    ("pm-vibration", "PM readings / vibration"),
    ("board", "Engineering board kiosk"),
    ("admin", "Admin console 13 sub-pages"),
    ("settings", "User settings, password, Telegram"),
    ("user-log", "User login history"),
    ("errors", "Error boundary, HttpErrorPage"),
    ("dev", "UiPlayground, sidebar states"),
]

I18N_NS = [
    "common", "nav", "admin", "home", "scheduling", "planning", "confirmation",
    "workOrders", "integration", "masterData", "manhours", "personnel", "reports",
    "portal", "board", "pmVibration", "userLog", "errors",
]


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "F4B084",
    font_size: int = 8,
) -> None:
    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def count_ts_files(folder: Path) -> int:
    if not folder.exists():
        return 0
    return len([f for f in folder.rglob("*.ts*") if ".test." not in f.name])


def scan_nav_permissions() -> list[tuple[str, str, str]]:
    nav_file = SRC / "components" / "layout" / "nav-config.ts"
    if not nav_file.exists():
        return []
    text = nav_file.read_text(encoding="utf-8", errors="replace")
    rows = []
    for m in re.finditer(
        r"to:\s*'([^']+)'.*?label:\s*'([^']+)'(?:.*?permission:\s*'([^']+)')?",
        text,
        re.DOTALL,
    ):
        rows.append((m.group(1), m.group(2), m.group(3) or "—"))
    return rows


def main() -> None:
    today = date.today().isoformat()

    feature_count = len(list((SRC / "features").iterdir())) if (SRC / "features").exists() else 0
    ui_count = count_ts_files(SRC / "components" / "ui")
    layout_count = count_ts_files(SRC / "components" / "layout")
    lib_count = count_ts_files(SRC / "lib")
    test_count = len(list(SRC.rglob("*.test.ts*")))
    e2e_count = len(list((FRONTEND / "e2e").glob("*.spec.ts")))
    i18n_files = len(list((SRC / "i18n" / "locales" / "en").glob("*.json")))

    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\nเอกสาร Frontend Design ฉบับละเอียด\n(React · Vite · TypeScript · Tailwind)"
    )
    run.bold = True
    run.font.size = Pt(20)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"วันที่จัดทำ: {today}\nเวอร์ชันเอกสาร: 1.0\n"
        f"แพ็กเกจ: PM-Pepsi-App/frontend"
    ).font.size = Pt(11)

    doc.add_paragraph()
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายสถาปัตยกรรม frontend ของ PM-Pepsi-App "
        "รวมเทคโนโลยี โครงสร้างโฟลเดอร์ routing navigation RBAC "
        "i18n state management UI components และแนวทาง build/deploy"
    )

    doc.add_page_break()

    # 1 Overview
    doc.add_heading("1. ภาพรวมสถาปัตยกรรม", level=1)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Framework", "React 19 + TypeScript"],
            ["Build tool", "Vite 8"],
            ["Routing", "React Router DOM 7"],
            ["Styling", "Tailwind CSS 4 + CVA"],
            ["UI primitives", "Radix UI + shadcn-style components"],
            ["Server state", "TanStack React Query v5"],
            ["Forms", "react-hook-form + Zod resolvers"],
            ["i18n", "i18next + react-i18next (EN default, TH toggle)"],
            ["Charts", "Chart.js + react-chartjs-2"],
            ["Calendar", "FullCalendar 6"],
            ["Dev port", "5173 (proxy /api → :4000)"],
            ["Routes", str(len(ROUTE_PAGES))],
            ["Feature modules", str(feature_count)],
            ["UI components", str(ui_count)],
            ["Lib utilities", str(lib_count)],
            ["Unit tests", str(test_count)],
            ["E2E specs", str(e2e_count)],
            ["i18n namespaces", str(len(I18N_NS))],
        ],
    )

    doc.add_heading("1.1 Layered Architecture", level=2)
    add_table(
        doc,
        ["ชั้น", "โฟลเดอร์", "หน้าที่"],
        [
            ["Entry", "src/main.tsx", "Providers, QueryClient, BrowserRouter"],
            ["App shell", "src/App.tsx", "Route tree + guards"],
            ["Features", "src/features/*/", "หน้าและ logic ตามโดเมน"],
            ["Components", "src/components/", "UI ใช้ซ้ำ (layout, scheduling, confirmation)"],
            ["Lib", "src/lib/", "API client, RBAC, nav, theme helpers"],
            ["Providers", "src/providers/", "Theme, Settings, I18n"],
            ["i18n", "src/i18n/locales/", "EN/TH JSON namespaces"],
            ["E2E", "e2e/*.spec.ts", "Playwright tests"],
        ],
    )

    doc.add_page_break()

    # 2 Provider tree
    doc.add_heading("2. Application Bootstrap & Providers", level=2)
    tree = [
        "StrictMode",
        "  QueryClientProvider (staleTime 60s, retry 1)",
        "    BrowserRouter",
        "      I18nProvider",
        "        SettingsProvider (public settings + branding)",
        "          ThemeProvider (light/dark/system)",
        "            AppErrorBoundary",
        "              App (Routes)",
        "            AppToaster (sonner)",
    ]
    p = doc.add_paragraph()
    r = p.add_run("\n".join(tree))
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.add_heading("2.1 Vite Proxy", level=2)
    doc.add_paragraph(
        "Development: VITE_API_URL ว่าง = same-origin — Vite proxy ส่ง /api ไป backend :4000. "
        "Production: IIS/nginx เสิร์ฟ static + reverse proxy /api"
    )

    doc.add_page_break()

    # 3 Routing
    doc.add_heading("3. Routing & Layout Guards", level=1)

    doc.add_heading("3.1 Guard hierarchy", level=2)
    add_bullets(
        doc,
        [
            "GuestOnly — /login เท่านั้น (redirect ถ้า login แล้ว)",
            "RequireAuth — ต้องมี session",
            "NavRouteGuard — ตรวจ permission ต่อ path (nav-route-permissions)",
            "RequireRole — admin zone (legacy + RBAC)",
            "AdminLayout — sidebar admin แยกจาก AppShell",
        ],
    )

    doc.add_heading("3.2 ตาราง Routes ทั้งหมด", level=2)
    add_table(
        doc,
        ["Path", "Component", "Layout", "คำอธิบาย"],
        [[p, c, l, d] for p, c, l, d in ROUTE_PAGES],
    )

    doc.add_page_break()

    # 4 Navigation
    doc.add_heading("4. Navigation & Sidebar", level=1)
    doc.add_paragraph(
        "เมนูหลักจาก nav-config.ts (fallback) หรือ API GET /api/v1/nav/menu — "
        "กรองด้วย filterNavForUser() ตาม user.permissions"
    )

    nav_rows = scan_nav_permissions()
    if nav_rows:
        doc.add_heading("4.1 เมนูหลัก (nav-config)", level=2)
        add_table(doc, ["Path", "Label", "Permission"], nav_rows[:25])

    doc.add_heading("4.2 Sidebar modes", level=2)
    add_bullets(
        doc,
        [
            "Expanded — แสดง label เต็ม (desktop)",
            "Rail — แค่ icon (tablet)",
            "Drawer — mobile < lg, ปิดอัตโนมัติเมื่อเปลี่ยนหน้า",
            "ค่าตั้ง: tbl_user_pref sidebar_mode — sync กับ Settings",
        ],
    )

    doc.add_heading("4.3 Command Palette", level=2)
    doc.add_paragraph("Ctrl+K / Cmd+K — cmdk component ค้นหาเมนูและ admin actions")

    doc.add_page_break()

    # 5 RBAC
    doc.add_heading("5. RBAC บน Frontend", level=1)
    add_table(
        doc,
        ["กลไก", "ไฟล์", "หน้าที่"],
        [
            ["usePermission(perm)", "lib/use-permission.ts", "hook ตรวจสิทธิ์"],
            ["CanPermission", "component wrapper", "ซ่อนปุ่ม/action"],
            ["nav-route-permissions", "lib/nav-route-permissions.ts", "map path → permission"],
            ["filterNavForUser", "lib/use-app-nav.ts", "กรอง sidebar"],
            ["/auth/me permissions[]", "จาก backend", "โหลดตอน login"],
        ],
    )

    doc.add_page_break()

    # 6 Feature modules
    doc.add_heading("6. Feature Modules", level=1)
    add_table(
        doc,
        ["โฟลเดอร์ features/", "คำอธิบาย"],
        FEATURE_MODULES,
    )

    doc.add_heading("6.1 Shared Components สำคัญ", level=2)
    add_table(
        doc,
        ["โฟลเดอร์", "ตัวอย่าง", "ใช้กับ"],
        [
            ["components/layout/", "AppShell, Sidebar, TopBar", "Shell ทุกหน้า"],
            ["components/scheduling/", "WorkOrderDetailDialog, MonthFullCalendar", "Calendar, WO"],
            ["components/confirmation/", "ConfirmQcPanel, MassConfirmBar", "Confirmation"],
            ["components/planning/", "PlanningAssignDialog", "Planning"],
            ["components/work-orders/", "WoConfirmationTable", "WO list"],
            ["components/ui/", "Button, Dialog, EmptyState, Table", "Design system"],
            ["components/admin/", "AdminTour, ConfirmPhraseDialog", "Admin zone"],
        ],
    )

    doc.add_page_break()

    # 7 State & API
    doc.add_heading("7. Data Fetching & API Client", level=1)

    doc.add_heading("7.1 React Query patterns", level=2)
    add_bullets(
        doc,
        [
            "useQuery — โหลดตาราง, KPI, calendar events",
            "useMutation — save, import, assign, confirm",
            "placeholderData: keepPreviousData — ตาราง/filter หนักไม่กระพริบ",
            "queryKey แยกตาม filter params",
            "invalidateQueries หลัง mutation สำเร็จ",
        ],
    )

    doc.add_heading("7.2 API client", level=2)
    add_bullets(
        doc,
        [
            "lib/fetch-api.ts — wrapper fetch + credentials include",
            "Bearer token จาก session หรือ cookie same-origin",
            "MaintenanceModeError — 503 MAINTENANCE แสดง banner",
            "AuthenticatedImage — โหลดรูป API ด้วย Bearer + blob URL",
            "query-load-error.tsx — UI error มาตรฐาน + retry",
        ],
    )

    doc.add_page_break()

    # 8 i18n & Theme
    doc.add_heading("8. Internationalization (i18n)", level=1)
    doc.add_paragraph(
        f"ค่าเริ่มต้น UI: ภาษาอังกฤษ — สลับ EN/TH ที่ TopBar และ Settings → Profile. "
        f"ไฟล์ locale: src/i18n/locales/{{en,th}}/*.json ({i18n_files} namespaces)"
    )
    add_table(doc, ["Namespace", "ใช้กับ"], [[ns, ""] for ns in I18N_NS])

    doc.add_heading("9. Theme & Branding", level=1)
    add_bullets(
        doc,
        [
            "ThemeProvider — light / dark / system จาก tbl_setting",
            "CSS variables — Pepsi red primary, system blue accent",
            "Admin branding — อัปโหลด logo/favicon → API → SettingsProvider",
            "Login page — พื้นหลังและโลโก้จาก public settings",
            "fullcalendar.css — สไตล์ปฏิทินแยก",
        ],
    )

    doc.add_page_break()

    # 10 Key UX patterns
    doc.add_heading("10. UX Patterns สำคัญ", level=1)
    add_table(
        doc,
        ["Pattern", "รายละเอียด"],
        [
            ["AppCard", "กล่องเนื้อหาหลักทุกหน้า"],
            ["EmptyState", "ไม่มีข้อมูล / ไม่มีสิทธิ์ / โหลดไม่สำเร็จ"],
            ["WorkOrderDetailDialog", "Modal WO หลายแท็บ — lazy load รูป"],
            ["Sticky table header", "ตารางยาวเลื่อนได้"],
            ["Mobile calendar drag", "long-press ~0.4s ก่อนลาก event"],
            ["Mass confirm ≤44", "ข้อจำกัด SAP ต่อชุด"],
            ["Joyride tour", "Admin tour + dashboard onboarding"],
            ["AppErrorBoundary", "จับ React crash — แสดง UnexpectedErrorPage"],
        ],
    )

    doc.add_heading("11. การทดสอบ Frontend", level=1)
    add_table(
        doc,
        ["ชั้น", "เครื่องมือ", "คำสั่ง"],
        [
            ["Unit / component", "Vitest + Testing Library", "npm test"],
            ["E2E smoke", "Playwright", "npm run test:e2e:smoke"],
            ["E2E all routes", "Playwright 76 cases", "npm run test:e2e:all-routes"],
            ["E2E viewport/locale", "Playwright", "test:e2e:viewport, locale"],
            ["UI audit", "audit-hardcode-mock.mjs", "npm run audit:ui"],
        ],
    )

    doc.add_heading("12. Build & Deploy", level=1)
    add_bullets(
        doc,
        [
            "npm run build — tsc -b && vite build → dist/",
            "Static files เสิร์ฟจาก IIS/nginx",
            "/api proxy ไป Express backend :4000",
            "Environment: VITE_API_URL (ว่าง = same-origin)",
            "Playwright: E2E_USE_DEV_SEED=1, PLAYWRIGHT_SKIP_WEBSERVER=1",
        ],
    )

    doc.add_heading("13. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "PM-Pepsi-App/frontend/src/App.tsx — route tree",
            "PM-Pepsi-App/frontend/src/components/layout/nav-config.ts",
            "docs/USER-MANUAL-TH.md — คู่มือผู้ใช้ทุกหน้า",
            "docs/customer-requirements/BACKEND-DESIGN-DETAILED-TH.docx",
            "docs/customer-requirements/UI-POLISH-PHASES.md",
            "from customer/tools/generate_frontend_design_docx.py",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Routes: {len(ROUTE_PAGES)}, Features: {feature_count}")


if __name__ == "__main__":
    main()
