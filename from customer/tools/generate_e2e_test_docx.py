#!/usr/bin/env python3
"""Generate PM Pepsi App — detailed E2E Test documentation (Word .docx)."""
from __future__ import annotations

import re
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
E2E_DIR = ROOT / "PM-Pepsi-App" / "frontend" / "e2e"
OUT = ROOT / "docs" / "customer-requirements" / "E2E-TEST-DETAILED-TH.docx"

DESCRIBE_RE = re.compile(
    r"""test\.describe(?:\.skip)?\s*\(\s*['"`]([^'"`]+)['"`]""",
)
TEST_RE = re.compile(
    r"""^\s*test(?:\.skip)?\s*\(\s*['"`]([^'"`]+)['"`]""",
    re.MULTILINE,
)
ROUTE_RE = re.compile(
    r"""\{\s*path:\s*['"`]([^'"`]+)['"`],\s*label:\s*['"`]([^'"`]+)['"`]\s*\}""",
)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "548235",
    font_size: int = 8,
    header_font_white: bool = True,
) -> None:
    from docx.shared import RGBColor

    if not rows:
        doc.add_paragraph("(ไม่มีข้อมูล)")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                if header_font_white:
                    r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_spec_file(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    describes = DESCRIBE_RE.findall(text)
    tests = TEST_RE.findall(text)
    dynamic_routes = len(ROUTE_RE.findall(text))
    locales_loop = 2 if "for (const locale of" in text else 1
    hot_path_loop = "HOT_PATH_PAGES" in text
    all_routes_loop = "ALL_APP_ROUTES" in text
    return {
        "file": path.name,
        "describes": describes,
        "static_tests": tests,
        "dynamic_note": (
            f"× {dynamic_routes or 'N'} routes × {locales_loop} locale(s)"
            if all_routes_loop or hot_path_loop
            else ""
        ),
        "uses_all_routes": all_routes_loop,
        "uses_hot_path": hot_path_loop,
    }


def load_all_app_routes() -> list[tuple[str, str]]:
    path = E2E_DIR / "helpers" / "all-app-routes.ts"
    text = path.read_text(encoding="utf-8")
    return ROUTE_RE.findall(text)


def count_playwright_tests() -> tuple[int, str]:
    frontend = ROOT / "PM-Pepsi-App" / "frontend"
    env = {**dict(__import__("os").environ), "E2E_USE_DEV_SEED": "1"}
    try:
        proc = subprocess.run(
            ["npx", "playwright", "test", "--list"],
            cwd=frontend,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
            env=env,
            shell=True,
        )
        out = (proc.stdout or "") + (proc.stderr or "")
        m = re.search(r"Total:\s*(\d+)\s*tests", out)
        if m:
            return int(m.group(1)), out.strip().splitlines()[-1] if out else ""
    except (subprocess.TimeoutExpired, OSError, FileNotFoundError):
        pass
    # Fallback: 38 routes × 2 + static suites
    routes = len(load_all_app_routes())
    return routes * 2 + 58, "(estimated — run playwright test --list locally)"


def main() -> None:
    today = date.today().isoformat()
    total_tests, total_line = count_playwright_tests()
    all_routes = load_all_app_routes()
    specs = sorted(E2E_DIR.glob("*.spec.ts"), key=lambda p: p.name)

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "PM Pepsi App\n"
        "เอกสาร E2E / System Test ฉบับละเอียด\n"
        "(Playwright · End-to-End · UI Regression)"
    )
    run.bold = True
    run.font.size = Pt(18)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"โครงการ: PM Dashboard & Monitoring — Pepsi-Cola (Thailand)\n"
        f"ผู้พัฒนา: S.Y. Interactive Development Limited\n"
        f"รหัสเอกสาร: PM-PEPSI-E2E-001\n"
        f"วันที่: {today} · เวอร์ชัน: 1.0\n"
        f"เครื่องมือ: Playwright {total_line or f'Total ~{total_tests} tests'}"
    ).font.size = Pt(10)

    doc.add_page_break()

    doc.add_heading("การควบคุมเอกสาร", level=1)
    add_table(
        doc,
        ["รายการ", "ค่า"],
        [
            ["ชื่อเอกสาร", "E2E / System Test Specification"],
            ["รหัส", "PM-PEPSI-E2E-001"],
            ["เวอร์ชัน", "1.0"],
            ["สถานะ", "Draft — Pre-UAT / ส่งมอบ"],
            ["คู่กับ", "UNIT-TEST-DETAILED-TH.docx · STRESS-TEST-DETAILED-TH.docx"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("1. บทนำ", level=1)
    doc.add_paragraph(
        "เอกสารฉบับนี้อธิบายการทดสอบ End-to-End (E2E) ของ PM-Pepsi-App ด้วย Playwright "
        "ครอบคลุม API smoke, UI smoke, regression ทุก route, viewport, locale และ portal — "
        "ใช้เป็นหลักฐาน V&V ก่อน UAT และ go-live ตาม ISO/IEC 29110"
    )

    doc.add_heading("1.1 ความสัมพันธ์ชั้นการทดสอบ", level=2)
    add_table(
        doc,
        ["ชั้น", "เครื่องมือ", "จำนวน (โดยประมาณ)", "วัตถุประสงค์"],
        [
            ["Unit", "Vitest (BE/FE)", "458+ cases", "Logic · parser · schema · component"],
            ["Integration API", "Supertest + Vitest", "รวมใน unit BE", "Route + DB mock"],
            ["E2E / System", "Playwright", f"{total_tests} tests", "Browser + API จริง · RBAC · UI"],
            ["Stress", "stress-api.ts", "ตาม script", "โหลด API · p95 latency"],
            ["UAT", "ชีตลูกค้า", "UAT-ROUND-3", "ยอมรับธุรกิจ"],
        ],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_heading("1.2 ขอบเขต E2E", level=2)
    add_bullets(
        doc,
        [
            "In scope: Chromium desktop (Playwright default project)",
            "In scope: Login session · ทุก authenticated route (admin seed)",
            "In scope: Console error audit · AppErrorBoundary · viewport 1280/1920/tablet",
            "In scope: EN/TH locale บน hot path",
            "Out of scope: ทดสอบ SAP file import จริงทุกรูปแบบ (ใช้ UAT มือ)",
            "Out of scope: Telegram Bot ใน browser (ใช้ UAT §F / webhook แยก)",
            "Out of scope: Mobile native apps",
        ],
    )

    doc.add_page_break()

    doc.add_heading("2. สภาพแวดล้อมทดสอบ", level=1)

    doc.add_heading("2.1 ความต้องการ", level=2)
    add_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Node.js", "20 LTS+"],
            ["PostgreSQL", "14+ · schema app · migrations ครบ"],
            ["Seed", "database/seeds/009_dev_auth_seed.sql → ADMIN01/admin"],
            ["Backend", ":4000 — npm run dev"],
            ["Frontend", ":5173 — npm run dev (proxy /api)"],
            ["Playwright", "@playwright/test ^1.60 — npx playwright install chromium"],
        ],
        header_fill="BDD7EE",
        header_font_white=False,
    )

    doc.add_heading("2.2 ตัวแปร Environment", level=2)
    add_table(
        doc,
        ["ตัวแปร", "ค่าแนะนำ", "หมายเหตุ"],
        [
            ["E2E_USE_DEV_SEED", "1", "ใช้ ADMIN01/admin อัตโนมัติ"],
            ["E2E_ADMIN_USER / PASSWORD", "ADMIN01 / admin", "ทางเลือก — ดู e2e/.env.example"],
            ["PLAYWRIGHT_BASE_URL", "http://localhost:5173", "Windows: ใช้ localhost ไม่ใช่ 127.0.0.1"],
            ["PLAYWRIGHT_API_URL", "http://127.0.0.1:4000", "API smoke / session seed"],
            ["PLAYWRIGHT_SKIP_WEBSERVER", "1", "เมื่อรัน dev server เองอยู่แล้ว"],
            ["CAPTURE_UAT_SCREENSHOTS", "1", "เปิด uat-screenshots.spec.ts"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("2.3 คำสั่งรัน (npm scripts)", level=2)
    add_table(
        doc,
        ["Script", "ไฟล์ / ขอบเขต", "ใช้เมื่อ"],
        [
            ["npm run test:e2e", "ทุก spec ใน e2e/", "Full regression"],
            ["npm run test:e2e:smoke", "api + app + confirmation-export", "CI เร็ว (~15 tests)"],
            ["npm run test:e2e:all-routes", "u4e-all-routes-console", "U4 Gate — 76 routes×locale"],
            ["npm run test:e2e:console", "u4e-console-pages", "หน้า 1–7 hot path"],
            ["npm run test:e2e:viewport", "u4e-viewport-pages", "1280/1920/tablet"],
            ["npm run test:e2e:locale", "u4e-locale-shell", "EN↔TH"],
            ["npm run test:e2e:portal", "u4f-portal", "Portal flow"],
            ["npm run test:e2e:admin-smoke", "admin-hub-smoke", "Admin hub"],
            ["npm run test:e2e:screenshots", "uat-screenshots", "PNG สำหรับเอกสาร UAT"],
            ["npm run test:e2e:ui", "Playwright UI mode", "Debug"],
        ],
        header_fill="548235",
    )

    add_code(
        doc,
        "# Terminal 1\n"
        "cd PM-Pepsi-App/backend && npm run dev\n\n"
        "# Terminal 2\n"
        "cd PM-Pepsi-App/frontend && npm run dev\n\n"
        "# Terminal 3\n"
        "cd PM-Pepsi-App/frontend\n"
        "$env:E2E_USE_DEV_SEED=\"1\"\n"
        "$env:PLAYWRIGHT_SKIP_WEBSERVER=\"1\"\n"
        "npm run test:e2e:smoke",
    )

    doc.add_page_break()

    doc.add_heading("3. สรุปชุดทดสอบ (Test Suites)", level=1)
    doc.add_paragraph(
        f"รวม {len(specs)} ไฟล์ spec · {total_tests} test cases (playwright test --list)"
    )

    suite_rows = []
    suite_meta = {
        "api-smoke.spec.ts": ("API", "4", "Health · public settings · login 401 · export 401"),
        "confirmation-export-api.spec.ts": ("API RBAC", "1", "Export scope ALL/OWN จาก permission"),
        "app-smoke.spec.ts": ("UI Smoke", "7", "Login · redirect · dashboard · WO · confirm · settings · roles"),
        "admin-tour.spec.ts": ("Admin UX", "3", "Joyride tour · tooltip · navigation"),
        "admin-hub-smoke.spec.ts": ("Admin", "1", "Console → roles → menu"),
        "u4e-console-pages.spec.ts": ("U4e Console", "7", "Hot path 1–7 + WO modal — no console errors"),
        "u4e-all-routes-console.spec.ts": (
            "U4e Full",
            f"{len(all_routes) * 2}",
            f"ทุก route × EN/TH — no crash/console",
        ),
        "u4e-viewport-pages.spec.ts": ("U4e Viewport", "~25", "1280×720 · 1920×1080 · tablet drawer"),
        "u4e-locale-shell.spec.ts": ("U4e i18n", "~8", "EN→TH hot path + sidebar"),
        "u4f-portal.spec.ts": ("U4f Portal", "2", "Login → portal → PM app"),
        "uat-screenshots.spec.ts": ("UAT Assets", "3", "PNG เมื่อ CAPTURE_UAT_SCREENSHOTS=1"),
    }
    for spec in specs:
        name = spec.name
        cat, count, desc = suite_meta.get(name, ("Other", "?", ""))
        suite_rows.append([name, cat, count, desc])

    add_table(
        doc,
        ["ไฟล์", "หมวด", "จำนวน", "คำอธิบาย"],
        suite_rows,
        header_fill="548235",
    )

    doc.add_page_break()

    doc.add_heading("4. รายละเอียดตามไฟล์", level=1)

    for spec in specs:
        info = parse_spec_file(spec)
        doc.add_heading(f"4.x {info['file']}", level=2)
        if info["describes"]:
            doc.add_paragraph("Describe blocks: " + " · ".join(info["describes"][:4]))
        if info["static_tests"]:
            add_table(
                doc,
                ["#", "Test case (static)"],
                [[str(i + 1), t] for i, t in enumerate(info["static_tests"][:20])],
                header_fill="D9E2F3",
                header_font_white=False,
            )
            if len(info["static_tests"]) > 20:
                doc.add_paragraph(f"... และอีก {len(info['static_tests']) - 20} static tests")
        if info["uses_all_routes"]:
            doc.add_paragraph(
                f"Dynamic: 1 test ต่อ route ใน ALL_APP_ROUTES ({len(all_routes)} routes) "
                f"× 2 locales (en, th) = {len(all_routes) * 2} tests"
            )
        if info["uses_hot_path"]:
            doc.add_paragraph(
                "Dynamic: 1 test ต่อ HOT_PATH_PAGES (6 หน้า) + viewport variants"
            )

    doc.add_page_break()

    doc.add_heading("5. รายการ Route ทั้งหมด (ALL_APP_ROUTES)", level=1)
    doc.add_paragraph(
        f"ใช้ใน u4e-all-routes-console.spec.ts — {len(all_routes)} routes · ทดสอบ EN + TH"
    )
    add_table(
        doc,
        ["#", "Label", "Path"],
        [[str(i + 1), label, path] for i, (path, label) in enumerate(all_routes)],
        header_fill="E2EFDA",
        header_font_white=False,
    )

    doc.add_page_break()

    doc.add_heading("6. เกณฑ์ผ่าน (Pass Criteria)", level=1)
    add_table(
        doc,
        ["Gate", "คำสั่ง", "เกณฑ์ผ่าน"],
        [
            ["Smoke (รวดเร็ว)", "npm run test:e2e:smoke", "100% passed"],
            ["U4 Console", "npm run test:e2e:console", "0 console errors · no ErrorBoundary"],
            ["U4 All routes", "npm run test:e2e:all-routes", f"{len(all_routes)*2}/76 passed"],
            ["U4 Viewport", "npm run test:e2e:viewport", "layout OK 1280/1920/tablet"],
            ["U4 Locale", "npm run test:e2e:locale", "EN/TH switch ไม่ crash"],
            ["U4 Portal", "npm run test:e2e:portal", "portal → PM module"],
            ["Full E2E", "npm run test:e2e", f"{total_tests}/{total_tests} passed"],
        ],
        header_fill="2F5496",
    )

    doc.add_heading("6.1 สิ่งที่ E2E ตรวจจับ", level=2)
    add_bullets(
        doc,
        [
            "HTTP 5xx / unhandled promise rejection → console collector",
            "AppErrorBoundary (Try again / ลองแสดงผลอีกครั้ง)",
            "main content visible (main, [role=main], .engineering-board)",
            "Public /board ไม่ต้อง login",
            "WO modal test — skip ถ้า DB ไม่มีแถว WO",
            "Screenshot/video/trace on failure (Playwright config)",
        ],
    )

    doc.add_heading("6.2 Pre-UAT Master (P7.1)", level=2)
    add_bullets(
        doc,
        [
            "npm test (backend + frontend unit)",
            "npm run build (frontend)",
            "npm run test:e2e:smoke (เมื่อ env พร้อม)",
            "npm run audit:ui — ไม่มี (mock) ใน production UI",
        ],
    )

    doc.add_page_break()

    doc.add_heading("7. Helpers & โครงสร้าง e2e/", level=1)
    add_table(
        doc,
        ["ไฟล์", "หน้าที่"],
        [
            ["helpers/auth.ts", "seedAdminSession · login API · cookie"],
            ["helpers/env.ts", "โหลด e2e/.env"],
            ["helpers/console-errors.ts", "collect pageerror + console.error"],
            ["helpers/all-app-routes.ts", "38 authenticated routes"],
            ["helpers/hot-path-pages.ts", "6 หน้า U4e priority"],
            ["helpers/locale-switch.ts", "สลับ EN/TH ใน test"],
            ["helpers/viewport.ts", "ขนาดจอมาตรฐาน"],
            ["helpers/screenshot-dir.ts", "docs/.../screenshots/u4*"],
        ],
        header_fill="D9E2F3",
        header_font_white=False,
    )

    doc.add_heading("8. แก้ปัญหาที่พบบ่อย", level=1)
    add_table(
        doc,
        ["อาการ", "สาเหตุ", "แก้"],
        [
            ["ECONNREFUSED :5173", "Frontend ไม่รัน", "npm run dev frontend"],
            ["Login skip ทั้งชุด", "ไม่มี E2E_USE_DEV_SEED", "ตั้ง env + seed 009"],
            ["health db fail", "PG/migration", "run-all-migrations.ps1"],
            ["localhost vs 127.0.0.1", "Vite bind", "PLAYWRIGHT_BASE_URL=localhost:5173"],
            ["WO modal skip", "DB ว่าง", "import IW37N หรือ seed demo"],
            ["Schema not ready", "migration ขาด", "รัน migration ล่าสุด"],
            ["Timeout 90s", "หน้าช้า/ API", "ตรวจ backend log · เพิ่ม timeout ชั่วคราว"],
        ],
        header_fill="FFF2CC",
        header_font_white=False,
    )

    doc.add_heading("9. รายงานผล", level=1)
    add_bullets(
        doc,
        [
            "Terminal: list reporter — passed/failed/skipped",
            "HTML report: npx playwright show-report (หลังรัน)",
            "Trace: on-first-retry — เปิดจาก report",
            "Screenshot: only-on-failure",
            "บันทึกในเอกสารส่งมอบ: วันที่ · commit · คำสั่ง · X/Y passed",
        ],
    )

    doc.add_heading("10. เอกสารอ้างอิง", level=1)
    add_bullets(
        doc,
        [
            "docs/customer-requirements/E2E-SMOKE.md",
            "docs/customer-requirements/PRE-UAT-UI-PHASES.md §U4e/U4f",
            "docs/customer-requirements/UNIT-TEST-DETAILED-TH.docx",
            "docs/customer-requirements/STRESS-TEST-DETAILED-TH.docx",
            "PM-Pepsi-App/frontend/playwright.config.ts",
            "docs/customer-requirements/UAT-ROUND-3-TH.md",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph("— จบเอกสาร E2E Test —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT} ({total_tests} tests catalogued)")


if __name__ == "__main__":
    main()
